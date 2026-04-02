"""Render DOCX files to PDF and page PNGs for visual review."""

from __future__ import annotations

import argparse
import json
import os
from pathlib import Path
import shutil
import subprocess
import sys
from typing import Any, Callable

from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTDIR = REPO_ROOT / "tmp" / "docs" / "render_docx"
SOFFICE_ENV = "LEGALPDF_SOFFICE_PATH"
POPPLER_ENV = "LEGALPDF_POPPLER_BIN_DIR"
_BENIGN_SOFFICE_STDERR_LINES = {
    "Could not find platform independent libraries <prefix>",
}


def _existing_path(raw_value: str) -> Path | None:
    candidate = str(raw_value or "").strip()
    if not candidate:
        return None
    path = Path(candidate).expanduser()
    if path.exists():
        return path.resolve()
    return None


def _prefer_soffice_cli(path: Path | None) -> Path | None:
    if path is None:
        return None
    resolved = path.expanduser().resolve()
    if resolved.name.lower() == "soffice.exe":
        cli_path = resolved.with_name("soffice.com")
        if cli_path.exists():
            return cli_path.resolve()
    return resolved


def _resolve_soffice_path(*, environment: dict[str, str] | None = None) -> Path | None:
    env = environment or os.environ
    override = _existing_path(env.get(SOFFICE_ENV, ""))
    if override and override.is_file():
        return _prefer_soffice_cli(override)
    discovered = shutil.which("soffice.com") or shutil.which("soffice.exe") or shutil.which("soffice")
    if discovered:
        return _prefer_soffice_cli(Path(discovered))

    local_appdata = str(env.get("LOCALAPPDATA", "") or "").strip()
    program_files = str(env.get("ProgramFiles", "") or "").strip()
    program_files_x86 = str(env.get("ProgramFiles(x86)", "") or "").strip()
    candidates = [
        Path(local_appdata) / "Programs" / "LibreOffice" / "program" / "soffice.com",
        Path(local_appdata) / "Programs" / "LibreOffice" / "program" / "soffice.exe",
        Path(program_files) / "LibreOffice" / "program" / "soffice.com",
        Path(program_files) / "LibreOffice" / "program" / "soffice.exe",
        Path(program_files_x86) / "LibreOffice" / "program" / "soffice.com",
        Path(program_files_x86) / "LibreOffice" / "program" / "soffice.exe",
        Path(r"C:\Program Files\LibreOffice\program\soffice.com"),
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.com"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return _prefer_soffice_cli(candidate)
    return None


def _resolve_pdftoppm_path(*, environment: dict[str, str] | None = None) -> Path | None:
    env = environment or os.environ
    override = _existing_path(env.get(POPPLER_ENV, ""))
    if override:
        if override.is_file():
            return override
        candidate = override / "pdftoppm.exe"
        if candidate.exists():
            return candidate.resolve()
        candidate = override / "pdftoppm"
        if candidate.exists():
            return candidate.resolve()

    discovered = shutil.which("pdftoppm.exe") or shutil.which("pdftoppm")
    if discovered:
        return Path(discovered).resolve()

    local_appdata = Path(str(env.get("LOCALAPPDATA", "") or "").strip() or ".")
    program_files = str(env.get("ProgramFiles", "") or "").strip()
    winget_links = local_appdata / "Microsoft" / "WinGet" / "Links" / "pdftoppm.exe"
    winget_package_candidates = sorted(
        (local_appdata / "Microsoft" / "WinGet" / "Packages").glob("oschwartz10612.Poppler*")
    )
    recursive_winget_candidates = [
        candidate.resolve()
        for package in winget_package_candidates
        for candidate in package.glob("**/pdftoppm.exe")
        if candidate.is_file()
    ]
    candidates = [
        winget_links,
        *recursive_winget_candidates,
        *(package / "Library" / "bin" / "pdftoppm.exe" for package in winget_package_candidates),
        *(package / "bin" / "pdftoppm.exe" for package in winget_package_candidates),
        Path(program_files) / "poppler" / "Library" / "bin" / "pdftoppm.exe",
        Path(program_files) / "Poppler" / "Library" / "bin" / "pdftoppm.exe",
        Path(program_files) / "poppler" / "bin" / "pdftoppm.exe",
        Path(program_files) / "Poppler" / "bin" / "pdftoppm.exe",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate.resolve()
    return None


def _load_pdf2image_convert() -> tuple[Callable[..., Any] | None, str]:
    try:
        from pdf2image import convert_from_path
    except Exception as exc:  # pragma: no cover - import error path is validated indirectly
        return None, str(exc)
    return convert_from_path, ""


def _build_docx_summary(docx_path: Path) -> dict[str, Any]:
    document = Document(str(docx_path))
    non_empty_paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return {
        "paragraph_count": len(document.paragraphs),
        "non_empty_paragraph_count": len(non_empty_paragraphs),
        "table_count": len(document.tables),
        "sample_text": non_empty_paragraphs[:5],
    }


def _normalize_soffice_stderr(stderr_text: str) -> tuple[str, list[str]]:
    lines = [line.strip() for line in str(stderr_text or "").splitlines() if line.strip()]
    ignored = [line for line in lines if line in _BENIGN_SOFFICE_STDERR_LINES]
    kept = [line for line in lines if line not in _BENIGN_SOFFICE_STDERR_LINES]
    return "\n".join(kept), ignored


def _convert_docx_to_pdf(
    *,
    input_path: Path,
    outdir: Path,
    soffice_path: Path,
) -> tuple[Path, dict[str, Any]]:
    outdir.mkdir(parents=True, exist_ok=True)
    output_pdf = outdir / f"{input_path.stem}.pdf"
    if output_pdf.exists():
        output_pdf.unlink()
    profile_dir = outdir / ".lo_profile"
    profile_dir.mkdir(parents=True, exist_ok=True)
    command = [
        str(soffice_path),
        f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(outdir),
        str(input_path),
    ]
    completed = subprocess.run(
        command,
        text=True,
        capture_output=True,
        check=False,
    )
    if completed.returncode != 0 or not output_pdf.exists():
        raise RuntimeError(str(completed.stderr or completed.stdout or "LibreOffice PDF conversion failed").strip())
    normalized_stderr, ignored_stderr = _normalize_soffice_stderr(completed.stderr)
    conversion: dict[str, Any] = {
        "command": command,
        "stdout": str(completed.stdout or "").strip(),
        "stderr": normalized_stderr,
    }
    if ignored_stderr:
        conversion["ignored_stderr"] = ignored_stderr
    return output_pdf, {
        **conversion,
    }


def _rasterize_pdf_to_pngs(
    *,
    pdf_path: Path,
    outdir: Path,
    pdftoppm_path: Path,
    convert_from_path: Callable[..., Any],
) -> list[Path]:
    outdir.mkdir(parents=True, exist_ok=True)
    poppler_dir = pdftoppm_path.parent
    images = convert_from_path(
        str(pdf_path),
        fmt="png",
        dpi=150,
        poppler_path=str(poppler_dir),
    )
    page_paths: list[Path] = []
    for index, image in enumerate(images, start=1):
        target = outdir / f"{pdf_path.stem}-{index:03d}.png"
        image.save(target, "PNG")
        page_paths.append(target.resolve())
    return page_paths


def render_docx(
    *,
    input_path: Path,
    outdir: Path,
    environment: dict[str, str] | None = None,
) -> dict[str, Any]:
    resolved_input = input_path.expanduser().resolve()
    resolved_outdir = outdir.expanduser().resolve()
    if not resolved_input.exists():
        raise FileNotFoundError(f"DOCX file not found: {resolved_input}")

    soffice_path = _resolve_soffice_path(environment=environment)
    pdftoppm_path = _resolve_pdftoppm_path(environment=environment)
    convert_from_path, pdf2image_error = _load_pdf2image_convert()
    docx_summary = _build_docx_summary(resolved_input)

    result: dict[str, Any] = {
        "status": "fallback",
        "input_docx": str(resolved_input),
        "outdir": str(resolved_outdir),
        "pdf_path": "",
        "page_image_paths": [],
        "page_count": 0,
        "soffice_path": str(soffice_path) if soffice_path else "",
        "pdftoppm_path": str(pdftoppm_path) if pdftoppm_path else "",
        "fallback_mode": "",
        "toolchain": {
            "soffice_available": soffice_path is not None,
            "pdftoppm_available": pdftoppm_path is not None,
            "pdf2image_available": convert_from_path is not None,
        },
        "docx_summary": docx_summary,
    }

    missing: list[str] = []
    if soffice_path is None:
        missing.append("soffice")
    if pdftoppm_path is None:
        missing.append("pdftoppm")
    if convert_from_path is None:
        missing.append("pdf2image")
        if pdf2image_error:
            result["toolchain"]["pdf2image_error"] = pdf2image_error
    if missing:
        result["fallback_mode"] = "structural_only_missing_" + "_".join(missing)
        result["message"] = (
            "Visual DOCX rendering is unavailable on this host; returning structural DOCX validation only."
        )
        summary_path = resolved_outdir / "render_docx.json"
        resolved_outdir.mkdir(parents=True, exist_ok=True)
        summary_path.write_text(json.dumps(result, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        result["summary_path"] = str(summary_path)
        return result

    try:
        pdf_path, conversion = _convert_docx_to_pdf(
            input_path=resolved_input,
            outdir=resolved_outdir,
            soffice_path=soffice_path,
        )
        page_paths = _rasterize_pdf_to_pngs(
            pdf_path=pdf_path,
            outdir=resolved_outdir,
            pdftoppm_path=pdftoppm_path,
            convert_from_path=convert_from_path,
        )
    except Exception as exc:
        result["fallback_mode"] = "structural_only_render_failed"
        result["message"] = str(exc)
        summary_path = resolved_outdir / "render_docx.json"
        resolved_outdir.mkdir(parents=True, exist_ok=True)
        summary_path.write_text(json.dumps(result, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        result["summary_path"] = str(summary_path)
        return result

    result.update(
        {
            "status": "ok",
            "pdf_path": str(pdf_path.resolve()),
            "page_image_paths": [str(path) for path in page_paths],
            "page_count": len(page_paths),
            "fallback_mode": "",
            "message": "DOCX rendered to PDF and PNG page images successfully.",
            "conversion": conversion,
        }
    )
    summary_path = resolved_outdir / "render_docx.json"
    summary_path.write_text(json.dumps(result, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    result["summary_path"] = str(summary_path)
    return result


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Render a DOCX file to PDF and PNG page images.")
    parser.add_argument("--input", required=True, type=Path, help="Path to the input DOCX file.")
    parser.add_argument(
        "--outdir",
        type=Path,
        default=DEFAULT_OUTDIR,
        help="Directory for the generated PDF, PNG pages, and JSON summary.",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(argv)
    result = render_docx(input_path=args.input, outdir=args.outdir)
    json.dump(result, sys.stdout, ensure_ascii=False, indent=2)
    sys.stdout.write("\n")
    if result["status"] == "ok":
        return 0
    return 2 if str(result.get("fallback_mode") or "").startswith("structural_only") else 1


if __name__ == "__main__":
    raise SystemExit(main())
