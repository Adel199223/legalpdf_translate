"""DOCX assembly utilities with atomic save semantics."""

from __future__ import annotations

import os
import hashlib
import json
import re
import unicodedata
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from statistics import median
from zipfile import ZipFile, ZipInfo

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .types import TargetLang

_LRM = "\u200e"
_RTL_LANG_BIDI_CODES = {
    "AR": "ar-SA",
    "HE": "he-IL",
    "FA": "fa-IR",
    "UR": "ur-PK",
}
_PLACEHOLDER_RE = re.compile(r"\[\[(.*?)\]\]", re.DOTALL)
_RTL_STRONG_BIDI_CATEGORIES = {"R", "AL"}
_LTR_STRONG_BIDI_CATEGORIES = {"L", "EN", "AN"}
_BIDI_CONTROL_CODEPOINTS = str.maketrans(
    "",
    "",
    "\u061c\u200e\u200f\u202a\u202b\u202c\u202d\u202e\u2066\u2067\u2068\u2069\ufeff",
)
_DOCX_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
_VISIBLE_WORD_RE = re.compile(r"\S+")
_PLACEHOLDER_TOKEN_SPAN_RE = re.compile(
    r"(?P<open>[\u2066\u2067\u2068]?)(?:\[\[(?P<token>.*?)\]\])(?P<close>\u2069?)",
    re.DOTALL,
)


@dataclass(frozen=True)
class _DocxVisibleContentSummary:
    paragraph_count: int
    visible_paragraph_count: int
    text_node_count: int
    text_char_count: int
    word_count: int


def sanitize_bidi_controls(text: str) -> str:
    if not text:
        return text
    return text.translate(_BIDI_CONTROL_CODEPOINTS)


def unwrap_internal_placeholders(text: str) -> str:
    if not text:
        return text
    return _PLACEHOLDER_RE.sub(lambda match: match.group(1), text)


def _is_rtl_target_lang(lang: TargetLang | str) -> bool:
    if isinstance(lang, TargetLang):
        code = lang.value
    else:
        code = str(lang)
    return code.strip().upper() in _RTL_LANG_BIDI_CODES


def _rtl_bidi_lang_code(lang: TargetLang | str) -> str:
    if isinstance(lang, TargetLang):
        code = lang.value
    else:
        code = str(lang)
    return _RTL_LANG_BIDI_CODES.get(code.strip().upper(), "ar-SA")


def _classify_directional_char(char: str) -> str:
    bidi = unicodedata.bidirectional(char)
    if bidi in _RTL_STRONG_BIDI_CATEGORIES:
        return "rtl"
    if bidi in _LTR_STRONG_BIDI_CATEGORIES:
        return "ltr"
    return "neutral"


def _nearest_strong_kind(segments: list[tuple[str, str]], index: int) -> str | None:
    for lookup in range(index - 1, -1, -1):
        kind = segments[lookup][0]
        if kind != "neutral":
            return kind
    for lookup in range(index + 1, len(segments)):
        kind = segments[lookup][0]
        if kind != "neutral":
            return kind
    return None


def _merge_directional_runs(segments: list[tuple[str, str]]) -> list[tuple[str, str]]:
    merged_segments: list[tuple[str, str]] = []
    for kind, chunk in segments:
        if not chunk:
            continue
        if merged_segments and merged_segments[-1][0] == kind:
            prev_kind, prev_chunk = merged_segments[-1]
            merged_segments[-1] = (prev_kind, f"{prev_chunk}{chunk}")
            continue
        merged_segments.append((kind, chunk))
    return merged_segments


def _rtl_boundary_space_runs(segments: list[tuple[str, str]]) -> list[tuple[str, str]]:
    """Keep separators in the paragraph direction, not in a Latin embedding.

    Word uses run-level rtl to resolve neutral characters too. An LTR run's
    trailing separator can therefore move to the wrong side of a name/number,
    visibly joining that token to the following Arabic word. Move only outer
    whitespace into RTL runs; internal spaces in a coherent Latin name/address
    stay LTR. Concatenation is byte-for-byte unchanged.
    """
    result: list[tuple[str, str]] = []
    for kind, chunk in segments:
        if kind != "ltr":
            result.append((kind, chunk))
            continue
        leading = len(chunk) - len(chunk.lstrip())
        end = len(chunk.rstrip())
        if not chunk.strip():
            result.append(("rtl", chunk))
            continue
        result.extend((("rtl", chunk[:leading]), ("ltr", chunk[leading:end]),
                       ("rtl", chunk[end:])))
    return _merge_directional_runs(result)


def _segment_directional_runs(text: str) -> tuple[list[tuple[str, str]], bool]:
    if not text:
        return [], False

    raw_segments: list[tuple[str, str]] = []
    current_kind = _classify_directional_char(text[0])
    current_chars = [text[0]]

    for char in text[1:]:
        kind = _classify_directional_char(char)
        if kind == current_kind:
            current_chars.append(char)
            continue
        raw_segments.append((current_kind, "".join(current_chars)))
        current_kind = kind
        current_chars = [char]
    raw_segments.append((current_kind, "".join(current_chars)))

    relabeled_segments: list[tuple[str, str]] = []
    for index, (kind, chunk) in enumerate(raw_segments):
        if kind != "neutral":
            relabeled_segments.append((kind, chunk))
            continue
        neighbor_kind = _nearest_strong_kind(raw_segments, index) or "rtl"
        relabeled_segments.append((neighbor_kind, chunk))

    merged_segments = _merge_directional_runs(relabeled_segments)
    has_rtl = any(kind == "rtl" for kind, _ in merged_segments)
    has_ltr = any(kind == "ltr" for kind, _ in merged_segments)
    return merged_segments, has_rtl and has_ltr


def _infer_atomic_run_kind(text: str) -> str:
    has_rtl = False
    has_ltr = False
    for char in text:
        kind = _classify_directional_char(char)
        if kind == "rtl":
            has_rtl = True
        elif kind == "ltr":
            has_ltr = True
    if has_rtl and not has_ltr:
        return "rtl"
    return "ltr"


def _segment_rtl_placeholder_aware_line(
    text: str,
    *,
    strip_bidi_controls: bool,
) -> tuple[list[tuple[str, str]], bool]:
    if not text:
        return [], False

    pieces: list[tuple[str, str]] = []
    cursor = 0
    for match in _PLACEHOLDER_TOKEN_SPAN_RE.finditer(text):
        if match.start() > cursor:
            plain = text[cursor : match.start()]
            if strip_bidi_controls:
                plain = sanitize_bidi_controls(plain)
            plain_runs, _ = _segment_directional_runs(plain)
            pieces.extend(plain_runs)

        token_core = match.group("token") or ""
        if strip_bidi_controls:
            token_text = sanitize_bidi_controls(token_core)
        else:
            token_text = f"{match.group('open') or ''}{token_core}{match.group('close') or ''}"
        if token_text:
            pieces.append((_infer_atomic_run_kind(token_core), token_text))
        cursor = match.end()

    if cursor < len(text):
        tail = text[cursor:]
        if strip_bidi_controls:
            tail = sanitize_bidi_controls(tail)
        tail_runs, _ = _segment_directional_runs(tail)
        pieces.extend(tail_runs)

    merged_segments = _merge_directional_runs(pieces)
    has_rtl = any(kind == "rtl" for kind, _ in merged_segments)
    has_ltr = any(kind == "ltr" for kind, _ in merged_segments)
    if has_rtl and has_ltr:
        merged_segments = _rtl_boundary_space_runs(merged_segments)
    return merged_segments, has_rtl and has_ltr


def _segment_rtl_placeholder_aware_runs(
    text: str,
    *,
    strip_bidi_controls: bool,
) -> tuple[list[tuple[str, str]], bool]:
    """Keep a complete Latin-only line in one directional sequence.

    Protected tokens are translation checks, not visual word boundaries. In an
    Arabic address paragraph, assigning the spaces between postcode and town to
    RTL reverses otherwise correct LTR tokens. Decide each explicit line before
    the mixed-language segmentation used for Arabic prose. No visible character
    or explicit line break is inserted, removed, or reordered.
    """
    # Older saved tokens may themselves span an explicit line break. Preserve
    # their existing atomic interpretation rather than exposing split [[...]]
    # delimiters while making the normal, single-line address decision below.
    if any("\n" in match.group("token") or "\r" in match.group("token")
           for match in _PLACEHOLDER_TOKEN_SPAN_RE.finditer(text)):
        return _segment_rtl_placeholder_aware_line(text, strip_bidi_controls=strip_bidi_controls)
    pieces: list[tuple[str, str]] = []
    for line in text.splitlines(keepends=True):
        visible = sanitize_bidi_controls(unwrap_internal_placeholders(line))
        kinds = {_classify_directional_char(char) for char in visible}
        if "ltr" in kinds and "rtl" not in kinds:
            unwrapped = unwrap_internal_placeholders(line)
            if strip_bidi_controls:
                unwrapped = sanitize_bidi_controls(unwrapped)
            pieces.append(("ltr", unwrapped))
        else:
            line_runs, _ = _segment_rtl_placeholder_aware_line(
                line, strip_bidi_controls=strip_bidi_controls,
            )
            pieces.extend(line_runs)
    kinds = {kind for kind, _ in pieces}
    return pieces, "rtl" in kinds and "ltr" in kinds


def _wrap_ltr_run_with_lrm(text: str) -> str:
    if not text.strip():
        return text
    if text.startswith(_LRM) and text.endswith(_LRM):
        return text
    leading_ws = len(text) - len(text.lstrip())
    trailing_ws = len(text) - len(text.rstrip())
    core_end = len(text) - trailing_ws if trailing_ws else len(text)
    core = text[leading_ws:core_end]
    if not core:
        return text
    return f"{text[:leading_ws]}{_LRM}{core}{_LRM}{text[core_end:]}"


def _add_rtl_flags(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is not None:
        p_pr.remove(bidi)
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.insert_element_before(
        bidi, "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
        "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
        "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
        "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
    )
    # rtl belongs in run properties, never in pPr.
    for invalid in list(p_pr.findall(qn("w:rtl"))):
        p_pr.remove(invalid)
    jc = p_pr.get_or_add_jc()
    if jc.get(qn("w:val")) != "center":
        jc.set(qn("w:val"), "start")


def _set_rtl_visual_alignment(paragraph, alignment: str) -> None:
    # Modern Word supports logical start/end. Legacy left/right values are
    # interpreted relative to bidi in Word's compatibility layout, so "right"
    # can render physically left after compatibilityMode is removed. Logical
    # values avoid that inversion without changing the document compatibility
    # contract. python-docx's older alignment enum lacks start/end: write XML.
    value = {"right": "start", "left": "end", "center": "center", "justify": "both"}[alignment]
    paragraph._p.get_or_add_pPr().get_or_add_jc().set(qn("w:val"), value)


def _set_rtl_run_props(run, *, bidi_lang: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "1")

    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), bidi_lang)
    lang.set(qn("w:bidi"), bidi_lang)


def _set_ltr_run_props(run, *, lang_code: str = "en-US") -> None:
    r_pr = run._r.get_or_add_rPr()
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "0")

    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), lang_code)


def _verify_non_empty_file(path: Path) -> None:
    if not path.exists():
        raise RuntimeError(f"DOCX file was not created: {path}")
    if path.stat().st_size <= 0:
        raise RuntimeError(f"DOCX file is empty: {path}")


def _verify_docx_readable(path: Path) -> None:
    try:
        Document(path)
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"DOCX validation failed for {path}: {exc}") from exc


def _read_docx_visible_content_summary(path: Path) -> _DocxVisibleContentSummary:
    try:
        with ZipFile(path, "r") as archive:
            document_xml = archive.read("word/document.xml")
    except KeyError as exc:
        raise RuntimeError(f"DOCX is missing word/document.xml: {path}") from exc
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"Failed reading DOCX body XML from {path}: {exc}") from exc

    try:
        root = ET.fromstring(document_xml)
    except ET.ParseError as exc:
        raise RuntimeError(f"DOCX body XML is malformed in {path}: {exc}") from exc

    # Editable tables are body content too.  Exclude headers/footer fields, not cells.
    paragraphs = root.findall(".//w:body//w:p", _DOCX_NS)
    paragraph_count = len(paragraphs)
    visible_paragraph_count = 0
    text_node_count = 0
    text_char_count = 0
    word_count = 0

    for paragraph in paragraphs:
        text_nodes = [node.text or "" for node in paragraph.findall(".//w:t", _DOCX_NS)]
        text_node_count += len(text_nodes)
        paragraph_text = "".join(text_nodes)
        if not paragraph_text.strip():
            continue
        visible_paragraph_count += 1
        text_char_count += len(paragraph_text.strip())
        word_count += len(_VISIBLE_WORD_RE.findall(paragraph_text))

    return _DocxVisibleContentSummary(
        paragraph_count=paragraph_count,
        visible_paragraph_count=visible_paragraph_count,
        text_node_count=text_node_count,
        text_char_count=text_char_count,
        word_count=word_count,
    )


def _verify_docx_visible_content(
    path: Path,
    *,
    stage: str,
    expected_visible_paragraphs: int,
    baseline: _DocxVisibleContentSummary | None = None,
) -> _DocxVisibleContentSummary:
    summary = _read_docx_visible_content_summary(path)
    if expected_visible_paragraphs > 0 and summary.visible_paragraph_count == 0:
        raise RuntimeError(
            f"DOCX content verification failed after {stage}: {path} has no visible paragraphs"
        )
    if expected_visible_paragraphs > 0 and summary.text_char_count == 0:
        raise RuntimeError(
            f"DOCX content verification failed after {stage}: {path} has no visible text"
        )
    if baseline is not None and baseline.visible_paragraph_count > 0 and summary.visible_paragraph_count == 0:
        raise RuntimeError(
            f"DOCX content verification failed after {stage}: {path} lost all visible paragraphs"
        )
    return summary


def _fsync_file(path: Path) -> None:
    try:
        with path.open("r+b") as handle:
            handle.flush()
            os.fsync(handle.fileno())
    except OSError as exc:
        raise RuntimeError(f"Failed to fsync temporary DOCX: {path}") from exc


def resolve_noncolliding_output_path(output_path: Path) -> Path:
    final_path = output_path.expanduser().resolve()
    if not final_path.exists():
        return final_path

    stem = final_path.stem
    suffix = final_path.suffix
    index = 1
    while True:
        candidate = final_path.with_name(f"{stem}_{index:02d}{suffix}")
        if not candidate.exists():
            return candidate
        index += 1


def save_document_atomic(
    document: Document,
    output_path: Path,
    *,
    verify_readable: bool = True,
) -> Path:
    final_path = resolve_noncolliding_output_path(output_path)
    tmp_path = final_path.with_name(f"{final_path.name}.tmp")

    final_path.parent.mkdir(parents=True, exist_ok=True)
    if tmp_path.exists():
        tmp_path.unlink()

    try:
        document.save(tmp_path)
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"Failed writing temporary DOCX: {tmp_path}") from exc

    _verify_non_empty_file(tmp_path)
    _fsync_file(tmp_path)
    if verify_readable:
        _verify_docx_readable(tmp_path)

    try:
        os.replace(tmp_path, final_path)
    except OSError as exc:
        raise RuntimeError(f"Failed to atomically replace DOCX at {final_path}") from exc

    _verify_non_empty_file(final_path)
    if verify_readable:
        _verify_docx_readable(final_path)
    return final_path


# Surgical regex patterns for removing compatibilityMode from settings.xml
# without re-serializing the entire XML (which can break namespace prefixes).
_COMPAT_SETTING_RE = re.compile(
    r'<w:compatSetting\b[^>]*\bw:name\s*=\s*"compatibilityMode"[^>]*/>'
    r"|"
    r'<w:compatSetting\b[^>]*\bw:name\s*=\s*"compatibilityMode"[^>]*>.*?</w:compatSetting>',
    re.DOTALL,
)
_EMPTY_COMPAT_RE = re.compile(
    r"<w:compat\b[^>]*/>\s*"
    r"|"
    r"<w:compat\b[^>]*>\s*</w:compat>\s*",
    re.DOTALL,
)


def _clone_zipinfo_for_rewrite(info: ZipInfo) -> ZipInfo:
    clone = ZipInfo(filename=info.filename, date_time=info.date_time)
    clone.compress_type = info.compress_type
    clone.comment = info.comment
    clone.extra = info.extra
    clone.create_system = info.create_system
    clone.create_version = info.create_version
    clone.extract_version = info.extract_version
    clone.internal_attr = info.internal_attr
    clone.external_attr = info.external_attr
    return clone


def _remove_compatibility_mode(docx_path: Path) -> None:
    """Remove compatibilityMode from word/settings.xml to avoid Word upgrade prompt.

    Uses regex surgery on the raw XML text so that namespace declarations,
    prefix mappings, and mc:Ignorable contracts are preserved byte-for-byte.
    """
    with ZipFile(docx_path, "r") as zin:
        if "word/settings.xml" not in zin.namelist():
            return
        settings_text = zin.read("word/settings.xml").decode("utf-8")

    modified = _COMPAT_SETTING_RE.sub("", settings_text)
    if modified == settings_text:
        return  # Nothing to remove.

    # Clean up empty <w:compat> element if all children were removed.
    modified = _EMPTY_COMPAT_RE.sub("", modified)

    tmp_path = docx_path.with_name(f"{docx_path.name}.compat_tmp")
    try:
        with ZipFile(docx_path, "r") as zin, ZipFile(tmp_path, "w") as zout:
            for item in zin.infolist():
                item_copy = _clone_zipinfo_for_rewrite(item)
                if item.filename == "word/settings.xml":
                    zout.writestr(item_copy, modified.encode("utf-8"))
                else:
                    zout.writestr(item_copy, zin.read(item.filename))
        os.replace(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)


_ROLE_STYLES = {
    "paragraph": "LegalPDF Body",
    "heading": "LegalPDF Heading",
    "header": "LegalPDF Header",
    "footer": "LegalPDF Footer Text",
    "address": "LegalPDF Address",
    "list_item": "LegalPDF List",
    "signature": "LegalPDF Signature",
    "reference": "LegalPDF Reference",
    "table_cell": "LegalPDF Table Text",
}
_ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_A4_SIZE_PT = (595.276, 841.89)
_PAGE_LABEL_RE = re.compile(
    r"(?:[0-9٠-٩۰-۹]+(?:\s*(?:/|of|de|sur|من)\s*[0-9٠-٩۰-۹]+)?"
    r"|(?:p[aá]g(?:ina)?\.?|page|صفحة|الصفحة)\s*:?\s*[0-9٠-٩۰-۹]+"
    r"(?:\s*(?:/|of|de|sur|من)\s*[0-9٠-٩۰-۹]+)?)",
    re.IGNORECASE,
)


@dataclass(frozen=True)
class _AssemblyPage:
    number: int
    text: str
    structure: dict | None
    structure_status: str


def _text_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _load_assembly_page(path: Path) -> _AssemblyPage:
    """Treat sidecars as optional hints, never as authority to discard TXT content."""
    number = int(path.stem.split("_")[1])
    text = path.read_text(encoding="utf-8")
    sidecar = path.with_suffix(".structure.json")
    if not sidecar.exists():
        return _AssemblyPage(number, text, None, "legacy_txt")
    try:
        from .document_structure import plain_text_from_structure, validate_page_structure

        if sidecar.stat().st_size > 8 * 1024 * 1024:
            raise ValueError("Structure sidecar exceeds the local evidence bound.")
        structure = validate_page_structure(json.loads(sidecar.read_text(encoding="utf-8")))
        if structure.page_number != number:
            raise ValueError("Structure belongs to a different source page.")
        expected_hash = _text_hash(text)
        if structure.translation_sha256 != expected_hash:
            raise ValueError("Saved translation and structure hashes differ.")
        if _text_hash(plain_text_from_structure(structure)) != expected_hash:
            raise ValueError("Structure does not contain the complete saved translation.")
        # A valid schema may still describe a sparse or impractically large table.
        # Bound allocations and reject interleaved table fragments, rather than
        # moving text around or silently rendering only part of a page.
        tables: dict[str, tuple[int, int]] = {}
        completed_tables: set[str] = set()
        previous_table = None
        for block in structure.blocks:
            table_id = block.table_id
            if previous_table and table_id != previous_table:
                completed_tables.add(previous_table)
            if table_id:
                if table_id in completed_tables:
                    raise ValueError("Table cells are not contiguous in source order.")
                rows, cols = tables.get(table_id, (0, 0))
                rows, cols = max(rows, block.row + 1), max(cols, block.col + 1)
                if rows > 200 or cols > 32 or rows * cols > 2000:
                    raise ValueError("Table geometry exceeds the editable layout bound.")
                tables[table_id] = (rows, cols)
            previous_table = table_id
        payload = structure.to_dict()
        from .layout_integration import load_rebuild_layout

        rebuild_layout = load_rebuild_layout(path, payload)
        if rebuild_layout is not None:
            payload.setdefault("metadata", {})["layout"] = rebuild_layout
        return _AssemblyPage(number, text, payload, "validated")
    except (OSError, ValueError, TypeError, KeyError, ImportError, AttributeError):
        # TXT remains the compatibility contract, including for old/partial runs.
        return _AssemblyPage(number, text, None, "invalid_sidecar_txt_fallback")


def _font_profile(lang: TargetLang | str) -> tuple[str, float]:
    return ("Arial", 11.0) if _is_rtl_target_lang(lang) else ("Times New Roman", 10.5)


def _validated_source_pair(previous: _AssemblyPage, current: _AssemblyPage, pages_dir: Path) -> list | None:
    """Bind complete original source pages to the validated target structures."""
    from .document_structure import PageStructure

    sources = []
    try:
        for page in (previous, current):
            if not page.structure:
                return None
            path = pages_dir / f"page_{page.number:04d}.source_structure.json"
            if not path.is_file():
                path = pages_dir / f"page_{page.number:04d}.source.json"
            if path.stat().st_size > 8 * 1024 * 1024:
                return None
            payload = json.loads(path.read_text(encoding="utf-8"))
            source = PageStructure.from_dict(payload.get("structure", payload))
            if (source.page_number != page.number or not source.source_file_sha256
                    or source.source_file_sha256 != page.structure["source_file_sha256"]
                    or source.source_sha256 != page.structure["source_sha256"]
                    or source.source_text_sha256 != source.source_sha256
                    or _text_hash(source.text) != source.source_sha256):
                return None
            source_rows = source.to_dict()["blocks"]
            target_rows = page.structure["blocks"]
            page_keys = ("width_pt", "height_pt", "uncertain", "document_start")
            if any(getattr(source, key) != page.structure.get(key) for key in page_keys):
                return None
            if len(source_rows) != len(target_rows):
                return None
            # Layout roles/positions must still refer to the same complete source
            # blocks; translated text is intentionally not an equality witness.
            keys = ("id", "role", "bbox", "table_id", "alignment", "uncertain", "document_start")
            if any(any(a.get(key) != b.get(key) for key in keys) for a, b in zip(source_rows, target_rows)):
                return None
            sources.append(source)
        return sources
    except (OSError, ValueError, TypeError, KeyError, AttributeError):
        return None


def _validated_furniture_bridge(previous: _AssemblyPage, current: _AssemblyPage, pages_dir: Path) -> dict | None:
    """Recheck full adjacent SOURCE evidence, not translated headers or flags.

    Legacy plain-adjacent links still work without source sidecars. Moving a
    fragment around page furniture is stricter: absent/stale source evidence
    leaves all blocks in their original editable order.
    """
    from .formatting_support import confirmed_source_continuation

    sources = _validated_source_pair(previous, current, pages_dir)
    if sources is None:
        return None
    try:
        proof = confirmed_source_continuation(*sources)
        if not proof or not proof["furniture_pairs"]:
            return None
        declared = current.structure.get("metadata", {}).get("continuation_bridge")
        if declared is not None and declared != proof:
            return None
        return proof
    except (OSError, ValueError, TypeError, KeyError, AttributeError):
        return None


def _reflow_compact_lists(document, pages: list[_AssemblyPage], source_map: list[dict],
                          pages_dir: Path, *, page_breaks: bool) -> None:
    """Move only proven leading list paragraphs ahead of repeated furniture.

    Work on complete already-written paragraphs so content, runs and furniture
    occurrences cannot be dropped or rewritten. Resolve final locators after
    movement, including aliases belonging to older sentence-continuation joins.
    """
    if page_breaks:
        return
    from .docx_furniture_reflow import confirmed_list_furniture_reflow, potential_list_furniture_boundary

    paragraphs = document.paragraphs
    mapped_nodes = []
    by_id = {}
    for page_map in source_map:
        for mapping in page_map["blocks"]:
            locator = mapping.get("location", {})
            if locator.get("kind") == "body_paragraph":
                node = paragraphs[locator["paragraph_index"]]._element
                mapped_nodes.append((locator, node))
                if mapping["block_id"] in by_id:
                    return  # Ambiguous cross-page IDs cannot authorize movement.
                by_id[mapping["block_id"]] = node
    admitted_furniture = set()
    moved = False
    for index in range(1, len(pages)):
        previous, current = pages[index - 1:index + 1]
        previous_map, current_map = source_map[index - 1:index + 1]
        if (not previous.structure or not current.structure or previous.number + 1 != current.number
                or previous_map.get("section_furniture_adopted") or current_map.get("section_furniture_adopted")
                or current.structure.get("document_start")
                or any(page_map["layout_status"] not in {"flow", "not_provided"}
                       or page_map["layout_review_required"] for page_map in (previous_map, current_map))
                or any(abs(a - b) > 1 for a, b in zip(_page_size(previous), _page_size(current)))
                or not potential_list_furniture_boundary(previous.structure, current.structure)):
            continue
        sources = _validated_source_pair(previous, current, pages_dir)
        proof = confirmed_list_furniture_reflow(*(source.to_dict() for source in sources)) if sources else None
        if proof:
            group = [by_id.get(identity) for identity in proof["current_list_ids"]]
            tail = by_id.get(proof["previous_tail_id"])
            packet = {by_id.get(identity) for identity in proof["previous_trailing_ids"] + proof["current_leading_ids"]}
            valid_nodes = tail is not None and all(node is not None for node in group) and None not in packet
            if valid_nodes:
                # A chained list can have older admitted furniture waiting after
                # its preceding item. Carry it too, but cross no other content,
                # semantic table, section/page break or unmapped empty paragraph.
                intervening = []
                node = tail.getnext()
                while node is not None and node is not group[0]:
                    intervening.append(node)
                    node = node.getnext()
                valid_nodes = (node is group[0] and bool(intervening)
                               and packet <= set(intervening)
                               and set(intervening) <= packet | admitted_furniture
                               and all(a.getnext() is b for a, b in zip(group, group[1:])))
            if valid_nodes:
                anchor = intervening[0]
                for node in group:
                    anchor.addprevious(node)
                admitted_furniture.update(packet)
                current_map.setdefault("list_furniture_reflows", []).append(proof)
                moved = True
                continue
        current_map["layout_review_required"] = True
        if "list_furniture_reflow_not_evaluated" not in current_map["layout_warnings"]:
            current_map["layout_warnings"].append("list_furniture_reflow_not_evaluated")
    if moved:
        final_indices = {paragraph._element: index for index, paragraph in enumerate(document.paragraphs)}
        for locator, node in mapped_nodes:
            locator["paragraph_index"] = final_indices[node]


def _set_font_properties(owner, *, name: str, size: float) -> None:
    """Set both Western and complex-script slots; do not inherit theme fonts."""
    owner.font.name = name
    owner.font.size = Pt(size)
    owner.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = owner._element.get_or_add_rPr()
    fonts = r_pr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme", "csTheme"):
        fonts.attrib.pop(qn(f"w:{attr}"), None)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), name)
    size_cs = r_pr.find(qn("w:szCs"))
    if size_cs is None:
        size_cs = OxmlElement("w:szCs")
        r_pr.append(size_cs)
    size_cs.set(qn("w:val"), str(round(size * 2)))


def _configure_styles(document, lang: TargetLang | str) -> None:
    name, size = _font_profile(lang)
    normal = document.styles["Normal"]
    _set_font_properties(normal, name=name, size=size)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.widow_control = True
    for role, style_name in _ROLE_STYLES.items():
        style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        _set_font_properties(style, name=name, size=size)
        fmt = style.paragraph_format
        fmt.line_spacing = 1.0
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(3 if role in {"paragraph", "list_item"} else 0)
        fmt.widow_control = True
        if role == "heading":
            style.font.bold = True
            style.font.cs_bold = True
            fmt.space_before = Pt(6)
            fmt.space_after = Pt(3)
            fmt.keep_with_next = True
            fmt.keep_together = True
        elif role == "signature":
            fmt.space_before = Pt(3)
            fmt.keep_together = True
        elif role == "list_item":
            # Keep literal legal numbering. Do not create a second Word number.
            if _is_rtl_target_lang(lang):
                fmt.right_indent = Cm(0.35)
            else:
                fmt.left_indent = Cm(0.35)
            fmt.first_line_indent = Cm(-0.35)


def _configure_section(section, size: tuple[float, float]) -> None:
    section.page_width, section.page_height = Pt(size[0]), Pt(size[1])
    section.left_margin = section.right_margin = Cm(1.7)
    section.top_margin = section.bottom_margin = Cm(1.5)
    section.header_distance = section.footer_distance = Cm(0.7)


def _page_size(page: _AssemblyPage) -> tuple[float, float]:
    if page.structure:
        size = (page.structure["width_pt"], page.structure["height_pt"])
        # Tiny image geometry cannot support the agreed readable margins/fonts.
        # Do not shrink text to fit it; use the documented A4 fallback instead.
        if min(size) >= 144:
            return size
    return _A4_SIZE_PT


def _add_page_number_footer(document, lang: TargetLang | str, *, section=None, append: bool = False) -> None:
    footer = (section or document.sections[0]).footer
    paragraph = footer.add_paragraph() if append else footer.paragraphs[0]
    paragraph.style = _ROLE_STYLES["footer"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    name, size = _font_profile(lang)
    _set_font_properties(run, name=name, size=size)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, cached, end):
        run._r.append(element)
    # PAGE is computed by Word's pagination/print engine. A global updateFields
    # request is unnecessary and can open the "fields may refer to other files"
    # modal, blocking read-only automation even though this file has no links.


def _append_block_text(paragraph, text: str, *, lang: TargetLang | str, strip_bidi_controls: bool,
                       bold: bool = False, italic: bool = False) -> int:
    name, size = _font_profile(lang)
    rtl = _is_rtl_target_lang(lang)
    if rtl:
        runs, mixed = _segment_rtl_placeholder_aware_runs(text, strip_bidi_controls=strip_bidi_controls)
    else:
        runs, mixed = [("ltr", sanitize_bidi_controls(text) if strip_bidi_controls else text)], False
    count = 0
    for kind, chunk in runs:
        if not chunk:
            continue
        if kind == "ltr" and mixed:
            chunk = _wrap_ltr_run_with_lrm(chunk)
        run = paragraph.add_run(chunk)
        _set_font_properties(run, name=name, size=size)
        if bold:
            run.bold = True
            run.font.cs_bold = True
        if italic:
            run.italic = True
            run.font.cs_italic = True
        if rtl:
            if kind == "ltr":
                _set_ltr_run_props(run)
            else:
                _set_rtl_run_props(run, bidi_lang=_rtl_bidi_lang_code(lang))
        count += 1
    return count


def _format_block_paragraph(paragraph, block: dict, lang: TargetLang | str,
                            *, physical_alignment: str | None = None,
                            in_layout_container: bool = False,
                            readable_body_spacing: bool = False) -> None:
    paragraph.style = _ROLE_STYLES[block["role"]]
    rtl = _is_rtl_target_lang(lang)
    if (readable_body_spacing and str(getattr(lang, "value", lang)).upper() == "AR"
            and block["role"] in {"paragraph", "list_item"} and not in_layout_container):
        # Explicit readable style fallback; source boxes do not prove line height.
        paragraph.paragraph_format.line_spacing = 1.10
    if rtl:
        _add_rtl_flags(paragraph)
    # A source prose/header alignment follows the target reading direction.
    # Physical layout-region and semantic-table alignment is a separate choice.
    alignment = physical_alignment or block.get("alignment")
    lang_code = str(getattr(lang, "value", lang)).strip().upper()
    arabic_leading_roles = {"paragraph", "list_item", "heading", "header", "signature"}
    if (physical_alignment is None and lang_code == "AR" and block["role"] in arabic_leading_roles
            and alignment == "left" and any(_classify_directional_char(char) == "rtl"
                                            for char in block.get("text", ""))):
        alignment = "right"
    if alignment in _ALIGNMENTS:
        if rtl:
            _set_rtl_visual_alignment(paragraph, alignment)
        else:
            paragraph.alignment = _ALIGNMENTS[alignment]
    if in_layout_container and lang_code == "AR" and block["role"] == "list_item":
        # Word clips a layout cell at its physical edge. The body list style's
        # RTL hanging indent puts a literal leading marker beyond that edge,
        # although PDF text extraction still reports the hidden character.
        # Override inherited hanging explicitly; do not change text/numbering
        # or the full-width list style and semantic data-table formatting.
        ind = paragraph._p.get_or_add_pPr().get_or_add_ind()
        for attr in ("left", "right", "hanging"):
            ind.set(qn(f"w:{attr}"), "0")
        ind.attrib.pop(qn("w:firstLine"), None)


def _signature_group_keep_next(blocks: list[dict], lang: TargetLang | str,
                               *, excluded_ids: set[str] | None = None,
                               alignment: str = "left") -> dict[str, bool]:
    """Unify a short mixed-script signature, not unrelated Latin-only labels.

    Source-confirmed adjacent signature rows with the same left anchor are one
    visual group. Adapting only their Arabic rows splits the Latin signer's name
    to the opposite edge. Keep the group right-aligned while its name run stays
    LTR. Never infer a group across a page, region, panel or document boundary.
    The optional centered mode provides cohesion only to the source-bound
    compact postpass; its caller must retain the existing centered alignment.
    """
    if str(getattr(lang, "value", lang)).strip().upper() != "AR" or alignment not in {"left", "center"}:
        return {}
    result: dict[str, bool] = {}
    group: list[dict] = []

    def finish() -> None:
        kinds = [{_classify_directional_char(char) for char in _visible_signature(block)} for block in group]
        if (2 <= len(group) <= 6 and any("rtl" in row for row in kinds)
                and any("ltr" in row and "rtl" not in row for row in kinds)):
            result.update((block["id"], index < len(group) - 1) for index, block in enumerate(group))
        group.clear()

    def anchor(bbox) -> float:
        return (bbox[0] + bbox[2]) / 2 if alignment == "center" else bbox[0]

    for block in blocks:
        bbox = block.get("bbox")
        visible = _visible_signature(block)
        eligible = (block.get("role") == "signature" and block.get("alignment") == alignment
                    and not block.get("uncertain") and bbox and 0 < len(visible) <= 160
                    and block["id"] not in (excluded_ids or set()))
        if not eligible:
            finish()
            continue
        if group and (block.get("document_start") or abs(anchor(bbox) - anchor(group[0]["bbox"])) > 9
                      or not -2 <= bbox[1] - group[-1]["bbox"][3] <= 24
                      or bbox[3] - group[0]["bbox"][1] > 96):
            finish()
        group.append(block)
    finish()
    return result


def _visible_signature(block: dict) -> str:
    return sanitize_bidi_controls(unwrap_internal_placeholders(block.get("text", "")))


def _force_page_boundary(document, size: tuple[float, float], *, change_size: bool = False) -> None:
    children = list(document._element.body)
    last_content = next((element for element in reversed(children) if element.tag != qn("w:sectPr")), None)
    if last_content is None:
        _configure_section(document.sections[-1], size)
        return
    if change_size or last_content.tag == qn("w:tbl"):
        _configure_section(document.add_section(WD_SECTION_START.NEW_PAGE), size)
    else:
        document.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)


def _table_column_widths(cells: list[dict], *, columns: int, width_pt: float,
                         source_width_pt: float) -> tuple[list[float], str]:
    """Use consistent source column intervals, not translated string lengths.

    A missing, overlapping, or merged-cell-looking interval is not reliable
    evidence of a column boundary. Fall back to equal editable columns in that
    case and leave text intact for human layout review.
    """
    fallback = ([width_pt / columns] * columns, "equal_width_fallback")
    bounds: dict[int, list[tuple[float, float]]] = {col: [] for col in range(columns)}
    for cell in cells:
        bbox = cell.get("bbox")
        if not bbox or not 0 <= bbox[0] < bbox[2] <= source_width_pt:
            return fallback
        bounds[cell["col"]].append((bbox[0], bbox[2]))
    if any(not intervals for intervals in bounds.values()):
        return fallback
    tolerance = max(3.0, source_width_pt * 0.015)
    intervals = []
    for col in range(columns):
        lefts = [left for left, _ in bounds[col]]
        if max(lefts) - min(lefts) > tolerance:
            return fallback
        intervals.append((median(lefts), max(right for _, right in bounds[col]), col))
    physical = sorted(intervals)
    order = [col for _, _, col in physical]
    if order not in (list(range(columns)), list(reversed(range(columns)))):
        return fallback
    if any(left[1] > right[0] + tolerance for left, right in zip(physical, physical[1:])):
        return fallback
    boundaries = [physical[0][0]]
    boundaries.extend((left[1] + right[0]) / 2 for left, right in zip(physical, physical[1:]))
    boundaries.append(physical[-1][1])
    extent = boundaries[-1] - boundaries[0]
    if extent <= 0:
        return fallback
    widths = [0.0] * columns
    for index, (_, _, col) in enumerate(physical):
        widths[col] = width_pt * (boundaries[index + 1] - boundaries[index]) / extent
    if min(widths) < min(18.0, width_pt / columns):
        return fallback
    return widths, "source_geometry"


def _format_table(table, *, rtl: bool, width_pt: float, cells: list[dict], source_width_pt: float) -> str:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths, geometry_status = _table_column_widths(cells, columns=len(table.columns), width_pt=width_pt,
                                                  source_width_pt=source_width_pt)
    for index, column in enumerate(table.columns):
        column.width = Pt(widths[index])
        for cell in column.cells:
            cell.width = Pt(widths[index])
    tbl_pr = table._tbl.tblPr
    if rtl:
        visual = OxmlElement("w:bidiVisual")
        visual.set(qn("w:val"), "1")
        tbl_pr.append(visual)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "B7B7B7")
        borders.append(border)
    tbl_pr.append(borders)
    margins = OxmlElement("w:tblCellMar")
    for edge in ("top", "left", "bottom", "right"):
        margin = OxmlElement(f"w:{edge}")
        margin.set(qn("w:w"), "60")
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    tbl_pr.append(margins)
    return geometry_status


def _validated_layout(structure: dict) -> tuple[dict | None, str, bool, list[str]]:
    """Layout is optional and cannot invalidate complete translation evidence."""
    metadata = structure.get("metadata", {})
    if "layout" not in metadata:
        return None, "not_provided", False, []
    try:
        from .document_layout import validate_page_layout

        layout = validate_page_layout(metadata["layout"], structure)
        status = layout["status"]
        warnings = list(layout.get("warnings", []))
        if status == "needs_review":
            return None, "needs_review_flow_fallback", True, warnings
        if status == "regions":
            width = _page_size(_AssemblyPage(0, "", structure, "validated"))[0] - Cm(3.4).pt
            for band in layout["bands"]:
                edges = band["column_edges_pt"]
                for region in band["regions"]:
                    start, span = region["column_start"], region["column_span"]
                    if width * (edges[start + span] - edges[start]) / (edges[-1] - edges[0]) < 36:
                        return None, "unsupported_layout_flow_fallback", True, [*warnings, "region_too_narrow"]
        return layout, status, bool(layout.get("review_required")), warnings
    except (ValueError, TypeError, KeyError, AttributeError, ImportError):
        return None, "invalid_layout_flow_fallback", True, ["invalid_or_unsupported_layout"]


def _format_layout_table(table, widths: list[float]) -> None:
    """Physical columns are never mirrored because their text is Arabic.

    This is a layout container, not a semantic data table. Allow natural row
    splitting and growth; fixed heights/cantSplit would clip long legal text or
    force an entire panel to the next page.
    """
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, column in enumerate(table.columns):
        column.width = Pt(widths[index])
        for cell in column.cells:
            cell.width = Pt(widths[index])
    props = table._tbl.tblPr
    for name in ("bidiVisual", "tblBorders", "tblCellMar"):
        for old in list(props.findall(qn(f"w:{name}"))):
            props.remove(old)
    bidi = OxmlElement("w:bidiVisual")
    bidi.set(qn("w:val"), "0")
    props.append(bidi)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    props.append(borders)
    margins = OxmlElement("w:tblCellMar")
    for edge in ("top", "left", "bottom", "right"):
        margin = OxmlElement(f"w:{edge}")
        margin.set(qn("w:w"), "0")
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    props.append(margins)
    for row in table.rows:
        row_props = row._tr.get_or_add_trPr()
        for name in ("cantSplit", "trHeight"):
            for old in list(row_props.findall(qn(f"w:{name}"))):
                row_props.remove(old)
        for cell in row.cells:
            valign = OxmlElement("w:vAlign")
            valign.set(qn("w:val"), "top")
            cell._tc.get_or_add_tcPr().append(valign)


def _set_layout_cell_style(cell, *, fill: str | None = None, padding_pt: float = 3.5) -> None:
    props = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        margin = OxmlElement(f"w:{edge}")
        margin.set(qn("w:w"), str(round(padding_pt * 20)))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    props.append(margins)
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), fill)
        props.append(shading)


def _container_paragraph(container):
    """Reuse only the mandatory empty cell paragraph, never a content block."""
    paragraphs = container.paragraphs
    if paragraphs and not paragraphs[-1].text and not list(paragraphs[-1]._p.iter(qn("w:br"))):
        last = paragraphs[-1]
        if last._p.getnext() is None:
            return last
    return container.add_paragraph("")


def _add_nested_layout_table(container, *, rows: int, cols: int):
    # A cell requires a final paragraph, but not an extra blank paragraph before
    # its first/nested table. python-docx supplies the required trailing one.
    paragraphs = container.paragraphs
    if paragraphs and not paragraphs[-1].text and paragraphs[-1]._p.getnext() is None:
        paragraph = paragraphs[-1]._p
        if not list(paragraph.iter(qn("w:br"))):
            paragraph.getparent().remove(paragraph)
    return container.add_table(rows=rows, cols=cols)


def _render_region_blocks(container, blocks: list[dict], *, panels: list[dict],
                          lang: TargetLang | str, strip_bidi_controls: bool,
                          width_pt: float, source_width_pt: float, location: dict,
                          mappings: dict[str, dict], physical_alignment: str | None = None) -> None:
    panel_starts = {panel["block_ids"][0]: panel for panel in panels}
    signature_groups = (_signature_group_keep_next(
        blocks, lang, excluded_ids={bid for panel in panels for bid in panel["block_ids"]},
    ) if physical_alignment is None else {})
    index = 0
    while index < len(blocks):
        block = blocks[index]
        panel = panel_starts.get(block["id"])
        if panel:
            count = len(panel["block_ids"])
            panel_table_index = len(container.tables)
            table = _add_nested_layout_table(container, rows=1, cols=1)
            _format_layout_table(table, [width_pt])
            cell = table.cell(0, 0)
            _set_layout_cell_style(cell, fill=panel.get("shading", {}).get("fill"), padding_pt=5.0)
            panel_location = {**location, "panel_id": panel["id"],
                              "table_path": [*location["table_path"],
                                             {"table_index": panel_table_index, "row": 0, "col": 0}]}
            _render_region_blocks(cell, blocks[index:index + count], panels=[], lang=lang,
                                  strip_bidi_controls=strip_bidi_controls, width_pt=max(1, width_pt - 10),
                                  source_width_pt=source_width_pt, location=panel_location,
                                  mappings=mappings, physical_alignment=physical_alignment)
            index += count
            continue
        index += 1
        mapping = mappings[block["id"]]
        role, text = block["role"], block["text"]
        if role == "footer" and _PAGE_LABEL_RE.fullmatch(sanitize_bidi_controls(text).strip()):
            mapping["location"] = {"kind": "generated_footer_page_field"}
            continue
        if role != "table_cell" and not sanitize_bidi_controls(unwrap_internal_placeholders(text)).strip():
            mapping["location"] = {"kind": "empty_block"}
            continue
        if role == "table_cell":
            cells = [block]
            while index < len(blocks) and blocks[index].get("table_id") == block["table_id"]:
                cells.append(blocks[index])
                index += 1
            table_index = len(container.tables)
            table = _add_nested_layout_table(container, rows=max(cell["row"] for cell in cells) + 1,
                                             cols=max(cell["col"] for cell in cells) + 1)
            geometry_status = _format_table(table, rtl=_is_rtl_target_lang(lang), width_pt=width_pt,
                                            cells=cells, source_width_pt=source_width_pt)
            for cell_block in cells:
                cell = table.cell(cell_block["row"], cell_block["col"])
                paragraph = cell.paragraphs[0]
                _format_block_paragraph(paragraph, cell_block, lang,
                                        physical_alignment=cell_block.get("alignment"))
                _append_block_text(paragraph, cell_block["text"], lang=lang,
                                   strip_bidi_controls=strip_bidi_controls,
                                   bold=cell_block.get("bold", False), italic=cell_block.get("italic", False))
                mappings[cell_block["id"]].update(
                    location={**location, "kind": "table_cell", "table_index": table_index,
                              "row": cell_block["row"], "col": cell_block["col"], "paragraph_index": 0,
                              "table_path": [*location["table_path"],
                                             {"table_index": table_index, "row": cell_block["row"],
                                              "col": cell_block["col"]}]},
                    table_geometry_status=geometry_status,
                )
            continue
        paragraph = _container_paragraph(container)
        explicit = "center" if block.get("alignment") == "center" else physical_alignment
        if block["id"] in signature_groups:
            explicit = "right"
        _format_block_paragraph(paragraph, block, lang, physical_alignment=explicit,
                                in_layout_container=True)
        if block["id"] in signature_groups:
            paragraph.paragraph_format.keep_with_next = signature_groups[block["id"]]
        _append_block_text(paragraph, text, lang=lang, strip_bidi_controls=strip_bidi_controls,
                           bold=block.get("bold", False), italic=block.get("italic", False))
        mapping["location"] = {**location, "kind": "layout_region_paragraph",
                               "paragraph_index": len(container.paragraphs) - 1}
    # Only the mandatory, visibly empty paragraph after a nested table is
    # compacted. This never reduces a text run's agreed readable font size.
    if container.paragraphs and not container.paragraphs[-1].text:
        tail = container.paragraphs[-1]
        tail.paragraph_format.space_before = tail.paragraph_format.space_after = Pt(0)
        tail.paragraph_format.line_spacing = Pt(1)


def _render_page_layout(document, layout: dict, blocks: list[dict], *, lang: TargetLang | str,
                        strip_bidi_controls: bool, size: tuple[float, float],
                        prior_page_content: bool) -> list[dict]:
    by_id = {block["id"]: block for block in blocks}
    mappings = {block["id"]: {"block_id": block["id"], "role": block["role"]} for block in blocks}
    available_width = size[0] - Cm(3.4).pt
    for band_index, band in enumerate(layout["bands"]):
        band_ids = [bid for region in band["regions"] for bid in region["block_ids"]]
        if any(by_id[bid].get("document_start") for bid in band_ids) and (band_index or prior_page_content):
            _force_page_boundary(document, size)
        edges = band["column_edges_pt"]
        scale = available_width / (edges[-1] - edges[0])
        widths = [(right - left) * scale for left, right in zip(edges, edges[1:])]
        table_index = len(document.tables)
        table = document.add_table(rows=1, cols=len(widths))
        _format_layout_table(table, widths)
        for region in sorted(band["regions"], key=lambda region: region["column_start"]):
            start, span = region["column_start"], region["column_span"]
            cell = table.cell(0, start)
            if span > 1:
                cell = cell.merge(table.cell(0, start + span - 1))
            region_width = sum(widths[start:start + span])
            cell.width = Pt(region_width)
            _set_layout_cell_style(cell)
            _render_region_blocks(
                cell, [by_id[bid] for bid in region["block_ids"]], panels=region.get("panels", []),
                lang=lang, strip_bidi_controls=strip_bidi_controls,
                width_pt=max(1, region_width - 7), source_width_pt=size[0],
                location={"layout_band_id": band["id"], "layout_region_id": region["id"],
                          "table_path": [{"table_index": table_index, "row": 0, "col": start}]},
                mappings=mappings, physical_alignment=region.get("alignment"),
            )
    if any("location" not in mapping for mapping in mappings.values()):
        raise RuntimeError("Layout assembly did not map every complete source block.")
    return list(mappings.values())


def _section_furniture_plan(pages: list[_AssemblyPage], pages_dir: Path, *, page_breaks: bool) -> dict:
    """Use independently bound complete sources; never infer repeats from targets."""
    from .section_furniture import plan_section_furniture

    pairs = []
    for page in pages:
        source_pair = _validated_source_pair(page, page, pages_dir) if page.structure else None
        pairs.append((source_pair[0].to_dict(), page.structure) if source_pair else None)
    return plan_section_furniture(pairs, page_breaks=page_breaks)


def _furniture_line_estimate(blocks: list[dict], width_pt: float, font_size: float) -> int:
    """Conservative wrap allowance, not a claim about Word's final pagination."""
    import math

    return sum(max(1, math.ceil(len(line) * font_size * 0.65 / max(1, width_pt)))
               for block in blocks
               for line in sanitize_bidi_controls(unwrap_internal_placeholders(block["text"])).split("\n"))


def _furniture_margins(plan: dict | None, size: tuple[float, float], lang: TargetLang | str) -> tuple[float, float]:
    parts = (plan or {}).get("parts", {})
    _, font_size = _font_profile(lang)
    top = bottom = Cm(1.5).pt
    width = size[0] - 2 * Cm(1.7).pt
    if parts.get("header"):
        lines = _furniture_line_estimate(parts["header"]["canonical_blocks"], width, font_size)
        top = max(top, Cm(0.7).pt + lines * font_size * 1.2 + 18)
    if parts.get("footer"):
        lines = _furniture_line_estimate(parts["footer"]["canonical_blocks"], width, font_size)
        # Keep the conservative wrapping allowance plus PAGE, but avoid a
        # second oversized blank reserve beyond those already reserved lines.
        # This is a style allowance, not a measured ink-to-body guarantee.
        bottom = max(bottom, Cm(0.7).pt + (lines + 1) * font_size * 1.2 + 6)
    return top, bottom


def _bound_furniture_reserves(plan: dict, pages: list[_AssemblyPage], lang: TargetLang | str) -> None:
    """Reject pathological translations before making header/footer parts."""
    for section in plan.get("sections", []):
        if not section["consolidated"]:
            continue
        size = _page_size(pages[section["start_index"]])
        top, bottom = _furniture_margins(section, size, lang)
        section["render_reserve"] = {"basis": "conservative_text_wrap_estimate", "top_margin_pt": top,
                                     "bottom_margin_pt": bottom, "header_body_gap_pt": 18,
                                     "footer_body_gap_pt": 6, "gap_policy": "readable_section_reserve_v2",
                                     "word_render_verified": False}
        if size[1] - top - bottom >= 144 and size[0] - 2 * Cm(1.7).pt >= 72:
            continue
        section["consolidated"] = False
        section["reserve_rejected"] = True
        plan["review_required"] = True
        for index in range(section["start_index"], section["end_index_exclusive"]):
            page = plan["pages"][index]
            page["adopted_header_ids"] = []
            page["adopted_footer_ids"] = []
            page["review_required"] = True
            page["warnings"].append("section_furniture_reserve_exceeds_page")


def _configure_furniture_section(document, section, plan: dict | None, *, size: tuple[float, float],
                                 lang: TargetLang | str, strip_bidi_controls: bool,
                                 section_index: int) -> dict[str, dict]:
    """Write independent header/footer parts and explicit many-to-one ID aliases."""
    _configure_section(section, size)
    section.different_first_page_header_footer = False
    for name in ("header", "footer", "first_page_header", "first_page_footer", "even_page_header", "even_page_footer"):
        part = getattr(section, name)
        part.is_linked_to_previous = False
        for child in list(part._element):
            part._element.remove(child)
        part.add_paragraph()
    for node in list(section._sectPr.findall(qn("w:pgNumType"))):
        section._sectPr.remove(node)  # PAGE continues; do not reset numbering.
    locations = {}
    parts = (plan or {}).get("parts", {})
    for kind in ("header", "footer"):
        definition = parts.get(kind)
        if not definition:
            continue
        part = getattr(section, kind)
        canonical_locations = {}
        blocks = definition["canonical_blocks"]
        for index, block in enumerate(blocks):
            paragraph = part.paragraphs[0] if index == 0 else part.add_paragraph()
            _format_block_paragraph(paragraph, block, lang)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            _append_block_text(paragraph, block["text"], lang=lang, strip_bidi_controls=strip_bidi_controls,
                               bold=block.get("bold", False), italic=block.get("italic", False))
            canonical_locations[block["id"]] = {"kind": f"section_{kind}", "section_index": section_index,
                                                "section_id": plan["section_id"], "part_uri": str(part.part.partname),
                                                "paragraph_index": index, "canonical_block_id": block["id"]}
        for alias in definition["aliases"]:
            locations[alias["block_id"]] = {
                "location": dict(canonical_locations[alias["canonical_block_id"]]),
                "furniture_alias": {**alias, "part": kind, "source_signature": definition["source_signature"]},
            }
    top, bottom = _furniture_margins(plan, size, lang)
    section.top_margin, section.bottom_margin = Pt(top), Pt(bottom)
    _add_page_number_footer(document, lang, section=section, append=bool(parts.get("footer")))
    return locations


def _start_furniture_section(document):
    """A section boundary replaces a page break; never emit both at one boundary."""
    content = [node for node in document._element.body if node.tag != qn("w:sectPr")]
    return document.add_section(WD_SECTION_START.NEW_PAGE) if content else document.sections[0]


def _apply_source_spacing(document, pages: list[_AssemblyPage], source_map: list[dict], pages_dir: Path) -> None:
    from .document_spacing import infer_page_spacing

    for page, page_map in zip(pages, source_map):
        if not page.structure:
            continue
        sources = _validated_source_pair(page, page, pages_dir)
        if not sources:
            continue
        mappings = {item["block_id"]: item for item in page_map["blocks"]}
        excluded = [identity for identity, item in mappings.items()
                    if item.get("location", {}).get("kind") != "body_paragraph" or item.get("joined_to_block_id")]
        spacing = infer_page_spacing(source_page=sources[0].to_dict(), translated_page=page.structure,
                                     excluded_block_ids=excluded)
        applied = []
        for identity, override in spacing.get("overrides", {}).items():
            current = mappings.get(identity, {}).get("location", {})
            prior = mappings.get(override["previous_block_id"], {}).get("location", {})
            if current.get("kind") != "body_paragraph" or prior.get("kind") != "body_paragraph":
                continue
            if current["paragraph_index"] != prior["paragraph_index"] + 1:
                continue
            paragraph = document.paragraphs[current["paragraph_index"]]
            previous = document.paragraphs[prior["paragraph_index"]]
            after = previous.paragraph_format.space_after
            if after is None:
                after = previous.style.paragraph_format.space_after
            # Word collapses adjacent paragraph spacing to max(before, after),
            # rather than adding them. Subtraction would under-deliver the gap.
            before_pt = max(0, override["desired_gap_pt"])
            paragraph.paragraph_format.space_before = Pt(before_pt)
            applied.append({"block_id": identity, **override, "applied_space_before_pt": before_pt,
                            "effective_gap_lower_bound_pt": max(before_pt, after.pt if after else 0)})
        page_map["source_spacing"] = {"status": spacing["status"], "line_height_evidence": "unavailable", "applied": applied}


def _apply_source_cohesion(document, pages: list[_AssemblyPage], source_map: list[dict],
                           pages_dir: Path, *, lang: TargetLang | str) -> None:
    """Keep two narrowly source-proven groups together, without text/alignment edits.

    This is a compact-body postpass, after final paragraph movement. It never
    crosses a source page, section, region, table, explicit break or joined alias.
    In particular, centered signatures are not subject to the older left-anchor
    Arabic realignment rule, and a list's later items are not chained together.
    """
    paragraphs = document.paragraphs
    for page, page_map in zip(pages, source_map):
        if (not page.structure or page.structure.get("uncertain")
                or page_map.get("layout_status") not in {"flow", "not_provided"}):
            continue
        sources = _validated_source_pair(page, page, pages_dir)
        if not sources:
            continue
        source_blocks = sources[0].to_dict()["blocks"]
        target_blocks = page.structure["blocks"]
        mappings = {item["block_id"]: item for item in page_map["blocks"]}
        excluded = {source["id"] for source, target in zip(source_blocks, target_blocks)
                    if any(block.get("uncertain") or block.get("table_id") or block.get("continuation_of")
                           for block in (source, target))}

        def adjacent(first_id: str, second_id: str):
            first, second = mappings.get(first_id, {}), mappings.get(second_id, {})
            a, b = first.get("location", {}), second.get("location", {})
            if (first_id in excluded or second_id in excluded or first.get("joined_to_block_id")
                    or second.get("joined_to_block_id") or a.get("kind") != "body_paragraph"
                    or b.get("kind") != "body_paragraph" or b["paragraph_index"] != a["paragraph_index"] + 1):
                return None
            prior, current = paragraphs[a["paragraph_index"]], paragraphs[b["paragraph_index"]]
            if (prior._p.getnext() is not current._p or prior._p.xpath(".//w:br|./w:pPr/w:sectPr")
                    or current._p.xpath("./w:pPr/w:sectPr") or current.paragraph_format.page_break_before):
                return None
            return prior, current

        centered = _signature_group_keep_next(target_blocks, lang, excluded_ids=excluded, alignment="center")
        applied = []
        for index, (source, target) in enumerate(zip(source_blocks[:-1], target_blocks[:-1])):
            following, next_target = source_blocks[index + 1], target_blocks[index + 1]
            if following.get("document_start") or next_target.get("document_start"):
                continue
            pair = adjacent(source["id"], following["id"])
            if pair is None:
                continue
            kind = None
            if centered.get(source["id"]) is True and following["id"] in centered:
                kind = "centered_signature"
            else:
                label = source["text"].strip()
                box, next_box = source.get("bbox"), following.get("bbox")
                if (source["role"] == "paragraph" and following["role"] == "list_item"
                        and 1 < len(label) <= 80 and "\n" not in label and label.endswith(":")
                        and 0 < len(_visible_signature(target)) <= 160
                        and re.match(r"^\s*(?:[•-]|[0-9]{1,4}[.)]|[A-Za-z][.)])\s+", following["text"])
                        and box and next_box and box[3] - box[1] <= 30
                        and -2 <= next_box[1] - box[3] <= 24 and abs(next_box[0] - box[0]) <= 24):
                    kind = "short_list_introduction"
            if kind:
                pair[0].paragraph_format.keep_with_next = True
                if kind == "centered_signature" and not centered[following["id"]]:
                    pair[1].paragraph_format.keep_with_next = False
                applied.append({"kind": kind, "block_id": source["id"], "next_block_id": following["id"],
                                "source_sha256": sources[0].source_sha256})
        if applied:
            page_map["source_cohesion"] = applied


def _write_source_map(result_path: Path, pages: list[dict], *, section_furniture: dict | None = None) -> None:
    """Source provenance is not the number of pages Word will actually render."""
    mapping_path = result_path.with_suffix(".source_map.json")
    payload = {
        "version": 1,
        "docx_sha256": hashlib.sha256(result_path.read_bytes()).hexdigest(),
        "source_page_count": len(pages),
        "rendered_page_count": None,
        "layout_review_required": any(page.get("layout_review_required", False) for page in pages),
        "pages": pages,
    }
    if section_furniture is not None:
        payload["section_furniture"] = section_furniture
    temporary = mapping_path.with_name(f"{mapping_path.name}.tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    _fsync_file(temporary)
    os.replace(temporary, mapping_path)


def assemble_docx(
    pages_dir: Path,
    output_path: Path,
    *,
    lang: TargetLang,
    page_breaks: bool,
    up_to_page: int | None = None,
    strip_bidi_controls: bool = True,
    verify_readable: bool = True,
    stats: dict[str, int] | None = None,
) -> Path:
    page_files = sorted(pages_dir.glob("page_*.txt"))
    if up_to_page is not None:
        page_files = [path for path in page_files if int(path.stem.split("_")[1]) <= up_to_page]
    if not page_files:
        raise RuntimeError(f"No page text files available for DOCX assembly: {pages_dir}")

    pages = [_load_assembly_page(path) for path in page_files]
    furniture_plan = _section_furniture_plan(pages, pages_dir, page_breaks=page_breaks)
    _bound_furniture_reserves(furniture_plan, pages, lang)
    furniture_sections = {item["section_id"]: item for item in furniture_plan.get("sections", [])}
    furniture_enabled = not page_breaks and any(item["consolidated"] for item in furniture_sections.values())
    furniture_locations: dict[str, dict] = {}
    active_furniture_id = None
    document = Document()
    if document.paragraphs and document.paragraphs[0].text == "":
        first = document.paragraphs[0]._element
        first.getparent().remove(first)
    _configure_styles(document, lang)
    current_size = _page_size(pages[0])
    _configure_section(document.sections[0], current_size)
    _add_page_number_footer(document, lang)
    source_map: list[dict] = []
    last_paragraph = None
    last_block: dict | None = None
    last_locator: dict | None = None
    body_paragraphs_by_block: dict[str, tuple] = {}
    for page_idx, page in enumerate(pages):
        structure = page.structure or {}
        furniture_page = furniture_plan.get("pages", [])[page_idx] if furniture_plan.get("pages") else {}
        planned_section = furniture_sections.get(furniture_page.get("section_id"))
        admitted = planned_section if planned_section and planned_section["consolidated"] else None
        furniture_id = admitted["section_id"] if admitted else None
        size = _page_size(page)
        dimensions_changed = any(abs(a - b) > 1 for a, b in zip(size, current_size))
        page_boundary = bool(page_idx and (page_breaks or structure.get("document_start") or dimensions_changed))
        furniture_boundary = bool(furniture_enabled and (page_idx == 0 or page_boundary or furniture_id != active_furniture_id))
        if furniture_boundary:
            section = _start_furniture_section(document)
            furniture_locations.update(_configure_furniture_section(
                document, section, admitted, size=size, lang=lang, strip_bidi_controls=strip_bidi_controls,
                section_index=len(document.sections) - 1,
            ))
            page_boundary = bool(page_idx)
            last_paragraph = last_block = last_locator = None
        elif page_boundary:
            _force_page_boundary(document, size, change_size=dimensions_changed)
        active_furniture_id = furniture_id
        current_size = size
        blocks = structure.get("blocks")
        if blocks is None:
            blocks = [{"id": f"legacy_p{page.number:04d}_line{index:04d}", "text": line, "role": "paragraph"}
                      for index, line in enumerate(page.text.split("\n"), 1)]
        layout, layout_status, layout_review, layout_warnings = _validated_layout(structure)
        page_map = {
            "source_page_number": page.number,
            "structure_status": page.structure_status,
            "source_sha256": structure.get("source_sha256"),
            "source_file_sha256": structure.get("source_file_sha256"),
            "source_text_sha256": structure.get("source_text_sha256"),
            "translation_sha256": _text_hash(page.text),
            "page_size_pt": list(size),
            "layout_status": layout_status,
            "layout_review_required": layout_review,
            "layout_warnings": layout_warnings,
            "source_block_coverage_status": "complete" if page.structure else "legacy_txt_only",
            "blocks": [],
        }
        if furniture_enabled:
            page_map["section_index"] = len(document.sections) - 1
            page_map["section_furniture_adopted"] = bool(admitted)
            page_map["section_furniture_id"] = furniture_id
            page_map["target_text_coverage_status"] = "canonical_furniture_with_recorded_aliases" if admitted else "complete"
        # A rejected structured plan remains a review result even if no interval
        # was admitted. Legacy TXT, explicit matching, and separately validated
        # notice regions do not acquire a fictitious furniture warning.
        furniture_review = bool(not page_breaks and page.structure is not None and layout_status != "regions"
                                and furniture_page.get("review_required"))
        if furniture_review:
            page_map["layout_review_required"] = True
            page_map["layout_warnings"] = list(dict.fromkeys(page_map["layout_warnings"] + furniture_page.get("warnings", [])))
        source_map.append(page_map)
        if layout is not None and layout_status == "regions":
            page_map["blocks"] = _render_page_layout(
                document, layout, blocks, lang=lang, strip_bidi_controls=strip_bidi_controls,
                size=size, prior_page_content=bool(page_idx and not page_boundary),
            )
            last_paragraph = last_block = last_locator = None
            continue
        previous = pages[page_idx - 1] if page_idx else None
        can_continue = bool(
            not page_breaks and not page_boundary and previous and previous.number + 1 == page.number
            and previous.structure and previous.structure.get("continuation_to_next")
            and not previous.structure.get("uncertain") and not structure.get("uncertain")
            and structure.get("continuation_from_previous") and not structure.get("document_start")
        )
        bridge = _validated_furniture_bridge(previous, page, pages_dir) if can_continue else None
        signature_groups = _signature_group_keep_next(blocks, lang) if not structure.get("uncertain") else {}
        block_index = 0
        rendered_on_page = False
        while block_index < len(blocks):
            block = blocks[block_index]
            block_index += 1
            role, text = block["role"], block["text"]
            mapping = {"block_id": block["id"], "role": role}
            page_map["blocks"].append(mapping)
            if admitted and block["id"] in furniture_locations:
                mapping.update(furniture_locations[block["id"]])
                continue
            if role == "footer" and _PAGE_LABEL_RE.fullmatch(sanitize_bidi_controls(text).strip()):
                mapping["location"] = {"kind": "generated_footer_page_field"}
                continue
            if role != "table_cell" and not sanitize_bidi_controls(unwrap_internal_placeholders(text)).strip():
                mapping["location"] = {"kind": "empty_block"}
                continue
            document_start = bool(block.get("document_start"))
            if document_start and (rendered_on_page or (page_idx and not page_boundary)):
                if furniture_enabled:
                    section = _start_furniture_section(document)
                    _configure_furniture_section(document, section, None, size=size, lang=lang,
                                                 strip_bidi_controls=strip_bidi_controls,
                                                 section_index=len(document.sections) - 1)
                    last_paragraph = last_block = last_locator = None
                else:
                    _force_page_boundary(document, size)
            if role == "table_cell":
                cells = [block]
                while block_index < len(blocks) and blocks[block_index].get("table_id") == block["table_id"]:
                    cells.append(blocks[block_index])
                    block_index += 1
                rows = max(cell["row"] for cell in cells) + 1
                cols = max(cell["col"] for cell in cells) + 1
                table_index = len(document.tables)
                table = document.add_table(rows=rows, cols=cols)
                geometry_status = _format_table(table, rtl=_is_rtl_target_lang(lang), width_pt=size[0] - Cm(3.4).pt,
                                                cells=cells, source_width_pt=size[0])
                for row in table.rows:
                    for cell in row.cells:
                        _format_block_paragraph(cell.paragraphs[0], {"role": "table_cell"}, lang)
                for index, cell_block in enumerate(cells):
                    cell_paragraph = table.cell(cell_block["row"], cell_block["col"]).paragraphs[0]
                    _format_block_paragraph(cell_paragraph, cell_block, lang)
                    _append_block_text(cell_paragraph, cell_block["text"], lang=lang,
                                       strip_bidi_controls=strip_bidi_controls,
                                       bold=cell_block.get("bold", False), italic=cell_block.get("italic", False))
                    cell_mapping = mapping if index == 0 else {"block_id": cell_block["id"], "role": "table_cell"}
                    cell_mapping["location"] = {"kind": "table_cell", "table_index": table_index,
                                                "row": cell_block["row"], "col": cell_block["col"], "paragraph_index": 0}
                    cell_mapping["table_geometry_status"] = geometry_status
                    if index:
                        page_map["blocks"].append(cell_mapping)
                last_paragraph = last_block = last_locator = None
            else:
                bridge_target = body_paragraphs_by_block.get(bridge["previous_block_id"]) if bridge else None
                bridge_join = bool(
                    bridge_target and block["id"] == bridge["current_block_id"] and not document_start
                    and role == "paragraph" and not block.get("uncertain")
                    and block.get("continuation_of") == bridge["previous_block_id"]
                    and block.get("alignment") == bridge_target[1].get("alignment")
                )
                join = bool(
                    can_continue and not rendered_on_page and not document_start and last_paragraph is not None
                    and last_block and role == last_block["role"] == "paragraph"
                    and not block.get("uncertain") and not last_block.get("uncertain")
                    and block.get("continuation_of") == last_block["id"]
                    and block.get("alignment") == last_block.get("alignment")
                )
                if join or bridge_join:
                    paragraph, joined_block, joined_locator = bridge_target if bridge_join else (last_paragraph, last_block, last_locator)
                    if paragraph.text and text and not paragraph.text[-1].isspace() and not text[0].isspace():
                        _append_block_text(paragraph, " ", lang=lang, strip_bidi_controls=strip_bidi_controls)
                    locator = dict(joined_locator)
                    mapping["joined_to_block_id"] = joined_block["id"]
                    if bridge_join:
                        mapping["continuation_evidence"] = "revalidated_identical_source_furniture"
                else:
                    paragraph = document.add_paragraph("")
                    _format_block_paragraph(paragraph, block, lang,
                                            physical_alignment="right" if block["id"] in signature_groups else None,
                                            readable_body_spacing=not page_breaks and page.structure is not None)
                    if block["id"] in signature_groups:
                        paragraph.paragraph_format.keep_with_next = signature_groups[block["id"]]
                    locator = {"kind": "body_paragraph", "paragraph_index": len(document.paragraphs) - 1}
                _append_block_text(paragraph, text, lang=lang, strip_bidi_controls=strip_bidi_controls,
                                   bold=block.get("bold", False), italic=block.get("italic", False))
                mapping["location"] = locator
                last_paragraph, last_block, last_locator = paragraph, block, locator
                if role == "paragraph":
                    body_paragraphs_by_block[block["id"]] = (paragraph, block, locator)
            rendered_on_page = True
        # Never join across an empty source page or unvalidated legacy evidence.
        if not rendered_on_page or page.structure is None:
            last_paragraph = last_block = last_locator = None

    if not page_breaks:
        _apply_source_spacing(document, pages, source_map, pages_dir)
    _reflow_compact_lists(document, pages, source_map, pages_dir, page_breaks=page_breaks)
    if not page_breaks:
        _apply_source_cohesion(document, pages, source_map, pages_dir, lang=lang)
    body_paragraphs = list(document._element.body.iter(qn("w:p")))
    paragraph_count = sum(bool("".join(node.text or "" for node in paragraph.iter(qn("w:t"))).strip())
                          for paragraph in body_paragraphs)
    if stats is not None:
        stats["paragraph_count"] = len(body_paragraphs)
        stats["run_count"] = len(list(document._element.body.iter(qn("w:r"))))
        stats["page_count"] = len(page_files)  # Legacy source-page count contract.
        stats["source_page_count"] = len(page_files)
        stats["structured_page_count"] = sum(page.structure is not None for page in pages)
        stats["structure_fallback_count"] = sum(page.structure is None for page in pages)
        stats["layout_region_page_count"] = sum(page["layout_status"] == "regions" for page in source_map)
        stats["layout_review_page_count"] = sum(page["layout_review_required"] for page in source_map)
        stats["section_furniture_page_count"] = sum(bool(page.get("section_furniture_adopted")) for page in source_map)
        stats["section_furniture_alias_count"] = sum("furniture_alias" in block for page in source_map for block in page["blocks"])
    result_path = save_document_atomic(document, output_path, verify_readable=verify_readable)
    saved_summary = _verify_docx_visible_content(
        result_path,
        stage="save",
        expected_visible_paragraphs=paragraph_count,
    )
    _remove_compatibility_mode(result_path)
    _verify_docx_visible_content(
        result_path,
        stage="compatibility rewrite",
        expected_visible_paragraphs=paragraph_count,
        baseline=saved_summary,
    )
    furniture_review_recorded = any(any(warning.startswith("section_furniture_") for warning in page["layout_warnings"])
                                    for page in source_map)
    _write_source_map(result_path, source_map, section_furniture=furniture_plan
                      if furniture_enabled or furniture_review_recorded else None)
    return result_path
