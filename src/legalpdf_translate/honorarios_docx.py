"""Deterministic DOCX generator for Requerimento de Honorarios."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from enum import Enum
import re
from pathlib import Path
import unicodedata

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Pt

from .docx_writer import resolve_noncolliding_output_path
from .user_profile import UserProfile

_PT_MONTHS = {
    1: "janeiro",
    2: "fevereiro",
    3: "março",
    4: "abril",
    5: "maio",
    6: "junho",
    7: "julho",
    8: "agosto",
    9: "setembro",
    10: "outubro",
    11: "novembro",
    12: "dezembro",
}
_INVALID_FILENAME_RE = re.compile(r"[^A-Za-z0-9._-]+")
_TRAILING_PREPOSITION_RE = re.compile(r"\b(?:de|do|da|dos|das)\s*$", re.IGNORECASE)
_PREPOSITION_TOKENS = {"de", "do", "da", "dos", "das"}
_INCOMPLETE_LOCATION_TOKENS = {
    "audiencia",
    "central",
    "civel",
    "civil",
    "comarca",
    "comercio",
    "competencia",
    "criminal",
    "execucao",
    "familia",
    "generica",
    "instrucao",
    "juizo",
    "local",
    "menores",
    "ministerio",
    "publico",
    "republica",
    "secao",
    "trabalho",
    "tribunal",
}


class HonorariosKind(str, Enum):
    TRANSLATION = "translation"
    INTERPRETATION = "interpretation"


@dataclass(slots=True)
class HonorariosDraft:
    case_number: str
    word_count: int
    case_entity: str
    case_city: str
    date_pt: str
    profile: UserProfile
    kind: HonorariosKind = HonorariosKind.TRANSLATION
    service_date: str = ""
    service_entity: str = ""
    service_city: str = ""
    use_service_location_in_honorarios: bool = False
    include_transport_sentence_in_honorarios: bool = True
    travel_km_outbound: float = 0.0
    travel_km_return: float = 0.0
    recipient_block: str = ""


def format_portuguese_date(value: date) -> str:
    return f"{value.day:02d} de {_PT_MONTHS[value.month]} de {value.year:04d}"


def build_honorarios_draft(
    *,
    case_number: str,
    word_count: int,
    case_entity: str,
    case_city: str,
    profile: UserProfile,
    today: date | None = None,
) -> HonorariosDraft:
    current_date = today or date.today()
    return HonorariosDraft(
        case_number=case_number.strip(),
        word_count=int(word_count),
        case_entity=case_entity.strip(),
        case_city=case_city.strip(),
        date_pt=format_portuguese_date(current_date),
        profile=profile,
    )


def build_interpretation_honorarios_draft(
    *,
    case_number: str,
    case_entity: str,
    case_city: str,
    service_date: str,
    profile: UserProfile,
    service_entity: str = "",
    service_city: str = "",
    use_service_location_in_honorarios: bool = False,
    include_transport_sentence_in_honorarios: bool = True,
    travel_km_outbound: float = 0.0,
    travel_km_return: float = 0.0,
    recipient_block: str = "",
    today: date | None = None,
) -> HonorariosDraft:
    current_date = today or date.today()
    cleaned_service_date = service_date.strip()
    cleaned_service_city = service_city.strip()
    cleaned_case_city = case_city.strip() or cleaned_service_city
    return HonorariosDraft(
        case_number=case_number.strip(),
        word_count=0,
        case_entity=case_entity.strip(),
        case_city=cleaned_case_city,
        date_pt=format_portuguese_date(current_date),
        profile=profile,
        kind=HonorariosKind.INTERPRETATION,
        service_date=cleaned_service_date,
        service_entity=service_entity.strip(),
        service_city=cleaned_service_city,
        use_service_location_in_honorarios=bool(use_service_location_in_honorarios),
        include_transport_sentence_in_honorarios=bool(include_transport_sentence_in_honorarios),
        travel_km_outbound=float(travel_km_outbound),
        travel_km_return=float(travel_km_return),
        recipient_block=recipient_block.strip(),
    )


def sanitize_case_number_for_filename(case_number: str) -> str:
    cleaned = _INVALID_FILENAME_RE.sub("_", case_number.strip())
    cleaned = cleaned.strip("._-")
    return cleaned or "sem_processo"


def default_honorarios_filename(
    case_number: str,
    today: date | None = None,
    *,
    kind: HonorariosKind = HonorariosKind.TRANSLATION,
) -> str:
    current_date = today or date.today()
    slug = sanitize_case_number_for_filename(case_number)
    if kind == HonorariosKind.INTERPRETATION:
        return f"Requerimento_Honorarios_Interpretacao_{slug}_{current_date:%Y%m%d}.docx"
    return f"Requerimento_Honorarios_{slug}_{current_date:%Y%m%d}.docx"


def _irs_sentence_fragment(value: str) -> str:
    cleaned = value.strip().rstrip(".")
    if cleaned.casefold() == "sem retenção":
        return "não tem retenção de IRS"
    return cleaned


def _irs_sentence_fragment_interpretation(value: str) -> str:
    cleaned = value.strip().rstrip(".")
    if cleaned.casefold() == "sem retenção":
        return "não está sujeito a retenção de IRS"
    return cleaned


def _format_service_date(value: str) -> str:
    cleaned = value.strip()
    if cleaned == "":
        return ""
    try:
        parsed = date.fromisoformat(cleaned)
    except ValueError:
        return cleaned
    return f"{parsed.day:02d}/{parsed.month:02d}/{parsed.year:04d}"


def _format_km_value(value: float) -> str:
    numeric = float(value)
    if numeric.is_integer():
        return str(int(numeric))
    return f"{numeric:.2f}".rstrip("0").rstrip(".").replace(".", ",")


def _normalize_text_for_match(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    without_accents = "".join(char for char in normalized if not unicodedata.combining(char))
    return " ".join(without_accents.split()).casefold()


def _entity_mentions_case_city(case_entity: str, case_city: str) -> bool:
    normalized_entity = _normalize_text_for_match(case_entity)
    normalized_city = _normalize_text_for_match(case_city)
    if not normalized_entity or not normalized_city:
        return False
    city_tokens = [re.escape(token) for token in normalized_city.split()]
    if not city_tokens:
        return False
    city_pattern = re.compile(r"\b" + r"\s+".join(city_tokens) + r"\b")
    return bool(city_pattern.search(normalized_entity))


def _entity_has_explicit_location_suffix(case_entity: str) -> bool:
    normalized_tokens = _normalize_text_for_match(case_entity).split()
    if len(normalized_tokens) < 2:
        return False
    if normalized_tokens[-1] in _INCOMPLETE_LOCATION_TOKENS:
        return False
    trailing_context = normalized_tokens[-4:-1]
    return any(token in _PREPOSITION_TOKENS for token in trailing_context)


def _complete_case_entity_with_city(case_entity: str, case_city: str) -> str:
    cleaned_entity = " ".join(case_entity.split())
    cleaned_city = " ".join(case_city.split())
    if not cleaned_entity or not cleaned_city:
        return cleaned_entity
    if _entity_mentions_case_city(cleaned_entity, cleaned_city):
        return cleaned_entity
    if _entity_has_explicit_location_suffix(cleaned_entity):
        return cleaned_entity
    separator = " " if _TRAILING_PREPOSITION_RE.search(cleaned_entity) else " de "
    return f"{cleaned_entity}{separator}{cleaned_city}"


def default_interpretation_recipient_block(case_entity: str, case_city: str = "") -> str:
    cleaned = _complete_case_entity_with_city(case_entity, case_city)
    cleaned_city = case_city.strip()
    if cleaned.casefold().startswith("ministério público de ".casefold()):
        return f"Exmo. Senhor Procurador do {cleaned}"
    if cleaned.casefold() == "ministério público":
        if cleaned_city:
            return f"Exmo. Senhor Procurador do Ministério Público de {cleaned_city}"
        return "Exmo. Senhor Procurador do Ministério Público"
    if cleaned:
        return f"Exmo. Senhor Procurador da República do {cleaned}"
    if cleaned_city:
        return f"Exmo. Senhor Procurador do Ministério Público de {cleaned_city}"
    return "Exmo. Senhor Procurador da República"


def _interpretation_service_location_phrase(draft: HonorariosDraft) -> str:
    if not draft.use_service_location_in_honorarios:
        return ""
    city = draft.service_city.strip()
    if city == "":
        return ""
    entity = draft.service_entity.strip()
    entity_norm = entity.casefold()
    if entity_norm in {"gnr", "psp"}:
        return f"na {entity} de {city}"
    return f"na cidade de {city}"


def _interpretation_travel_destination(draft: HonorariosDraft) -> str:
    if draft.service_city.strip():
        return draft.service_city.strip()
    return draft.case_city.strip()


def _interpretation_one_way_distance(draft: HonorariosDraft) -> float:
    if float(draft.travel_km_outbound) > 0:
        return float(draft.travel_km_outbound)
    return float(draft.travel_km_return)


def _translation_recipient_line(draft: HonorariosDraft) -> str:
    entity = _complete_case_entity_with_city(draft.case_entity, draft.case_city)
    if entity.casefold() == "ministério público":
        return "Exmo. Sr(a). Procurador(a) da república do Ministério Público"
    return "Exmo. Sr(a). Procurador(a) da república do " + entity


def _translation_paragraph_texts(draft: HonorariosDraft) -> list[tuple[str, str]]:
    profile = draft.profile
    return [
        (f"Número de processo: {draft.case_number}", "left"),
        ("", "left"),
        (_translation_recipient_line(draft), "address"),
        ("", "left"),
        (f"Nome: {profile.document_name}", "left"),
        (f"Morada: {profile.postal_address}", "left"),
        ("", "left"),
        (
            "Venho por este meio requerer o pagamento dos honorários devidos, em virtude de ter sido nomeado "
            "tradutor no âmbito do processo acima identificado.",
            "left",
        ),
        (f"O documento traduzido contém {draft.word_count} palavras.", "left"),
        (
            f"Este serviço inclui a taxa IVA de {profile.iva_text} e {_irs_sentence_fragment(profile.irs_text)}.",
            "left",
        ),
        (f"O Pagamento deverá ser efetuado para o seguinte IBAN: {profile.iban}", "left"),
        ("", "left"),
        ("Melhores cumprimentos,", "left"),
        ("", "left"),
        ("Espera deferimento,", "center"),
        ("", "left"),
        (f"{draft.case_city}, {draft.date_pt}", "center"),
        ("", "left"),
        (profile.document_name, "center"),
    ]


def _interpretation_paragraph_texts(draft: HonorariosDraft) -> list[tuple[str, str]]:
    profile = draft.profile
    recipient_block = draft.recipient_block.strip() or default_interpretation_recipient_block(
        draft.case_entity,
        draft.case_city,
    )
    recipient_lines = [line.strip() for line in recipient_block.splitlines() if line.strip()]
    location_phrase = _interpretation_service_location_phrase(draft)
    date_phrase = f"no dia {_format_service_date(draft.service_date)}"
    if location_phrase:
        date_phrase = f"{date_phrase}, {location_phrase}"
    body_text = (
        "Venho, por este meio, requerer o pagamento dos honorários devidos, em virtude de ter sido nomeado "
        f"intérprete no âmbito do processo acima identificado, {date_phrase}."
    )
    if draft.include_transport_sentence_in_honorarios:
        travel_destination = _interpretation_travel_destination(draft)
        body_text = (
            body_text[:-1]
            + ", bem como o pagamento das despesas de transporte entre "
            f"{profile.travel_origin_label} e {travel_destination}, tendo percorrido "
            f"{_format_km_value(_interpretation_one_way_distance(draft))} km em cada sentido."
        )
    paragraphs: list[tuple[str, str]] = [
        (f"Número de processo: {draft.case_number}", "left"),
        ("", "left"),
    ]
    for line in recipient_lines:
        paragraphs.append((line, "address"))
    paragraphs.extend(
        [
            ("", "left"),
            (f"Nome: {profile.document_name}", "left"),
            (f"Morada: {profile.postal_address}", "left"),
            ("", "left"),
            (body_text, "left"),
            (
                f"Este serviço inclui a taxa de IVA de {profile.iva_text} e {_irs_sentence_fragment_interpretation(profile.irs_text)}.",
                "left",
            ),
            ("", "left"),
            (f"O Pagamento deverá ser efetuado para o seguinte IBAN: {profile.iban}", "left"),
            ("", "left"),
            ("Melhores cumprimentos,", "left"),
            ("", "left"),
            ("Espera deferimento,", "center"),
            ("", "left"),
            (f"{draft.case_city}, {draft.date_pt}", "center"),
            ("", "left"),
            ("O Requerente,", "center"),
            ("", "left"),
            (profile.document_name, "center"),
        ]
    )
    return paragraphs


def build_honorarios_paragraph_texts(draft: HonorariosDraft) -> list[tuple[str, str]]:
    if draft.kind == HonorariosKind.INTERPRETATION:
        return _interpretation_paragraph_texts(draft)
    return _translation_paragraph_texts(draft)


def _set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)


def _configure_page(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def generate_honorarios_docx(draft: HonorariosDraft, output_path: Path) -> Path:
    requested_output = output_path.expanduser().resolve()
    requested_output.parent.mkdir(parents=True, exist_ok=True)
    if requested_output.exists():
        fallback_output = requested_output.parent / default_honorarios_filename(
            draft.case_number,
            kind=draft.kind,
        )
        output = resolve_noncolliding_output_path(fallback_output)
    else:
        output = requested_output

    document = Document()
    _configure_page(document)
    _set_default_font(document)

    for text, kind in build_honorarios_paragraph_texts(draft):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        if kind == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "address":
            paragraph.paragraph_format.left_indent = Cm(9.0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12)

    document.save(output)
    return output
