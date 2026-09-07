"""Windows Word automation helpers used by the Qt app."""

from __future__ import annotations

from dataclasses import dataclass, replace
import json
import math
import re
from datetime import UTC, datetime
import os
from pathlib import Path
import shutil
from shutil import which
import subprocess
from tempfile import mkdtemp
import time

from docx import Document

_WORD_PHASE_PREFIX = "LEGALPDF_WORD_PHASE:"
_WORD_HELPER_PID_PREFIX = "LEGALPDF_WORD_HELPER_PID:"
_WORD_HELPER_OWNER_PREFIX = "LEGALPDF_WORD_HELPER_OWNER:"
_WORD_HELPER_OWNER = "app_owned"
_WORD_READINESS_CACHE_TTL_SECONDS = 60.0
_WORD_READINESS_CACHE: dict[str, tuple[float, dict[str, object]]] = {}


@dataclass(frozen=True, slots=True)
class WordAutomationResult:
    ok: bool
    action: str
    message: str
    stdout: str = ""
    stderr: str = ""
    command: tuple[str, ...] = ()
    failure_code: str = ""
    details: str = ""
    elapsed_ms: int = 0
    failure_phase: str = ""
    helper_pid: int = 0
    cleanup_attempted: bool = False
    cleanup_succeeded: bool = False
    cleanup_details: str = ""


def _is_windows_host() -> bool:
    return os.name == "nt"


def _resolve_powershell_path() -> str | None:
    if not _is_windows_host():
        return None
    system_root = os.environ.get("SystemRoot", r"C:\Windows").strip() or r"C:\Windows"
    preferred = Path(system_root) / "System32" / "WindowsPowerShell" / "v1.0" / "powershell.exe"
    if preferred.exists():
        return str(preferred)
    discovered = which("powershell.exe") or which("powershell")
    if discovered:
        return discovered
    return None


def _resolve_winword_path() -> str | None:
    if not _is_windows_host():
        return None
    candidates: list[Path] = []
    for env_name in ("ProgramFiles", "ProgramW6432", "ProgramFiles(x86)"):
        root = str(os.environ.get(env_name, "") or "").strip()
        if root == "":
            continue
        base = Path(root)
        candidates.extend(
            [
                base / "Microsoft Office" / "root" / "Office16" / "WINWORD.EXE",
                base / "Microsoft Office" / "Office16" / "WINWORD.EXE",
                base / "Microsoft Office" / "root" / "Office15" / "WINWORD.EXE",
                base / "Microsoft Office" / "Office15" / "WINWORD.EXE",
            ]
        )
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    discovered = which("WINWORD.EXE") or which("winword.exe") or which("winword")
    if discovered:
        return discovered
    return None


def _quote_powershell_single(value: str) -> str:
    return value.replace("'", "''")


def _word_phase_marker(phase: str) -> str:
    return f"Write-Output '{_WORD_PHASE_PREFIX}{_quote_powershell_single(phase)}'"


def _word_helper_header() -> list[str]:
    return [
        f"Write-Output '{_WORD_HELPER_OWNER_PREFIX}{_WORD_HELPER_OWNER}'",
        f'Write-Output ("{_WORD_HELPER_PID_PREFIX}" + $PID)',
    ]


def _normalize_process_text(value: str | bytes | None) -> str:
    if value is None:
        return ""
    if isinstance(value, bytes):
        return value.decode("utf-8", errors="replace").strip()
    return str(value).strip()


def _strip_word_helper_markers(value: str) -> str:
    lines = []
    for raw_line in str(value or "").splitlines():
        line = raw_line.strip()
        if (
            line.startswith(_WORD_PHASE_PREFIX)
            or line.startswith(_WORD_HELPER_PID_PREFIX)
            or line.startswith(_WORD_HELPER_OWNER_PREFIX)
        ):
            continue
        lines.append(raw_line)
    return "\n".join(lines).strip()


def _extract_last_word_phase(*texts: str) -> str:
    last_phase = ""
    for text in texts:
        for raw_line in str(text or "").splitlines():
            line = raw_line.strip()
            if line.startswith(_WORD_PHASE_PREFIX):
                last_phase = line[len(_WORD_PHASE_PREFIX) :].strip()
    return last_phase


def _extract_word_helper_pid(*texts: str) -> int:
    for text in texts:
        for raw_line in str(text or "").splitlines():
            line = raw_line.strip()
            if not line.startswith(_WORD_HELPER_PID_PREFIX):
                continue
            raw_value = line[len(_WORD_HELPER_PID_PREFIX) :].strip()
            try:
                return int(raw_value)
            except (TypeError, ValueError):
                continue
    return 0


def _build_word_powershell_script(docx_path: Path, *, align_right_and_save: bool) -> str:
    resolved = str(docx_path.expanduser().resolve())
    quoted = _quote_powershell_single(resolved)
    action_block = ""
    if align_right_and_save:
        action_block = "\n".join(
            [
                "$doc.Range().ParagraphFormat.Alignment = 2",
                "$doc.Save()",
            ]
        )
    return "\n".join(
        [
            "$ErrorActionPreference = 'Stop'",
            *_word_helper_header(),
            f"$target = [System.IO.Path]::GetFullPath('{quoted}')",
            "$word = $null",
            _word_phase_marker("get_active_word"),
            "try {",
            "    $word = [Runtime.InteropServices.Marshal]::GetActiveObject('Word.Application')",
            "} catch {",
            "}",
            "if ($null -eq $word) {",
            f"    {_word_phase_marker('launch_word')}",
            "    $word = New-Object -ComObject Word.Application",
            "}",
            "$word.Visible = $true",
            "$doc = $null",
            _word_phase_marker("scan_documents"),
            "foreach ($candidate in @($word.Documents)) {",
            "    if ([string]::Equals($candidate.FullName, $target, [System.StringComparison]::OrdinalIgnoreCase)) {",
            "        $doc = $candidate",
            "        break",
            "    }",
            "}",
            "if ($null -eq $doc) {",
            f"    {_word_phase_marker('open_document')}",
            "    $doc = $word.Documents.Open($target)",
            "}",
            _word_phase_marker("activate_document"),
            "$doc.Activate()",
            "$word.Activate()",
            action_block,
            _word_phase_marker("complete"),
            "Write-Output 'OK'",
        ]
    )


def _build_powershell_command(docx_path: Path, *, align_right_and_save: bool) -> tuple[str, ...] | None:
    powershell = _resolve_powershell_path()
    if powershell is None:
        return None
    script = _build_word_powershell_script(docx_path, align_right_and_save=align_right_and_save)
    return (
        powershell,
        "-Sta",
        "-NoLogo",
        "-NoProfile",
        "-ExecutionPolicy",
        "Bypass",
        "-Command",
        script,
    )


def _build_pdf_export_powershell_script(docx_path: Path, pdf_path: Path, *, state_path: Path | None = None) -> str:
    from .word_pdf_script import build_pdf_script
    from .word_pdf_control import runtime_root

    executable = _resolve_winword_path()
    return build_pdf_script(docx_path, pdf_path, state_path=state_path or runtime_root() / "active.json",
                            word_executable=Path(executable) if executable else None)


def _build_pdf_preflight_powershell_script(*, state_path: Path | None = None) -> str:
    from .word_pdf_script import build_pdf_script
    from .word_pdf_control import runtime_root

    executable = _resolve_winword_path()
    return build_pdf_script(None, None, state_path=state_path or runtime_root() / "active.json",
                            word_executable=Path(executable) if executable else None)


def _classify_pdf_failure(raw_message: str, *, action: str) -> str:
    lowered = raw_message.casefold()
    if "timed out" in lowered:
        return "timeout"
    if (
        "powershell is unavailable" in lowered
        or "powershell unavailable" in lowered
        or "powershell.exe is unavailable" in lowered
    ):
        return "powershell_missing"
    if (
        "class not registered" in lowered
        or "invalid class string" in lowered
        or "cannot create activex component" in lowered
        or "no com class identified" in lowered
    ):
        return "word_unavailable"
    if (
        "server execution failed" in lowered
        or "co_e_server_exec_failure" in lowered
        or "0x80080005" in lowered
        or "retrieving the com class factory" in lowered
    ):
        return "com_launch_failed"
    if action == "export_pdf":
        return "export_failed"
    return "unknown"


def _pdf_failure_message(failure_code: str) -> str:
    return {
        "powershell_missing": "PowerShell is unavailable for Word PDF export.",
        "word_unavailable": "Microsoft Word is unavailable for PDF export on this computer.",
        "com_launch_failed": "Microsoft Word could not be started for PDF export.",
        "timeout": "Word PDF export timed out.",
        "export_failed": "Microsoft Word could not export the PDF.",
        "verification_failed": "Word PDF export verification failed.",
        "export_busy": "Another Word PDF operation is running. Try again when it finishes.",
        "cleanup_unconfirmed": "Word export cleanup could not be confirmed. Save your Word work before closing Word; see details for recovery.",
        "ownership_rejected": "A separate Word export instance could not be verified.",
        "ownership_unproven": "A separate Word export instance could not be verified.",
        "word_window_unavailable": "Word started, but its export window could not be connected safely.",
        "bootstrap_changed": "Word's temporary export document was not blank or changed. It was not closed automatically.",
        "cleanup_ambiguous": "Word export cleanup is uncertain. Save your Word work before closing Word and retrying.",
        "host_policy_blocked": "Windows policy blocked the Word export helper. Export the saved DOCX manually; no security settings were changed.",
        "unknown": "Word PDF export failed.",
    }.get(failure_code, "Word PDF export failed.")


def _build_pdf_failure_details(
    *,
    failure_code: str,
    elapsed_ms: int,
    raw_message: str,
    stdout: str,
    stderr: str,
    command: tuple[str, ...],
    failure_phase: str = "",
    helper_pid: int = 0,
    cleanup_attempted: bool = False,
    cleanup_succeeded: bool = False,
    cleanup_details: str = "",
) -> str:
    lines = [
        f"Failure code: {failure_code or 'unknown'}",
        f"Elapsed: {elapsed_ms} ms",
    ]
    if failure_phase:
        lines.append(f"Failure phase: {failure_phase}")
    if helper_pid > 0:
        lines.append(f"Helper PID: {helper_pid}")
    if cleanup_attempted:
        lines.append(f"Cleanup attempted: yes; succeeded: {'yes' if cleanup_succeeded else 'no'}")
    if cleanup_details:
        lines.extend(["", "Cleanup:", cleanup_details])
    if command:
        lines.append(f"Host command: {command[0]}")
    if raw_message:
        lines.extend(["", "Raw diagnostic:", raw_message])
    if stderr and stderr != raw_message:
        lines.extend(["", "stderr:", stderr])
    if stdout and stdout != raw_message:
        lines.extend(["", "stdout:", stdout])
    return "\n".join(lines).strip()


def _build_pdf_failure_result(
    *,
    action: str,
    raw_message: str,
    stdout: str = "",
    stderr: str = "",
    command: tuple[str, ...] = (),
    elapsed_ms: int = 0,
    failure_code: str | None = None,
    failure_phase: str = "",
    helper_pid: int = 0,
    cleanup_attempted: bool = False,
    cleanup_succeeded: bool = False,
    cleanup_details: str = "",
) -> WordAutomationResult:
    resolved_code = failure_code or _classify_pdf_failure(raw_message, action=action)
    return WordAutomationResult(
        ok=False,
        action=action,
        message=_pdf_failure_message(resolved_code),
        stdout=stdout,
        stderr=stderr,
        command=command,
        failure_code=resolved_code,
        details=_build_pdf_failure_details(
            failure_code=resolved_code,
            elapsed_ms=elapsed_ms,
            raw_message=raw_message,
            stdout=stdout,
            stderr=stderr,
            command=command,
            failure_phase=failure_phase,
            helper_pid=helper_pid,
            cleanup_attempted=cleanup_attempted,
            cleanup_succeeded=cleanup_succeeded,
            cleanup_details=cleanup_details,
        ),
        elapsed_ms=elapsed_ms,
        failure_phase=failure_phase,
        helper_pid=helper_pid,
        cleanup_attempted=cleanup_attempted,
        cleanup_succeeded=cleanup_succeeded,
        cleanup_details=cleanup_details,
    )


def _resolve_taskkill_path() -> str | None:
    if not _is_windows_host():
        return None
    system_root = os.environ.get("SystemRoot", r"C:\Windows").strip() or r"C:\Windows"
    preferred = Path(system_root) / "System32" / "taskkill.exe"
    if preferred.exists():
        return str(preferred)
    discovered = which("taskkill.exe") or which("taskkill")
    if discovered:
        return discovered
    return None


def _terminate_windows_process_tree(pid: int) -> tuple[bool, str]:
    if not _is_windows_host() or int(pid) <= 0:
        return False, "Process-tree cleanup is only available on Windows."
    taskkill = _resolve_taskkill_path()
    if taskkill is None:
        return False, "taskkill.exe is unavailable."
    creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    try:
        completed = subprocess.run(
            [taskkill, "/PID", str(int(pid)), "/T", "/F"],
            capture_output=True,
            text=True,
            check=False,
            timeout=10,
            creationflags=creationflags,
        )
    except subprocess.TimeoutExpired as exc:
        return False, f"taskkill timed out after {exc.timeout} seconds."
    stdout = _normalize_process_text(completed.stdout)
    stderr = _normalize_process_text(completed.stderr)
    lowered = f"{stdout}\n{stderr}".casefold()
    success = completed.returncode == 0 or "not found" in lowered or "no running instance" in lowered
    details = stderr or stdout or "taskkill returned no output."
    return success, details


def _cleanup_working_tree_best_effort(
    path: Path,
    *,
    retries: int = 10,
    delay_seconds: float = 0.2,
) -> None:
    for attempt in range(max(1, int(retries))):
        try:
            shutil.rmtree(path)
            return
        except FileNotFoundError:
            return
        except OSError:
            if attempt >= max(1, int(retries)) - 1:
                return
            time.sleep(max(0.0, float(delay_seconds)))


def serialize_word_automation_result(result: WordAutomationResult) -> dict[str, object]:
    return {
        "ok": bool(result.ok),
        "action": str(result.action or "").strip(),
        "message": str(result.message or "").strip(),
        "stdout": str(result.stdout or "").strip(),
        "stderr": str(result.stderr or "").strip(),
        "command": list(result.command),
        "failure_code": str(result.failure_code or "").strip(),
        "details": str(result.details or "").strip(),
        "elapsed_ms": int(result.elapsed_ms),
        "failure_phase": str(result.failure_phase or "").strip(),
        "helper_pid": int(result.helper_pid or 0),
        "cleanup_attempted": bool(result.cleanup_attempted),
        "cleanup_succeeded": bool(result.cleanup_succeeded),
        "cleanup_details": str(result.cleanup_details or "").strip(),
    }


def _run_command(
    *,
    action: str,
    command: tuple[str, ...],
    success_message: str,
    timeout_seconds: float = 20,
) -> WordAutomationResult:
    creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    started_at = time.perf_counter()
    try:
        process = subprocess.Popen(
            list(command),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            creationflags=creationflags,
        )
        stdout_raw, stderr_raw = process.communicate(timeout=timeout_seconds)
    except subprocess.TimeoutExpired as exc:
        elapsed_ms = int((time.perf_counter() - started_at) * 1000)
        stdout_raw = _normalize_process_text(getattr(exc, "stdout", ""))
        stderr_raw = _normalize_process_text(getattr(exc, "stderr", ""))
        helper_pid = _extract_word_helper_pid(stdout_raw, stderr_raw) or int(getattr(process, "pid", 0) or 0)
        failure_phase = _extract_last_word_phase(stdout_raw, stderr_raw)
        cleanup_attempted = False
        cleanup_succeeded = False
        cleanup_details = ""
        if action in {"export_pdf", "pdf_preflight"} and helper_pid > 0 and _is_windows_host():
            cleanup_attempted = True
            cleanup_succeeded, cleanup_details = _terminate_windows_process_tree(helper_pid)
        if getattr(process, "poll", lambda: None)() is None:
            try:
                process.kill()
            except Exception:  # noqa: BLE001
                pass
        try:
            extra_stdout, extra_stderr = process.communicate(timeout=2)
            stdout_raw = "\n".join(filter(None, [stdout_raw, _normalize_process_text(extra_stdout)])).strip()
            stderr_raw = "\n".join(filter(None, [stderr_raw, _normalize_process_text(extra_stderr)])).strip()
        except Exception:  # noqa: BLE001
            pass
        failure_phase = _extract_last_word_phase(stdout_raw, stderr_raw) or failure_phase
        stdout = _strip_word_helper_markers(stdout_raw)
        stderr = _strip_word_helper_markers(stderr_raw)
        raw_message = stderr or stdout or str(exc)
        if action in {"export_pdf", "pdf_preflight"}:
            return _build_pdf_failure_result(
                action=action,
                raw_message=raw_message,
                stdout=stdout,
                stderr=stderr,
                command=command,
                elapsed_ms=elapsed_ms,
                failure_code="timeout",
                failure_phase=failure_phase,
                helper_pid=helper_pid,
                cleanup_attempted=cleanup_attempted,
                cleanup_succeeded=cleanup_succeeded,
                cleanup_details=cleanup_details,
            )
        return WordAutomationResult(
            ok=False,
            action=action,
            message=str(exc),
            stdout=stdout,
            stderr=stderr,
            command=command,
            elapsed_ms=elapsed_ms,
            failure_phase=failure_phase,
            helper_pid=helper_pid,
            cleanup_attempted=cleanup_attempted,
            cleanup_succeeded=cleanup_succeeded,
            cleanup_details=cleanup_details,
        )
    except Exception as exc:  # noqa: BLE001
        elapsed_ms = int((time.perf_counter() - started_at) * 1000)
        raw_message = str(exc)
        if action in {"export_pdf", "pdf_preflight"}:
            return _build_pdf_failure_result(
                action=action,
                raw_message=raw_message,
                command=command,
                elapsed_ms=elapsed_ms,
            )
        return WordAutomationResult(
            ok=False,
            action=action,
            message=raw_message,
            command=command,
            elapsed_ms=elapsed_ms,
        )
    elapsed_ms = int((time.perf_counter() - started_at) * 1000)
    stdout_raw = _normalize_process_text(stdout_raw)
    stderr_raw = _normalize_process_text(stderr_raw)
    stdout = _strip_word_helper_markers(stdout_raw)
    stderr = _strip_word_helper_markers(stderr_raw)
    helper_pid = _extract_word_helper_pid(stdout_raw, stderr_raw) or int(getattr(process, "pid", 0) or 0)
    failure_phase = _extract_last_word_phase(stdout_raw, stderr_raw)
    if process.returncode != 0:
        raw_message = stderr or stdout or "Word automation failed."
        if action in {"export_pdf", "pdf_preflight"}:
            return _build_pdf_failure_result(
                action=action,
                raw_message=raw_message,
                stdout=stdout,
                stderr=stderr,
                command=command,
                elapsed_ms=elapsed_ms,
                failure_phase=failure_phase,
                helper_pid=helper_pid,
            )
        return WordAutomationResult(
            ok=False,
            action=action,
            message=raw_message,
            stdout=stdout,
            stderr=stderr,
            command=command,
            elapsed_ms=elapsed_ms,
            failure_phase=failure_phase,
            helper_pid=helper_pid,
        )
    return WordAutomationResult(
        ok=True,
        action=action,
        message=success_message,
        stdout=stdout,
        stderr=stderr,
        command=command,
        elapsed_ms=elapsed_ms,
        failure_phase=failure_phase,
        helper_pid=helper_pid,
    )


def _run_word_action(docx_path: Path, *, align_right_and_save: bool) -> WordAutomationResult:
    action = "align_right_and_save" if align_right_and_save else "open"
    resolved = docx_path.expanduser().resolve()
    if not _is_windows_host():
        return WordAutomationResult(
            ok=False,
            action=action,
            message="Word automation is available only on Windows.",
        )
    if not resolved.exists():
        return WordAutomationResult(
            ok=False,
            action=action,
            message=f"DOCX not found: {resolved}",
        )
    command = _build_powershell_command(resolved, align_right_and_save=align_right_and_save)
    if command is None:
        return WordAutomationResult(
            ok=False,
            action=action,
            message="PowerShell is unavailable for Word automation.",
        )
    success_message = "Word document opened." if not align_right_and_save else "Word document aligned right and saved."
    return _run_command(
        action=action,
        command=command,
        success_message=success_message,
    )


def open_docx_in_word(docx_path: Path) -> WordAutomationResult:
    return _run_word_action(docx_path, align_right_and_save=False)


def align_right_and_save_docx_in_word(docx_path: Path) -> WordAutomationResult:
    return _run_word_action(docx_path, align_right_and_save=True)


def _safe_word_token(value: object) -> str:
    text = str(value or "")
    return text if re.fullmatch(r"[a-z_]{1,64}", text) else ""


def _run_pdf_command(
    *, action: str, command: tuple[str, ...], state_path: Path, timeout_seconds: float,
) -> WordAutomationResult:
    """Run only the owned helper; never kill a Word server or its process tree.

    PowerShell persists bounded, content-free state before COM calls, so Windows
    pipe buffering and cleanup failures cannot erase the primary failure.
    """
    from .word_pdf_control import read_state

    started = time.perf_counter()
    timed_out = False
    helper_io_failed = False
    process = None
    helper_stopped = False
    stderr = b""
    try:
        process = subprocess.Popen(
            list(command), stdout=subprocess.DEVNULL, stderr=subprocess.PIPE,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
        try:
            _, stderr = process.communicate(timeout=timeout_seconds)
        except (subprocess.TimeoutExpired, OSError, ValueError) as exc:
            timed_out = isinstance(exc, subprocess.TimeoutExpired)
            helper_io_failed = not timed_out
            # Popen.kill on Windows uses our retained process handle, not a
            # textual PID possibly recycled into another user process.
            try:
                if process.poll() is None:
                    process.kill()
                _, stderr = process.communicate(timeout=2)
                helper_stopped = True
            except (subprocess.TimeoutExpired, OSError, ValueError):
                pass
    except (OSError, ValueError):
        if process is None:
            state_path.write_text(json.dumps({
                "status": "failed", "cleanup_status": "confirmed",
                "phase": "start_helper", "failure_code": "com_launch_failed",
                "launch_attempted": False,
            }), encoding="utf-8")
        else:
            timed_out = True
    state = read_state(state_path)
    if helper_stopped and state.get("status") != "invalid":
        # Recovery can prove our helper is gone even if it timed out before
        # recording its PID. This does NOT claim the Word server has exited.
        state["parent_helper_stopped"] = True
        state_path.write_text(json.dumps(state), encoding="utf-8")
    if not timed_out and not helper_io_failed and process is not None and state.get("status") == "starting":
        # The script must persist a phase successfully BEFORE every COM call.
        # A terminal helper that never wrote its initial state never reached
        # activation. Do not quarantine a PowerShell setup/parse failure as a
        # potentially live Word operation. Raw stderr is never user-visible.
        diagnostic = _normalize_process_text(stderr)
        code = _classify_pdf_failure(diagnostic, action=action)
        if any(token in diagnostic.casefold() for token in ("constrainedlanguage", "constrained language", "application control")):
            code = "host_policy_blocked"
        state = {"status":"failed", "cleanup_status":"confirmed", "launch_attempted":False,
                 "phase":"helper_setup", "failure_code":code}
        state_path.write_text(json.dumps(state), encoding="utf-8")
    elapsed = int((time.perf_counter() - started) * 1000)
    phase = _safe_word_token(state.get("primary_failure_phase")) or _safe_word_token(state.get("phase"))
    cleanup = _safe_word_token(state.get("cleanup_status"))
    status = _safe_word_token(state.get("status"))
    code = _safe_word_token(state.get("failure_code"))
    confirmed = cleanup == "confirmed"
    cleanup_summary = ("No Word activation occurred." if state.get("launch_attempted") is False
                       else "Owned Word instance exit confirmed.") if confirmed else (
                           "Word cleanup is unconfirmed; existing user sessions were not terminated.")
    ok = bool(process is not None and not timed_out and not helper_io_failed and process.returncode == 0
              and status == "succeeded" and confirmed)
    details = [
        f"Phase: {phase or 'unknown'}",
        f"Word ownership: {_safe_word_token(state.get('ownership')) or 'unknown'}",
        f"Word cleanup: {cleanup or 'unconfirmed'}",
    ]
    if process is not None and not timed_out and not helper_io_failed:
        details.append(f"Helper exit code: {process.returncode}")
    hresult = str(state.get("primary_hresult") or "")
    if re.fullmatch(r"0x[0-9A-Fa-f]{8}", hresult):
        details.append(f"HRESULT: {hresult}")
    if timed_out or helper_io_failed:
        details.append(f"Owned helper stopped: {'yes' if helper_stopped else 'unconfirmed'}")
        details.append("No Word process was force-terminated; retry requires confirmed cleanup.")
    if ok:
        return WordAutomationResult(
            ok=True, action=action,
            message="Word PDF export preflight passed." if action == "pdf_preflight" else "Word document exported to PDF.",
            command=command[:1], details="\n".join(details), elapsed_ms=elapsed,
            helper_pid=int(process.pid), cleanup_attempted=True, cleanup_succeeded=True,
            cleanup_details=cleanup_summary,
        )
    if timed_out:
        code = "timeout"
    elif helper_io_failed:
        code = "export_failed"
    elif not code:
        code = "cleanup_unconfirmed" if not confirmed else "export_failed"
    result = _build_pdf_failure_result(
        action=action, raw_message="\n".join(details), command=command[:1],
        elapsed_ms=elapsed, failure_code=code, failure_phase=phase,
        helper_pid=int(process.pid) if process is not None else 0,
        cleanup_attempted=cleanup not in {"", "not_started"} or timed_out or helper_io_failed,
        cleanup_succeeded=confirmed and not timed_out and not helper_io_failed,
        cleanup_details=cleanup_summary,
    )
    return result


class _WordExportAborted(Exception):
    def __init__(self, result: WordAutomationResult):
        self.result = result


def _run_pdf_operation(
    *, action: str, docx_path: Path | None = None, pdf_path: Path | None = None,
    timeout_seconds: float,
) -> WordAutomationResult:
    from .word_pdf_control import WordPdfBusy, WordPdfRecoveryRequired, word_pdf_slot
    from .word_pdf_artifacts import staged_pdf_export

    if not _is_windows_host():
        return WordAutomationResult(ok=False, action=action, message="Word automation is available only on Windows.")
    powershell = _resolve_powershell_path()
    if powershell is None:
        return _build_pdf_failure_result(action=action, raw_message="PowerShell is unavailable.", failure_code="powershell_missing")
    if not math.isfinite(timeout_seconds) or timeout_seconds <= 0:
        return _build_pdf_failure_result(action=action, raw_message="A positive finite export deadline is required.", failure_code="export_failed")
    started = time.perf_counter()
    result = None
    artifact_published = False
    try:
        with word_pdf_slot() as state_path:
            def dispatch(source=None, target=None):
                remaining = timeout_seconds - (time.perf_counter() - started)
                if remaining <= 0:
                    return _build_pdf_failure_result(action=action, raw_message="Deadline expired before Word startup.", failure_code="timeout", failure_phase="prepare_files")
                script = (_build_pdf_preflight_powershell_script(state_path=state_path) if source is None
                          else _build_pdf_export_powershell_script(source, target, state_path=state_path))
                script_path = state_path.with_suffix(".ps1")
                # UTF-8 BOM is required by Windows PowerShell 5.1 for accents;
                # -File also avoids the Windows command-length ceiling.
                script_path.write_text(script, encoding="utf-8-sig")
                state_path.write_text(json.dumps({"status":"starting", "launch_attempted":True}), encoding="utf-8")
                command = (powershell, "-Sta", "-NonInteractive", "-NoLogo", "-NoProfile",
                           "-ExecutionPolicy", "Bypass", "-File", str(script_path))
                return _run_pdf_command(action=action, command=command, state_path=state_path, timeout_seconds=remaining)

            if docx_path is None:
                result = dispatch()
            else:
                pdf_path.parent.mkdir(parents=True, exist_ok=True)
                with staged_pdf_export(docx_path, pdf_path) as (source, target):
                    result = dispatch(source, target)
                    if not result.ok:
                        raise _WordExportAborted(result)
                artifact_published = True
                result = replace(result, elapsed_ms=int((time.perf_counter() - started) * 1000))
    except _WordExportAborted as exc:
        result = exc.result
    except WordPdfBusy as exc:
        result = _build_pdf_failure_result(action=action, raw_message=str(exc), failure_code="export_busy", failure_phase="acquire_export_lock")
    except WordPdfRecoveryRequired as exc:
        result = _build_pdf_failure_result(action=action, raw_message=str(exc), failure_code="cleanup_unconfirmed", failure_phase="recover_previous_export")
    except (OSError, ValueError):
        if artifact_published and result is not None and result.ok:
            # Promotion already completed and Word exited safely. A later lock
            # release error cannot truthfully turn this into an artifact failure
            # or claim the previous PDF was retained. The next operation still
            # has to acquire the OS lock before doing anything.
            return replace(result, details=result.details + "\nThe PDF was verified and saved; releasing the export lock reported an error.")
        failure = _build_pdf_failure_result(
            action=action, raw_message="Could not safely prepare or verify a fresh PDF. The original DOCX and any previous PDF were preserved.",
            failure_code="verification_failed", failure_phase="verify_artifacts",
            elapsed_ms=int((time.perf_counter() - started) * 1000),
        )
        result = replace(failure, helper_pid=result.helper_pid, cleanup_attempted=result.cleanup_attempted,
                         cleanup_succeeded=result.cleanup_succeeded, cleanup_details=result.cleanup_details) if result else failure
    if not result.ok:
        clear_word_pdf_export_readiness_cache()
    return result


def export_docx_to_pdf_in_word(docx_path: Path, pdf_path: Path, *, timeout_seconds: float = 45.0) -> WordAutomationResult:
    return _run_pdf_operation(action="export_pdf", docx_path=docx_path.expanduser().absolute(),
                              pdf_path=pdf_path.expanduser().absolute(), timeout_seconds=float(timeout_seconds))


def probe_word_pdf_export_support(*, timeout_seconds: float = 12.0) -> WordAutomationResult:
    return _run_pdf_operation(action="pdf_preflight", timeout_seconds=float(timeout_seconds))


def _write_word_export_canary_docx(docx_path: Path) -> None:
    document = Document()
    document.add_heading("LegalPDF Translate", level=1)
    document.add_paragraph("Word PDF export canary")
    document.add_paragraph("This temporary document verifies Gmail finalization PDF readiness.")
    document.save(docx_path)


def run_word_pdf_export_canary(*, timeout_seconds: float = 45.0, temp_root: Path | None = None) -> WordAutomationResult:
    action = "pdf_export_canary"
    root = temp_root.expanduser().resolve() if isinstance(temp_root, Path) else None
    working_dir = Path(mkdtemp(prefix="legalpdf_word_export_canary_", dir=str(root) if root else None)).expanduser().resolve()
    try:
        docx_path = working_dir / "honorarios_canary.docx"
        pdf_path = working_dir / "honorarios_canary.pdf"
        _write_word_export_canary_docx(docx_path)
        export_result = export_docx_to_pdf_in_word(docx_path, pdf_path, timeout_seconds=timeout_seconds)
        if not export_result.ok:
            return WordAutomationResult(
                ok=False,
                action=action,
                message=export_result.message,
                stdout=export_result.stdout,
                stderr=export_result.stderr,
                command=export_result.command,
                failure_code=export_result.failure_code,
                details=export_result.details,
                elapsed_ms=export_result.elapsed_ms,
                failure_phase=export_result.failure_phase,
                helper_pid=export_result.helper_pid,
                cleanup_attempted=export_result.cleanup_attempted,
                cleanup_succeeded=export_result.cleanup_succeeded,
                cleanup_details=export_result.cleanup_details,
            )
        from .word_pdf_artifacts import validate_pdf
        try:
            validate_pdf(pdf_path, expected_text="Word PDF export canary")
        except (OSError, ValueError):
            return _build_pdf_failure_result(
                action=action,
                raw_message="The export canary did not produce a readable PDF containing its expected text.",
                command=export_result.command,
                elapsed_ms=export_result.elapsed_ms,
                failure_code="verification_failed",
                failure_phase="verify_pdf_content",
                helper_pid=export_result.helper_pid,
            )
        return WordAutomationResult(
            ok=True,
            action=action,
            message="Word PDF export canary passed.",
            stdout=export_result.stdout,
            stderr=export_result.stderr,
            command=export_result.command,
            details="\n".join(filter(None, (
                export_result.details,
                "PDF pages and expected canary text verified.",
            ))),
            elapsed_ms=export_result.elapsed_ms,
            helper_pid=export_result.helper_pid,
            cleanup_attempted=export_result.cleanup_attempted,
            cleanup_succeeded=export_result.cleanup_succeeded,
            cleanup_details=export_result.cleanup_details,
        )
    finally:
        _cleanup_working_tree_best_effort(working_dir)


def assess_word_pdf_export_readiness(
    *,
    cache_scope: object | None = None,
    launch_timeout_seconds: float = 12.0,
    export_timeout_seconds: float = 45.0,
    force_refresh: bool = False,
    cache_ttl_seconds: float = _WORD_READINESS_CACHE_TTL_SECONDS,
    temp_root: Path | None = None,
) -> dict[str, object]:
    scope = str(cache_scope or "global").strip() or "global"
    now = time.time()
    cached_entry = _WORD_READINESS_CACHE.get(scope)
    if not force_refresh and cached_entry and (now - cached_entry[0]) < float(cache_ttl_seconds):
        payload = dict(cached_entry[1])
        payload["used_cache"] = True
        return payload

    launch_preflight = probe_word_pdf_export_support(timeout_seconds=float(launch_timeout_seconds))
    if launch_preflight.ok:
        export_canary = run_word_pdf_export_canary(
            timeout_seconds=float(export_timeout_seconds),
            temp_root=temp_root,
        )
    else:
        export_canary = WordAutomationResult(
            ok=False,
            action="pdf_export_canary",
            message="Word export canary was skipped because launch preflight failed.",
            failure_code=launch_preflight.failure_code or "launch_preflight_failed",
            details=launch_preflight.details,
            elapsed_ms=launch_preflight.elapsed_ms,
            failure_phase=launch_preflight.failure_phase,
            helper_pid=launch_preflight.helper_pid,
            cleanup_attempted=launch_preflight.cleanup_attempted,
            cleanup_succeeded=launch_preflight.cleanup_succeeded,
            cleanup_details=launch_preflight.cleanup_details,
        )

    finalization_ready = bool(launch_preflight.ok and export_canary.ok)
    effective = export_canary if launch_preflight.ok else launch_preflight
    payload = {
        "ok": finalization_ready,
        "finalization_ready": finalization_ready,
        "failure_code": str(effective.failure_code or "").strip(),
        "message": str(effective.message or "").strip(),
        "details": str(effective.details or "").strip(),
        "elapsed_ms": int(effective.elapsed_ms),
        "failure_phase": str(effective.failure_phase or "").strip(),
        "launch_preflight": serialize_word_automation_result(launch_preflight),
        "export_canary": serialize_word_automation_result(export_canary),
        "preflight": serialize_word_automation_result(launch_preflight),
        "last_checked_at": datetime.now(UTC).replace(microsecond=0).isoformat(),
        "cache_ttl_seconds": int(max(1.0, float(cache_ttl_seconds))),
        "used_cache": False,
    }
    _WORD_READINESS_CACHE[scope] = (now, dict(payload))
    return payload


def clear_word_pdf_export_readiness_cache(*, scope_prefix: object | None = None) -> int:
    prefix = str(scope_prefix or "").strip()
    if prefix == "":
        removed = len(_WORD_READINESS_CACHE)
        _WORD_READINESS_CACHE.clear()
        return removed
    removed = 0
    for scope in list(_WORD_READINESS_CACHE):
        if not scope.startswith(prefix):
            continue
        removed += 1
        _WORD_READINESS_CACHE.pop(scope, None)
    return removed
