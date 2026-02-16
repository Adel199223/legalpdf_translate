import re
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

from docx import Document

import legalpdf_translate.docx_writer as docx_writer
from legalpdf_translate.docx_writer import (
    _remove_compatibility_mode,
    assemble_docx,
    sanitize_bidi_controls,
)
from legalpdf_translate.output_normalize import normalize_output_text
from legalpdf_translate.types import TargetLang


def _write_page(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def test_assemble_docx_has_no_empty_paragraphs_and_page_break(tmp_path: Path) -> None:
    pages_dir = tmp_path / "pages"
    _write_page(pages_dir / "page_0001.txt", "Line A\n\nLine B")
    _write_page(pages_dir / "page_0002.txt", "Line C")
    out = tmp_path / "out.docx"

    assemble_docx(pages_dir, out, lang=TargetLang.EN, page_breaks=True)

    doc = Document(out)
    texts = [p.text for p in doc.paragraphs]
    assert "" not in texts
    assert texts == ["Line A", "Line B", "Line C"]
    assert 'w:type="page"' in doc.element.xml


def test_assemble_docx_arabic_sets_rtl_bidi_flags(tmp_path: Path) -> None:
    pages_dir = tmp_path / "pages"
    _write_page(pages_dir / "page_0001.txt", "النص \u2066[[ABC-123]]\u2069")
    out = tmp_path / "ar.docx"

    assemble_docx(pages_dir, out, lang=TargetLang.AR, page_breaks=False)

    doc = Document(out)
    paragraph_xml = doc.paragraphs[0]._p.xml
    assert "<w:bidi" in paragraph_xml
    assert "<w:rtl" in paragraph_xml


def test_save_document_atomic_fsyncs_temp_file(tmp_path: Path, monkeypatch) -> None:
    pages_dir = tmp_path / "pages"
    _write_page(pages_dir / "page_0001.txt", "Line A")
    out = tmp_path / "out.docx"

    calls = {"count": 0}
    original_fsync = docx_writer.os.fsync

    def _counting_fsync(fd: int) -> None:
        calls["count"] += 1
        original_fsync(fd)

    monkeypatch.setattr(docx_writer.os, "fsync", _counting_fsync)
    assemble_docx(pages_dir, out, lang=TargetLang.EN, page_breaks=False)

    assert out.exists()
    assert calls["count"] >= 1


def test_sanitize_bidi_controls_removes_expected_codepoints_only() -> None:
    text = "A \u2066[[ABC-123]]\u2069 \u200eB\u200f \u202aC\u202e \u061cD\ufeff."

    sanitized = sanitize_bidi_controls(text)

    assert sanitized == "A [[ABC-123]] B C D."


def test_assemble_docx_strips_bidi_controls_from_document_xml(tmp_path: Path) -> None:
    pages_dir = tmp_path / "pages"
    _write_page(
        pages_dir / "page_0001.txt",
        "x\u2066[[TOK]]\u2069 y \u200ez\u200f \u202aA\u202e \u061cB\ufeff",
    )
    out = tmp_path / "clean.docx"

    assemble_docx(pages_dir, out, lang=TargetLang.AR, page_breaks=False)

    with ZipFile(out) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    for marker in ("\u061c", "\u200e", "\u200f", "\u202a", "\u202b", "\u202c", "\u202d", "\u202e", "\u2066", "\u2067", "\u2068", "\u2069", "\ufeff"):
        assert marker not in document_xml
    assert "<w:bidi" in document_xml
    assert "<w:rtl" in document_xml


def test_assemble_docx_can_preserve_bidi_controls_when_requested(tmp_path: Path) -> None:
    pages_dir = tmp_path / "pages"
    _write_page(pages_dir / "page_0001.txt", "token \u2066[[ABC-123]]\u2069")
    out = tmp_path / "preserve.docx"

    assemble_docx(
        pages_dir,
        out,
        lang=TargetLang.AR,
        page_breaks=False,
        strip_bidi_controls=False,
    )

    with ZipFile(out) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "\u2066" in document_xml
    assert "\u2069" in document_xml


def test_docx_no_standalone_list_marker_paragraph(tmp_path: Path) -> None:
    raw = "1.\nLe texte\n2.\nAutre texte"
    normalized = normalize_output_text(raw, lang=TargetLang.FR)
    pages_dir = tmp_path / "pages"
    _write_page(pages_dir / "page_0001.txt", normalized)
    out = tmp_path / "merged.docx"
    assemble_docx(pages_dir, out, lang=TargetLang.FR, page_breaks=False)
    doc = Document(out)
    texts = [p.text for p in doc.paragraphs]
    assert "1." not in texts
    assert any(t.startswith("1. Le texte") for t in texts)
    assert any(t.startswith("2. Autre texte") for t in texts)


def _assert_ignorable_prefixes_declared(xml_text: str) -> None:
    """Every prefix listed in an Ignorable attribute must have an xmlns declaration."""
    ignorable_match = re.search(r'\bIgnorable\s*=\s*"([^"]*)"', xml_text)
    if not ignorable_match:
        return
    prefixes = ignorable_match.group(1).split()
    for prefix in prefixes:
        assert re.search(rf'\bxmlns:{re.escape(prefix)}\s*=', xml_text), (
            f"Prefix '{prefix}' is listed in Ignorable but has no xmlns declaration"
        )


def _assert_all_xml_parts_wellformed(docx_path: Path) -> None:
    """Parse every XML part in the DOCX ZIP to assert well-formedness."""
    with ZipFile(docx_path) as archive:
        for name in archive.namelist():
            if name.endswith(".xml") or name.endswith(".rels"):
                data = archive.read(name)
                ET.fromstring(data)  # raises on malformed XML


def test_generated_docx_has_no_compatibility_mode(tmp_path: Path) -> None:
    pages_dir = tmp_path / "pages"
    _write_page(pages_dir / "page_0001.txt", "Hello world")
    out = tmp_path / "compat.docx"
    assemble_docx(pages_dir, out, lang=TargetLang.EN, page_breaks=False)
    with ZipFile(out) as archive:
        if "word/settings.xml" in archive.namelist():
            settings_xml = archive.read("word/settings.xml").decode("utf-8")
            assert "compatibilityMode" not in settings_xml
            _assert_ignorable_prefixes_declared(settings_xml)
    _assert_all_xml_parts_wellformed(out)


# -- Synthetic settings.xml tests for _remove_compatibility_mode --

_SETTINGS_WITH_COMPAT = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:settings xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:o="urn:schemas-microsoft-com:office:office"
  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
  xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
  xmlns:v="urn:schemas-microsoft-com:vml"
  xmlns:w10="urn:schemas-microsoft-com:office:word"
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
  xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
  xmlns:sl="http://schemas.openxmlformats.org/schemaLibrary/2006/main"
  mc:Ignorable="w14 w15">
  <w:zoom w:percent="100"/>
  <w:compat>
    <w:compatSetting w:name="compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="15"/>
    <w:compatSetting w:name="overrideTableStyleFontSizeAndJustification" w:uri="http://schemas.microsoft.com/office/word" w:val="1"/>
  </w:compat>
</w:settings>"""

_SETTINGS_COMPAT_ONLY_CHILD = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:settings xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
  mc:Ignorable="w14">
  <w:zoom w:percent="100"/>
  <w:compat>
    <w:compatSetting w:name="compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="15"/>
  </w:compat>
</w:settings>"""

_SETTINGS_NO_COMPAT = """\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:settings xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
  mc:Ignorable="w14">
  <w:zoom w:percent="100"/>
</w:settings>"""


def _make_docx_with_settings(tmp_path: Path, settings_xml: str) -> Path:
    """Build a minimal DOCX from a real one, replacing word/settings.xml."""
    # First create a real DOCX via python-docx so the ZIP has all required parts.
    pages_dir = tmp_path / "pages"
    _write_page(pages_dir / "page_0001.txt", "Test content")
    base = tmp_path / "base.docx"
    assemble_docx(pages_dir, base, lang=TargetLang.EN, page_breaks=False)

    # Now replace word/settings.xml with the synthetic content.
    patched = tmp_path / "patched.docx"
    with ZipFile(base, "r") as zin, ZipFile(patched, "w") as zout:
        for item in zin.infolist():
            if item.filename == "word/settings.xml":
                zout.writestr(item, settings_xml.encode("utf-8"))
            else:
                zout.writestr(item, zin.read(item.filename))
    return patched


def test_compat_removal_preserves_all_namespaces(tmp_path: Path) -> None:
    docx = _make_docx_with_settings(tmp_path, _SETTINGS_WITH_COMPAT)
    _remove_compatibility_mode(docx)

    with ZipFile(docx) as archive:
        result = archive.read("word/settings.xml").decode("utf-8")

    assert "compatibilityMode" not in result
    # The other compatSetting sibling must survive.
    assert "overrideTableStyleFontSizeAndJustification" in result
    # All namespace declarations must survive.
    assert 'xmlns:w14=' in result
    assert 'xmlns:w15=' in result
    assert 'mc:Ignorable=' in result
    _assert_ignorable_prefixes_declared(result)
    ET.fromstring(result.encode("utf-8"))  # Must be well-formed.


def test_compat_removal_noop_when_absent(tmp_path: Path) -> None:
    docx = _make_docx_with_settings(tmp_path, _SETTINGS_NO_COMPAT)

    with ZipFile(docx) as archive:
        before = archive.read("word/settings.xml")

    _remove_compatibility_mode(docx)

    with ZipFile(docx) as archive:
        after = archive.read("word/settings.xml")

    assert before == after  # Byte-for-byte identical; no rewrite occurred.


def test_compat_removal_removes_empty_compat_element(tmp_path: Path) -> None:
    docx = _make_docx_with_settings(tmp_path, _SETTINGS_COMPAT_ONLY_CHILD)
    _remove_compatibility_mode(docx)

    with ZipFile(docx) as archive:
        result = archive.read("word/settings.xml").decode("utf-8")

    assert "compatibilityMode" not in result
    assert "<w:compat" not in result  # Empty wrapper also removed.
    assert 'xmlns:w14=' in result     # Namespaces preserved.
    _assert_ignorable_prefixes_declared(result)
    ET.fromstring(result.encode("utf-8"))


def test_compat_removal_keeps_compat_with_other_children(tmp_path: Path) -> None:
    docx = _make_docx_with_settings(tmp_path, _SETTINGS_WITH_COMPAT)
    _remove_compatibility_mode(docx)

    with ZipFile(docx) as archive:
        result = archive.read("word/settings.xml").decode("utf-8")

    assert "compatibilityMode" not in result
    assert "<w:compat" in result  # Parent kept because it has another child.
    assert "overrideTableStyleFontSizeAndJustification" in result
