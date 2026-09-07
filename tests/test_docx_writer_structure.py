from __future__ import annotations

import hashlib
import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pytest

from legalpdf_translate.document_structure import PageStructure, StructureBlock, text_sha256
from legalpdf_translate.docx_writer import assemble_docx, sanitize_bidi_controls
from legalpdf_translate.joblog_flow import count_words_from_docx
from legalpdf_translate.types import TargetLang


def _save_page(folder: Path, number: int, blocks: list[dict], **flags) -> tuple[Path, dict]:
    folder.mkdir(exist_ok=True)
    items = [StructureBlock(id=f"p{number:04d}_b{index:04d}", **block) for index, block in enumerate(blocks, 1)]
    structure = PageStructure(page_number=number, source_sha256=text_sha256("synthetic source"),
                              blocks=items, **flags)
    structure.translation_sha256 = text_sha256(structure.text)
    path = folder / f"page_{number:04d}.txt"
    path.write_text(structure.text, encoding="utf-8")
    payload = structure.to_dict()
    path.with_suffix(".structure.json").write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    return path, payload


def _mapping(output: Path) -> dict:
    return json.loads(output.with_suffix(".source_map.json").read_text(encoding="utf-8"))


@pytest.mark.parametrize(("lang", "font_name", "half_points"), [
    (TargetLang.EN, "Times New Roman", "21"),
    (TargetLang.FR, "Times New Roman", "21"),
    (TargetLang.AR, "Arial", "22"),
])
def test_compact_profile_sets_explicit_fonts_margins_and_dynamic_footer(tmp_path, lang, font_name, half_points):
    pages = tmp_path / "pages"
    _save_page(pages, 1, [{"text": "النص Text"}])
    output = assemble_docx(pages, tmp_path / "out.docx", lang=lang, page_breaks=False)
    doc = Document(output)
    section = doc.sections[0]
    assert section.page_width.pt == pytest.approx(595.276, abs=0.05)
    assert section.page_height.pt == pytest.approx(841.89, abs=0.05)
    assert section.left_margin == section.right_margin
    assert section.left_margin.cm == pytest.approx(1.7, abs=0.002)
    assert section.top_margin == section.bottom_margin
    assert section.top_margin.cm == pytest.approx(1.5, abs=0.002)
    assert doc.styles["Normal"].paragraph_format.line_spacing == 1.0
    assert doc.styles["LegalPDF Body"].paragraph_format.space_after.pt == 3
    for run in doc.paragraphs[0].runs:
        fonts = run._r.rPr.find(qn("w:rFonts"))
        assert all(fonts.get(qn(f"w:{slot}")) == font_name for slot in ("ascii", "hAnsi", "cs", "eastAsia"))
        assert run._r.rPr.find(qn("w:sz")).get(qn("w:val")) == half_points
        assert run._r.rPr.find(qn("w:szCs")).get(qn("w:val")) == half_points
    with ZipFile(output) as archive:
        footer = archive.read("word/footer1.xml").decode()
        assert " PAGE " in footer
        for kind in ("begin", "separate", "end"):
            assert f'w:fldCharType="{kind}"' in footer
        assert 'w:updateFields' not in archive.read("word/settings.xml").decode()


def test_legacy_txt_rebuild_is_compact_and_source_map_is_not_a_rendered_page_count(tmp_path):
    pages = tmp_path / "pages"
    pages.mkdir()
    (pages / "page_0001.txt").write_text("Alpha\nBeta", encoding="utf-8")
    (pages / "page_0002.txt").write_text("Gamma", encoding="utf-8")
    stats = {}
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False, stats=stats)
    doc = Document(output)
    assert [p.text for p in doc.paragraphs] == ["Alpha", "Beta", "Gamma"]
    assert 'w:type="page"' not in doc._element.xml
    mapping = _mapping(output)
    assert mapping["source_page_count"] == 2
    assert mapping["rendered_page_count"] is None
    assert mapping["docx_sha256"] == hashlib.sha256(output.read_bytes()).hexdigest()
    assert all(page["structure_status"] == "legacy_txt" for page in mapping["pages"])
    assert stats["page_count"] == stats["source_page_count"] == 2
    assert stats["structure_fallback_count"] == 2


@pytest.mark.parametrize("corruption", ["version", "hash", "partial", "page", "malformed", "sparse_table"])
def test_untrusted_or_stale_structure_falls_back_to_complete_txt(tmp_path, corruption):
    pages = tmp_path / "pages"
    path, payload = _save_page(pages, 1, [{"text": "First complete paragraph."}, {"text": "Second obligation."}])
    if corruption == "version":
        payload["version"] = 200
    elif corruption == "hash":
        payload["translation_sha256"] = "0" * 64
    elif corruption == "partial":
        payload["blocks"].pop()  # Leave TXT hash untouched: independently hash block contents.
    elif corruption == "page":
        payload["page_number"] = 2
    elif corruption == "sparse_table":
        for index, block in enumerate(payload["blocks"]):
            block.update(role="table_cell", table_id="t", row=199, col=20 + index)
    serialized = "{bad" if corruption == "malformed" else json.dumps(payload)
    path.with_suffix(".structure.json").write_text(serialized, encoding="utf-8")
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False)
    assert [p.text for p in Document(output).paragraphs] == ["First complete paragraph.", "Second obligation."]
    assert _mapping(output)["pages"][0]["structure_status"] == "invalid_sidecar_txt_fallback"


def test_roles_keep_alignment_emphasis_literal_numbering_and_substantive_footer(tmp_path):
    pages = tmp_path / "pages"
    _save_page(pages, 1, [
        {"text": "قرار المحكمة", "role": "heading", "alignment": "center", "bold": True},
        {"text": "Rua Luís de Camões", "role": "address", "alignment": "left"},
        {"text": "1. الحضور واجب", "role": "list_item"},
        {"text": "التوقيع", "role": "signature", "italic": True},
        {"text": "Referência: ABC-1", "role": "reference"},
        {"text": "Page 1 of 2", "role": "footer"},
        {"text": "Contact: example.invalid — retention 10 days", "role": "footer"},
        {"text": "Page 1 of 2", "role": "paragraph"},
    ])
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    doc = Document(output)
    assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert doc.paragraphs[0].style.paragraph_format.keep_with_next is True
    assert doc.paragraphs[0].runs[0].bold is True
    assert doc.paragraphs[0].runs[0].font.cs_bold is True
    assert doc.paragraphs[1]._p.pPr.jc.get(qn("w:val")) == "end"
    assert doc.paragraphs[1].style.paragraph_format.space_after.pt == 0
    assert sanitize_bidi_controls(doc.paragraphs[2].text) == "1. الحضور واجب"
    assert doc.paragraphs[3].runs[0].italic is True
    assert len(doc.paragraphs) == 7
    assert "retention 10 days" in doc.paragraphs[-2].text
    assert doc.paragraphs[-1].text == "Page 1 of 2"
    blocks = _mapping(output)["pages"][0]["blocks"]
    assert blocks[5]["location"]["kind"] == "generated_footer_page_field"
    assert blocks[6]["location"]["kind"] == "body_paragraph"


@pytest.mark.parametrize(("alignment", "expected"), [(None, "start"), ("right", "start"),
                                                   ("left", "end"), ("center", "center"),
                                                   ("justify", "both")])
def test_arabic_visual_alignment_uses_unambiguous_logical_values_in_body_and_table(tmp_path, alignment, expected):
    pages = tmp_path / "pages"
    _save_page(pages, 1, [
        {"text": "نص عربي", "alignment": alignment},
        {"text": "خلية عربية", "role": "table_cell", "table_id": "t", "row": 0, "col": 0,
         "alignment": alignment},
    ])
    output = assemble_docx(pages, tmp_path / "aligned.docx", lang=TargetLang.AR, page_breaks=False)
    document = Document(output)
    for index, paragraph in enumerate([document.paragraphs[0], document.tables[0].cell(0, 0).paragraphs[0]]):
        p_pr = paragraph._p.pPr
        # python-docx can reopen/read/rewrite all text, but its legacy alignment
        # enum does not expose Office 2010 start/end. Inspect the native XML.
        assert paragraph.text
        assert p_pr.jc.get(qn("w:val")) == ("start" if index == 0 and alignment == "left" else expected)
        assert p_pr.find(qn("w:rtl")) is None
        assert p_pr.find(qn("w:bidi")).get(qn("w:val")) == "1"
        assert list(p_pr).index(p_pr.find(qn("w:bidi"))) < list(p_pr).index(p_pr.jc)
    assert document.tables[0]._tbl.tblPr.find(qn("w:bidiVisual")) is not None


@pytest.mark.parametrize("role", ["paragraph", "list_item", "heading", "header", "footer", "address", "signature", "reference"])
@pytest.mark.parametrize("lang", [TargetLang.AR, TargetLang.EN, TargetLang.FR])
def test_arabic_prose_and_label_roles_adapt_source_left_alignment(tmp_path, role, lang):
    pages = tmp_path / "pages"
    text = "1. نص عربي [[João Exemplo]]"
    source_path, _ = _save_page(pages, 1, [{"text": text, "role": role, "alignment": "left"}])
    sidecar = source_path.with_suffix(".structure.json")
    original_source = sidecar.read_bytes()
    output = assemble_docx(pages, tmp_path / "flow.docx", lang=lang, page_breaks=False)
    paragraph = Document(output).paragraphs[0]
    expected = "start" if role in {"paragraph", "list_item", "heading", "header", "signature"} else "end"
    assert paragraph._p.pPr.jc.get(qn("w:val")) == (expected if lang == TargetLang.AR else "left")
    assert sanitize_bidi_controls(paragraph.text) == (text.replace("[[", "").replace("]]", "") if lang == TargetLang.AR else text)
    assert sidecar.read_bytes() == original_source


def test_editable_table_is_visible_counted_and_mapped_without_text_loss(tmp_path):
    pages = tmp_path / "pages"
    _save_page(pages, 1, [
        {"text": "Name", "role": "table_cell", "table_id": "t1", "row": 0, "col": 0, "bold": True},
        {"text": "Date", "role": "table_cell", "table_id": "t1", "row": 0, "col": 1, "bold": True},
        {"text": "Synthetic Person", "role": "table_cell", "table_id": "t1", "row": 1, "col": 0},
        {"text": "01/02/2026", "role": "table_cell", "table_id": "t1", "row": 1, "col": 1},
    ])
    stats = {}
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.FR, page_breaks=False, stats=stats)
    doc = Document(output)
    assert len(doc.tables) == 1
    assert [[cell.text for cell in row.cells] for row in doc.tables[0].rows] == [
        ["Name", "Date"], ["Synthetic Person", "01/02/2026"]]
    assert count_words_from_docx(output) == 5
    assert stats["paragraph_count"] == 4
    assert "w:tblBorders" in doc.tables[0]._tbl.xml
    assert all(block["location"]["kind"] == "table_cell" for block in _mapping(output)["pages"][0]["blocks"])


@pytest.mark.parametrize("geometry", ["consistent", "missing", "overlap", "variable_left", "tiny"])
def test_table_widths_follow_reliable_source_columns_and_fall_back_on_ambiguity(tmp_path, geometry):
    pages = tmp_path / "pages"
    boxes = [(20, 20, 120, 40), (120, 20, 520, 40), (20, 40, 120, 70), (120, 40, 520, 70)]
    if geometry == "missing":
        boxes[2] = None
    elif geometry == "overlap":
        boxes[0] = (20, 20, 400, 40)
    elif geometry == "variable_left":
        boxes[2] = (45, 40, 120, 70)
    elif geometry == "tiny":
        boxes = [(20, 20, 22, 40), (22, 20, 520, 40), (20, 40, 22, 70), (22, 40, 520, 70)]
    _save_page(pages, 1, [{"text": text, "role": "table_cell", "table_id": "t1", "row": index // 2,
                           "col": index % 2, "bbox": boxes[index]}
                          for index, text in enumerate(["Label", "Long retained prose", "Label 2", "More retained prose"])])
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.FR, page_breaks=False)
    table = Document(output).tables[0]
    widths = [column.width.pt for column in table.columns]
    assert widths[1] / widths[0] == pytest.approx(4.0 if geometry == "consistent" else 1.0, abs=0.005)
    for row in table.rows:
        assert [cell.width.pt for cell in row.cells] == pytest.approx(widths)
    expected = "source_geometry" if geometry == "consistent" else "equal_width_fallback"
    assert all(block["table_geometry_status"] == expected for block in _mapping(output)["pages"][0]["blocks"])


@pytest.mark.parametrize("confirm", [True, False])
def test_only_explicit_confirmed_continuations_join_and_keep_source_mapping(tmp_path, confirm):
    pages = tmp_path / "pages"
    _save_page(pages, 1, [{"text": "The person must"}], continuation_to_next=confirm)
    _save_page(pages, 2, [{"text": "appear at the hearing.", "continuation_of": "p0001_b0001"}],
               continuation_from_previous=True)
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False)
    doc = Document(output)
    assert [p.text for p in doc.paragraphs] == (["The person must appear at the hearing."] if confirm else
                                              ["The person must", "appear at the hearing."])
    mapping = _mapping(output)
    first, second = [page["blocks"][0] for page in mapping["pages"]]
    assert (first["location"] == second["location"]) is confirm
    assert ("joined_to_block_id" in second) is confirm


def _save_furniture_bridge_pages(folder: Path) -> tuple[dict, dict]:
    from legalpdf_translate.formatting_support import link_source_continuations
    folder.mkdir()
    sources = {}
    for number in (1, 2):
        source = PageStructure(page_number=number, source_sha256="a" * 64, source_file_sha256="b" * 64, blocks=[
            StructureBlock(f"p{number:04d}_b0001", "Tribunal Judicial", role="header", bbox=(200, 90, 400, 112)),
            StructureBlock(f"p{number:04d}_b0002", "A parte deve comparecer na audiência de julgamento marcada para" if number == 1 else "o dia indicado na notificação.",
                           bbox=(30, 760, 540, 785) if number == 1 else (30, 150, 540, 175), alignment="justify"),
            StructureBlock(f"p{number:04d}_b0003", "Largo da Justiça, 1", role="footer", bbox=(200, 802, 400, 810)),
            StructureBlock(f"p{number:04d}_b0004", "Telef: 210000000; E-mail: tribunal@example.invalid", role="footer", bbox=(190, 812, 410, 820)),
        ])
        source.source_sha256 = source.source_text_sha256 = text_sha256(source.text)
        sources[number] = source
    sources = link_source_continuations(sources)
    targets = {}
    for number, source in sources.items():
        (folder / f"page_{number:04d}.source_structure.json").write_text(json.dumps(source.to_dict()), encoding="utf-8")
        target = source.to_dict()
        target["blocks"][0]["text"] = "Tribunal judiciaire" if number == 1 else "Tribunal judiciaire répété"
        target["blocks"][1]["text"] = "La partie doit comparaître à l’audience fixée pour" if number == 1 else "le jour indiqué dans la notification."
        text = "\n".join(block["text"] for block in target["blocks"])
        target["translation_sha256"] = text_sha256(text)
        path = folder / f"page_{number:04d}.txt"
        path.write_text(text, encoding="utf-8")
        path.with_suffix(".structure.json").write_text(json.dumps(target), encoding="utf-8")
        targets[number] = target
    return sources, targets


@pytest.fixture
def legacy_sentence_furniture(monkeypatch):
    # Isolate the old source-proof sentence bridge when section consolidation
    # is unavailable. Production-on joins are covered in section writer tests.
    monkeypatch.setattr("legalpdf_translate.docx_writer._section_furniture_plan",
                        lambda *a, **k: {"sections": [], "pages": []})


@pytest.mark.usefixtures("legacy_sentence_furniture")
def test_confirmed_sentence_joins_earlier_body_while_all_furniture_and_source_mappings_remain(tmp_path):
    pages = tmp_path / "pages"
    sources, targets = _save_furniture_bridge_pages(pages)
    # Independently reviewed source links do not need rewritten/fabricated derived
    # flags or bridge metadata: complete local evidence is recomputed by writer.
    targets[2]["metadata"].pop("continuation_bridge")
    (pages / "page_0002.structure.json").write_text(json.dumps(targets[2]), encoding="utf-8")
    before_sources = [(pages / f"page_{number:04d}.source_structure.json").read_bytes() for number in (1, 2)]
    output = assemble_docx(pages, tmp_path / "joined.docx", lang=TargetLang.FR, page_breaks=False)
    paragraphs = [paragraph.text for paragraph in Document(output).paragraphs]
    assert paragraphs[1] == "La partie doit comparaître à l’audience fixée pour le jour indiqué dans la notification."
    assert len(paragraphs) == 7  # Eight retained source blocks, one joined paragraph.
    assert paragraphs.count("Largo da Justiça, 1") == 2
    assert paragraphs.count("Telef: 210000000; E-mail: tribunal@example.invalid") == 2
    assert "Tribunal judiciaire" in paragraphs and "Tribunal judiciaire répété" in paragraphs
    first, second = _mapping(output)["pages"]
    assert len(first["blocks"]) == len(second["blocks"]) == 4
    assert first["blocks"][1]["location"] == second["blocks"][1]["location"]
    assert second["blocks"][1]["joined_to_block_id"] == "p0001_b0002"
    assert second["blocks"][1]["continuation_evidence"] == "revalidated_identical_source_furniture"
    assert all(block["location"]["kind"] == "body_paragraph" for page in (first, second) for block in page["blocks"])
    assert before_sources == [(pages / f"page_{number:04d}.source_structure.json").read_bytes() for number in (1, 2)]


@pytest.mark.parametrize("guard", ["missing_source", "corrupt_source_hash", "wrong_source_file", "altered_geometry",
                                  "forged_bridge", "explicit_breaks", "document_start", "block_start"])
@pytest.mark.usefixtures("legacy_sentence_furniture")
def test_furniture_bridge_requires_full_bound_source_proof_and_layout_permission(tmp_path, guard):
    pages = tmp_path / "pages"
    sources, targets = _save_furniture_bridge_pages(pages)
    if guard == "missing_source":
        (pages / "page_0001.source_structure.json").unlink()
    elif guard == "corrupt_source_hash":
        payload = sources[1].to_dict()
        payload["blocks"][0]["text"] = "Different source header"
        (pages / "page_0001.source_structure.json").write_text(json.dumps(payload))
    elif guard == "wrong_source_file":
        targets[2]["source_file_sha256"] = "c" * 64
    elif guard == "altered_geometry":
        targets[2]["blocks"][0]["bbox"] = [200, 350, 400, 375]
    elif guard == "forged_bridge":
        targets[2]["metadata"]["continuation_bridge"]["previous_block_id"] = "p0001_b0001"
    elif guard == "document_start":
        targets[2]["document_start"] = True
    elif guard == "block_start":
        targets[2]["blocks"][0]["document_start"] = True
    (pages / "page_0002.structure.json").write_text(json.dumps(targets[2]))
    output = assemble_docx(pages, tmp_path / "unjoined.docx", lang=TargetLang.FR, page_breaks=guard == "explicit_breaks")
    mapping = _mapping(output)["pages"][1]["blocks"][1]
    assert "joined_to_block_id" not in mapping
    assert len([paragraph for paragraph in Document(output).paragraphs if paragraph.text.strip()]) == 8


@pytest.mark.parametrize("guard", ["explicit_breaks", "document_start", "block_start", "uncertain", "page_uncertain", "missing_page"])
def test_continuations_cannot_cross_layout_or_evidence_boundaries(tmp_path, guard):
    pages = tmp_path / "pages"
    _save_page(pages, 1, [{"text": "The person must"}], continuation_to_next=True)
    number = 3 if guard == "missing_page" else 2
    _save_page(pages, number, [{"text": "appear at the hearing.", "continuation_of": "p0001_b0001",
                               "uncertain": guard == "uncertain", "document_start": guard == "block_start"}],
               continuation_from_previous=True, document_start=guard == "document_start", uncertain=guard == "page_uncertain")
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=guard == "explicit_breaks")
    doc = Document(output)
    assert len([p for p in doc.paragraphs if p.text.strip()]) == 2
    assert "joined_to_block_id" not in _mapping(output)["pages"][1]["blocks"][0]
    if guard in {"explicit_breaks", "document_start", "block_start"}:
        assert 'w:type="page"' in doc._element.xml


def test_source_page_size_and_table_document_boundary_create_valid_sections(tmp_path):
    pages = tmp_path / "pages"
    _save_page(pages, 1, [{"text": "Table fact", "role": "table_cell", "table_id": "t1", "row": 0, "col": 0}],
               width_pt=612, height_pt=792)
    _save_page(pages, 2, [{"text": "Independent document."}], document_start=True)
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False)
    doc = Document(output)
    assert len(doc.sections) == 2
    assert doc.sections[0].page_width.pt == 612
    assert doc.sections[0].page_height.pt == 792
    assert doc.sections[1].page_width.pt == pytest.approx(595.276, abs=0.05)
    assert doc.sections[1].footer.is_linked_to_previous
    assert len(doc.tables) == 1
    assert doc.paragraphs[-1].text == "Independent document."


@pytest.mark.parametrize("name", ["[[Jo]]ã[[o Gon]]ç[[alves]]", "[[José Luís de Camões]]",
                                  "[[Jose\u0301 D’Ávila-São]]"])
def test_arabic_latin_accented_names_are_coherent_uniform_font_runs(tmp_path, name):
    pages = tmp_path / "pages"
    _save_page(pages, 1, [{"text": f"الاسم: {name}، للحضور"}])
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    paragraph = Document(output).paragraphs[0]
    expected_name = name.replace("[[", "").replace("]]", "")
    latin = [run for run in paragraph.runs if run._r.rPr.find(qn("w:rtl")).get(qn("w:val")) == "0"]
    assert len(latin) == 1
    assert expected_name in sanitize_bidi_controls(latin[0].text)
    assert latin[0].font.name == "Arial"
    assert latin[0].font.size.pt == 11
    assert expected_name in sanitize_bidi_controls(paragraph.text)
