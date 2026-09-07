from __future__ import annotations

from copy import deepcopy
import json
from pathlib import Path

from docx import Document
import pytest

from legalpdf_translate.document_structure import PageStructure, StructureBlock, text_sha256
from legalpdf_translate.docx_furniture_reflow import confirmed_list_furniture_reflow
import legalpdf_translate.docx_writer as writer
from legalpdf_translate.types import TargetLang


@pytest.fixture
def legacy_list_postpass(monkeypatch):
    # These tests retain the pre-consolidation fallback contract. Production-on
    # section consolidation and body/alias conservation have their own suite.
    monkeypatch.setattr(writer, "_section_furniture_plan", lambda *a, **k: {"sections": [], "pages": []})


def _page(number, body, **flags):
    blocks = [dict(text="Tribunal Judicial de Cidade", role="header", bbox=(200, 90, 400, 112))]
    blocks.extend(body)
    blocks.extend([
        dict(text="Largo da Justiça, 1", role="footer", bbox=(200, 802, 400, 810)),
        dict(text="Telef: 210000000 - E-mail: tribunal@example.invalid", role="footer", bbox=(190, 812, 410, 820)),
    ])
    source = PageStructure(page_number=number, source_sha256="a" * 64, source_file_sha256="b" * 64,
                           blocks=[StructureBlock(f"p{number:04d}_b{i:04d}", **block)
                                   for i, block in enumerate(blocks, 1)], **flags)
    source.source_sha256 = source.source_text_sha256 = text_sha256(source.text)
    return source.to_dict()


def _pair(kind="number", *, next_number=2):
    starts = {"number": ("1. Primeira condição;", "2. Segunda condição;", "3. Terceira condição."),
              "letter": ("a) Primeira condição;", "b) Segunda condição;", "c) Terceira condição."),
              "dash": ("- Primeira condição;", "- Segunda condição;", "- Terceira condição."),
              "bullet": ("• Primeira condição, e", "• Segunda condição;", "• Terceira condição."),
              "intro": ("Devem verificar-se todas as seguintes condições:", "1. Primeira condição;", "2. Segunda condição.")}
    first, second, third = starts[kind]
    return [_page(1, [dict(text=first, role="paragraph" if kind == "intro" else "list_item",
                           bbox=(70, 740, 530, 780), alignment="left")]),
            _page(next_number, [dict(text=second, role="list_item", bbox=(70, 150, 530, 175), alignment="left"),
                                dict(text=third, role="list_item", bbox=(70, 185, 530, 210), alignment="left"),
                                dict(text="Outro parágrafo independente.", role="paragraph", bbox=(70, 230, 530, 260))])]


def _refresh(source):
    source["source_sha256"] = source["source_text_sha256"] = text_sha256("\n".join(b["text"] for b in source["blocks"]))


def _save(folder, sources, *, lang=TargetLang.EN):
    folder.mkdir(exist_ok=True)
    targets = []
    for source in sources:
        _refresh(source)
        number = source["page_number"]
        (folder / f"page_{number:04d}.source_structure.json").write_text(json.dumps(source), encoding="utf-8")
        target = deepcopy(source)
        for block in target["blocks"]:
            # Different translations of repeated source furniture must survive.
            if block["role"] in {"header", "footer"}:
                block["text"] += f" [target page {number}]"
            elif lang == TargetLang.AR:
                marker = block["text"].split(" ", 1)[0] if block["role"] == "list_item" else ""
                block["text"] = f"{marker} نص عربي مع [[Élodie Martins]] {block['id']}".strip()
                if source["blocks"][1]["role"] == "paragraph" and block["id"] == source["blocks"][1]["id"]:
                    block["text"] += ":"
        text = "\n".join(block["text"] for block in target["blocks"])
        target["translation_sha256"] = text_sha256(text)
        path = folder / f"page_{number:04d}.txt"
        path.write_text(text, encoding="utf-8")
        path.with_suffix(".structure.json").write_text(json.dumps(target), encoding="utf-8")
        targets.append(target)
    return targets


def _read(output):
    doc = Document(output)
    mapping = json.loads(output.with_suffix(".source_map.json").read_text(encoding="utf-8"))
    return doc, mapping


def _mapped_xml(output):
    doc, mapping = _read(output)
    return {b["block_id"]: doc.paragraphs[b["location"]["paragraph_index"]]._element.xml
            for page in mapping["pages"] for b in page["blocks"] if b["location"]["kind"] == "body_paragraph"}


@pytest.mark.parametrize("kind", ["number", "letter", "dash", "bullet", "intro"])
def test_source_proof_distinguishes_list_continuation_from_sentence_joining(kind):
    sources = _pair(kind)
    before = deepcopy(sources)
    proof = confirmed_list_furniture_reflow(*sources)
    assert proof["kind"] == ("introduced_list" if kind == "intro" else "continued_list")
    assert proof["current_list_ids"] == ["p0002_b0002", "p0002_b0003"]
    assert len(proof["furniture_pairs"]) == 3
    assert sources == before


@pytest.mark.parametrize("lang", [TargetLang.AR, TargetLang.EN, TargetLang.FR])
@pytest.mark.usefixtures("legacy_list_postpass")
@pytest.mark.parametrize("kind", ["intro", "bullet"])
def test_compact_list_reflow_preserves_every_paragraph_run_and_source_locator(tmp_path, monkeypatch, lang, kind):
    pages = tmp_path / "pages"
    targets = _save(pages, _pair(kind), lang=lang)
    before_files = {p.name: p.read_bytes() for p in pages.iterdir()}
    with monkeypatch.context() as patch:
        patch.setattr(writer, "_reflow_compact_lists", lambda *args, **kwargs: None)
        control = writer.assemble_docx(pages, tmp_path / "control.docx", lang=lang, page_breaks=False)
    output = writer.assemble_docx(pages, tmp_path / "reflow.docx", lang=lang, page_breaks=False)
    doc, mapping = _read(output)
    assert _mapped_xml(output) == _mapped_xml(control)  # Includes all runs, fonts, markers and paragraph properties.
    assert {p.name: p.read_bytes() for p in pages.iterdir()} == before_files
    by_id = {b["block_id"]: b for p in mapping["pages"] for b in p["blocks"]}
    ordered = sorted(by_id, key=lambda identity: by_id[identity]["location"]["paragraph_index"])
    assert ordered == ["p0001_b0001", "p0001_b0002", "p0002_b0002", "p0002_b0003",
                       "p0001_b0003", "p0001_b0004", "p0002_b0001", "p0002_b0004", "p0002_b0005", "p0002_b0006"]
    assert len(by_id) == len(doc.paragraphs) == sum(len(p["blocks"]) for p in targets)
    assert [b["block_id"] for p in mapping["pages"] for b in p["blocks"]] == [b["id"] for p in targets for b in p["blocks"]]
    assert len(mapping["pages"][1]["list_furniture_reflows"]) == 1
    assert not mapping["layout_review_required"]
    assert all("joined_to_block_id" not in block for block in by_id.values())
    assert all(paragraph.paragraph_format.keep_with_next is not True for paragraph in doc.paragraphs)


@pytest.mark.usefixtures("legacy_list_postpass")
def test_chained_lists_carry_each_previous_furniture_packet_to_the_end_once(tmp_path):
    sources = []
    for number in (1, 2, 3):
        body = [dict(text=f"{number}. Condição consecutiva;", role="list_item", bbox=(70, 150 if number > 1 else 740, 530, 780))]
        sources.append(_page(number, body))
    pages = tmp_path / "pages"
    _save(pages, sources)
    output = writer.assemble_docx(pages, tmp_path / "chain.docx", lang=TargetLang.EN, page_breaks=False)
    doc, mapping = _read(output)
    assert [p.text for p in doc.paragraphs[1:4]] == [f"{n}. Condição consecutiva;" for n in (1, 2, 3)]
    assert len(doc.paragraphs) == 12
    assert sum(len(p.get("list_furniture_reflows", [])) for p in mapping["pages"]) == 2
    assert len(_mapped_xml(output)) == 12


@pytest.mark.parametrize("guard", ["source_text", "source_geometry", "contact_changed", "legal_footer", "page_uncertain",
    "block_uncertain", "document_start", "block_start", "wrong_file", "different_size", "low_head", "high_tail",
    "number_reset", "number_gap", "marker_change", "intro_number_reset", "heading", "signature", "reference", "table",
    "no_contact", "missing_bbox", "duplicate_id"])
def test_source_proof_fails_closed_for_ambiguous_or_substantive_boundaries(guard):
    a, b = _pair("intro" if guard == "intro_number_reset" else "number")
    if guard == "source_text": b["blocks"][0]["text"] += " alterado"
    elif guard == "source_geometry": b["blocks"][0]["bbox"] = (200, 350, 400, 375)
    elif guard == "contact_changed": b["blocks"][-1]["text"] = "Telef: 220000000"
    elif guard == "legal_footer":
        for p in (a, b): p["blocks"][-1]["text"] = "Telef: 210000000; deve pagar no prazo de 10 dias."
    elif guard == "page_uncertain": b["uncertain"] = True
    elif guard == "block_uncertain": b["blocks"][1]["uncertain"] = True
    elif guard == "document_start": b["document_start"] = True
    elif guard == "block_start": b["blocks"][0]["document_start"] = True
    elif guard == "wrong_file": b["source_file_sha256"] = "c" * 64
    elif guard == "different_size": b["width_pt"] += 10
    elif guard == "low_head": b["blocks"][1]["bbox"] = (70, 400, 530, 430)
    elif guard == "high_tail": a["blocks"][1]["bbox"] = (70, 300, 530, 330)
    elif guard in {"number_reset", "intro_number_reset"}: b["blocks"][1]["text"] = "4. Reinício não demonstrado;"
    elif guard == "number_gap": b["blocks"][2]["text"] = "9. Salto não demonstrado;"
    elif guard == "marker_change": b["blocks"][1]["text"] = "b) Outro marcador;"
    elif guard in {"heading", "signature", "reference", "table"}:
        row = dict(b["blocks"][1], id="p0002_b0099", role="table_cell" if guard == "table" else guard)
        b["blocks"].insert(1, row)
    elif guard == "no_contact":
        for p in (a, b): p["blocks"][-1]["text"] = "Largo da Justiça, 2"
    elif guard == "missing_bbox": b["blocks"][1]["bbox"] = None
    elif guard == "duplicate_id": b["blocks"][1]["id"] = a["blocks"][1]["id"]
    assert confirmed_list_furniture_reflow(a, b) is None


@pytest.mark.parametrize("guard", ["missing_source", "stale_source", "altered_target_geometry", "altered_target_page_sizes",
                                  "legacy", "explicit_breaks", "gap", "document_start"])
@pytest.mark.usefixtures("legacy_list_postpass")
def test_writer_does_not_reflow_without_complete_binding_or_compact_permission(tmp_path, monkeypatch, guard):
    pages = tmp_path / "pages"
    sources = _pair(next_number=3 if guard == "gap" else 2)
    if guard == "document_start": sources[1]["document_start"] = True
    _save(pages, sources)
    if guard == "missing_source": (pages / "page_0001.source_structure.json").unlink()
    elif guard == "stale_source":
        source = json.loads((pages / "page_0001.source_structure.json").read_text())
        source["blocks"][0]["text"] += " corrupted"
        (pages / "page_0001.source_structure.json").write_text(json.dumps(source))
    elif guard == "altered_target_geometry":
        path = pages / "page_0002.structure.json"
        target = json.loads(path.read_text())
        target["blocks"][0]["bbox"] = [210, 90, 410, 112]
        path.write_text(json.dumps(target))
    elif guard == "altered_target_page_sizes":
        # Equal target sizes evade a target-to-target comparison, but neither
        # may replace the independently bound source-page geometry.
        for path in pages.glob("*.structure.json"):
            if path.name.endswith(".source_structure.json"):
                continue
            target = json.loads(path.read_text())
            target["width_pt"] += 10
            target["height_pt"] += 10
            path.write_text(json.dumps(target))
    elif guard == "legacy":
        for path in pages.glob("*.structure.json"): path.unlink()
    with monkeypatch.context() as patch:
        patch.setattr(writer, "_reflow_compact_lists", lambda *args, **kwargs: None)
        control = writer.assemble_docx(pages, tmp_path / "control.docx", lang=TargetLang.EN, page_breaks=guard == "explicit_breaks")
    output = writer.assemble_docx(pages, tmp_path / "guarded.docx", lang=TargetLang.EN, page_breaks=guard == "explicit_breaks")
    assert [p._element.xml for p in Document(output).paragraphs] == [p._element.xml for p in Document(control).paragraphs]
    _, mapping = _read(output)
    assert all(not page.get("list_furniture_reflows") for page in mapping["pages"])
    if guard in {"missing_source", "stale_source", "altered_target_geometry", "altered_target_page_sizes"}:
        assert mapping["layout_review_required"]
        assert "list_furniture_reflow_not_evaluated" in mapping["pages"][1]["layout_warnings"]


@pytest.mark.parametrize("key,value", [("width_pt", 620), ("height_pt", 880), ("uncertain", True), ("document_start", True)])
def test_full_source_pair_binds_page_geometry_and_page_level_flags(tmp_path, key, value):
    pages_dir = tmp_path / "pages"
    _save(pages_dir, _pair())
    path = pages_dir / "page_0001.structure.json"
    target = json.loads(path.read_text())
    target[key] = value
    path.write_text(json.dumps(target))
    pages = [writer._load_assembly_page(path) for path in sorted(pages_dir.glob("*.txt"))]
    assert writer._validated_source_pair(*pages, pages_dir) is None


@pytest.mark.usefixtures("legacy_list_postpass")
def test_translated_intro_without_colon_only_under_triggers_and_does_not_change_source(tmp_path):
    pages = tmp_path / "pages"
    targets = _save(pages, _pair("intro"))
    target = targets[0]
    target["blocks"][1]["text"] = target["blocks"][1]["text"].removesuffix(":")
    text = "\n".join(block["text"] for block in target["blocks"])
    target["translation_sha256"] = text_sha256(text)
    (pages / "page_0001.txt").write_text(text, encoding="utf-8")
    (pages / "page_0001.structure.json").write_text(json.dumps(target), encoding="utf-8")
    before = {path.name: path.read_bytes() for path in pages.iterdir()}
    output = writer.assemble_docx(pages, tmp_path / "under_trigger.docx", lang=TargetLang.EN, page_breaks=False)
    assert not any(page.get("list_furniture_reflows") for page in _read(output)[1]["pages"])
    assert {path.name: path.read_bytes() for path in pages.iterdir()} == before


@pytest.mark.parametrize("status", ["regions", "needs_review_flow_fallback", "invalid_layout_flow_fallback"])
def test_postpass_never_moves_region_or_fallback_layout_pages(tmp_path, status):
    pages_dir = tmp_path / "pages"
    _save(pages_dir, _pair())
    pages = [writer._load_assembly_page(path) for path in sorted(pages_dir.glob("*.txt"))]
    document = Document()
    mappings = []
    for page in pages:
        rows = []
        for block in page.structure["blocks"]:
            document.add_paragraph(block["text"])
            rows.append({"block_id": block["id"], "location": {"kind": "body_paragraph", "paragraph_index": len(document.paragraphs) - 1}})
        mappings.append({"blocks": rows, "layout_status": status, "layout_review_required": status != "regions", "layout_warnings": []})
    before = document._element.xml
    writer._reflow_compact_lists(document, pages, mappings, pages_dir, page_breaks=False)
    assert document._element.xml == before
    assert all(not page.get("list_furniture_reflows") for page in mappings)
