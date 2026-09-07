from __future__ import annotations

from copy import deepcopy
import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
import pytest

from legalpdf_translate.document_structure import PageStructure, StructureBlock, text_sha256
import legalpdf_translate.docx_writer as writer
from legalpdf_translate.types import TargetLang


def _page(number, *, header="Tribunal Judicial de Cidade", contact="court@example.invalid", start=False,
          signature=False, body=None):
    rows = []
    if signature:
        rows.append(dict(text="Assinado por Pessoa", role="signature", bbox=(30, 20, 300, 40), document_start=start))
    if header:
        rows.append(dict(text=header, role="header", bbox=(100, 90, 500, 110), alignment="center"))
    rows.extend(body or [dict(text=f"Corpo completo da página {number}.", role="paragraph", bbox=(40, 150, 550, 170)),
                             dict(text=f"Outra obrigação {number}.", role="paragraph", bbox=(40, 200, 550, 220))])
    if contact:
        rows.extend([dict(text="Largo da Justiça, 1", role="footer", bbox=(160, 800, 410, 810), alignment="center"),
                     dict(text=f"Telef: 210000000 - E-mail: {contact}", role="footer", bbox=(100, 812, 500, 822), alignment="center")])
    source = PageStructure(page_number=number, source_sha256="a" * 64, source_file_sha256="b" * 64,
                           document_start=start,
                           blocks=[StructureBlock(f"p{number:04d}_b{index:04d}", **row) for index, row in enumerate(rows, 1)])
    source.source_sha256 = source.source_text_sha256 = text_sha256(source.text)
    return source.to_dict()


def _save(folder, sources, *, variants=False, arabic=False):
    folder.mkdir(exist_ok=True)
    targets = []
    for source in sources:
        number = source["page_number"]
        target = deepcopy(source)
        for row in target["blocks"]:
            if variants and row["role"] in {"header", "footer"}:
                row["text"] += f" variant {number}"
            elif arabic and row["role"] == "paragraph":
                row["text"] = "هذا نص عربي مع [[João Exemplo]] " + row["id"]
        text = "\n".join(row["text"] for row in target["blocks"])
        target["translation_sha256"] = text_sha256(text)
        (folder / f"page_{number:04d}.source_structure.json").write_text(json.dumps(source), encoding="utf-8")
        path = folder / f"page_{number:04d}.txt"
        path.write_text(text, encoding="utf-8")
        path.with_suffix(".structure.json").write_text(json.dumps(target), encoding="utf-8")
        targets.append(target)
    return targets


def _read(path):
    return Document(path), json.loads(path.with_suffix(".source_map.json").read_text(encoding="utf-8"))


def _mapped_text(doc, row):
    loc = row["location"]
    if loc["kind"] == "body_paragraph":
        return doc.paragraphs[loc["paragraph_index"]].text
    part = getattr(doc.sections[loc["section_index"]], loc["kind"].removeprefix("section_"))
    assert str(part.part.partname) == loc["part_uri"]
    return part.paragraphs[loc["paragraph_index"]].text


def test_repeated_furniture_uses_real_parts_and_explicit_complete_aliases(tmp_path):
    folder = tmp_path / "pages"
    targets = _save(folder, [_page(1, start=True, signature=True), _page(2)], variants=True)
    before = {p.name: p.read_bytes() for p in folder.iterdir()}
    output = writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False)
    doc, mapping = _read(output)
    assert len(doc.sections) == 1
    assert len(doc.sections[0].header.paragraphs) == 1
    assert len(doc.sections[0].footer.paragraphs) == 3  # Two contacts and PAGE.
    assert " PAGE " in doc.sections[0].footer._element.xml
    assert all(not getattr(doc.sections[0], part).is_linked_to_previous for part in ("header", "footer"))
    assert "Assinado por Pessoa" in [p.text for p in doc.paragraphs]
    assert not any("Tribunal" in p.text or "E-mail" in p.text for p in doc.paragraphs)
    rows = [r for p in mapping["pages"] for r in p["blocks"]]
    expected = {b["id"]: b for p in targets for b in p["blocks"]}
    assert len(rows) == len(expected) == len({r["block_id"] for r in rows})
    for row in rows:
        if "furniture_alias" in row:
            alias = row["furniture_alias"]
            assert alias["original_target_text"] == expected[row["block_id"]]["text"]
            assert _mapped_text(doc, row) == expected[alias["canonical_block_id"]]["text"]
        else:
            assert _mapped_text(doc, row) == expected[row["block_id"]]["text"]
    assert sum(row.get("furniture_alias", {}).get("target_variant", False) for row in rows) == 3
    assert mapping["layout_review_required"]
    assert all("section_furniture_target_variant_standardized" in p["layout_warnings"] for p in mapping["pages"])
    assert all(not p.get("list_furniture_reflows") for p in mapping["pages"])
    assert before == {p.name: p.read_bytes() for p in folder.iterdir()}
    assert mapping["section_furniture"]["policy"] == "source_section_furniture_v1"


@pytest.mark.parametrize("middle", ["none", "headerless", "footerless", "changed_contact"])
def test_changed_or_absent_parts_create_unlinked_sections_without_double_breaks(tmp_path, middle):
    sources = [_page(1, start=True), _page(2)]
    settings = {"none": dict(header=None, contact=None), "headerless": dict(header=None),
                "footerless": dict(contact=None), "changed_contact": dict(contact="different@example.invalid")}[middle]
    sources += [_page(3, **settings), _page(4, **settings)]
    sources += [_page(5, start=True), _page(6)]
    folder = tmp_path / "pages"
    _save(folder, sources)
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False))
    assert len(doc.sections) == 3
    assert not doc._element.xpath(".//w:br[@w:type='page']")
    assert len(doc._element.xpath(".//w:pPr/w:sectPr")) == 2
    for section in doc.sections:
        assert not section.header.is_linked_to_previous
        assert not section.footer.is_linked_to_previous
        assert not section.different_first_page_header_footer
        assert section._sectPr.find(qn("w:pgNumType")) is None
        assert " PAGE " in section.footer._element.xml
    middle_section = doc.sections[1]
    assert bool(middle_section.header.paragraphs[0].text) is (middle not in {"none", "headerless"})
    assert ("E-mail" in middle_section.footer._element.xml) is (middle not in {"none", "footerless"})
    assert "different@example.invalid" in middle_section.footer._element.xml if middle == "changed_contact" else True
    assert all(_mapped_text(doc, b) for p in mapping["pages"] for b in p["blocks"])


def test_partial_run_does_not_inherit_an_unrepresented_previous_header(tmp_path):
    folder = tmp_path / "pages"
    _save(folder, [_page(7, header=None), _page(8, header=None)])
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False))
    assert len(doc.sections) == 1 and not doc.sections[0].header.paragraphs[0].text
    assert not mapping["section_furniture"]["sections"][0]["parts"]["header"]
    assert " PAGE " in doc.sections[0].footer._element.xml


@pytest.mark.parametrize("guard", ["explicit", "missing_source", "source_hash", "target_geometry", "single_page"])
def test_unproved_or_explicit_page_matching_keeps_every_furniture_occurrence(tmp_path, guard):
    folder = tmp_path / "pages"
    _save(folder, [_page(1), _page(2)] if guard != "single_page" else [_page(1)])
    if guard == "missing_source":
        (folder / "page_0002.source_structure.json").unlink()
    if guard == "source_hash":
        p = folder / "page_0002.source_structure.json"
        s = json.loads(p.read_text()); s["blocks"][0]["text"] += " changed"
        p.write_text(json.dumps(s))
    if guard == "target_geometry":
        p = folder / "page_0002.structure.json"
        s = json.loads(p.read_text()); s["width_pt"] += 10
        p.write_text(json.dumps(s))
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.EN,
                                            page_breaks=guard == "explicit"))
    if guard in {"explicit", "single_page"}:
        assert "section_furniture" not in mapping
    else:
        assert mapping["layout_review_required"]
        assert not any(s["consolidated"] for s in mapping["section_furniture"]["sections"])
    assert sum("Tribunal" in p.text for p in doc.paragraphs) == (1 if guard == "single_page" else 2)
    assert not any("furniture_alias" in b for p in mapping["pages"] for b in p["blocks"])


@pytest.mark.parametrize("lang,multiplier", [(TargetLang.AR, 1.10), (TargetLang.EN, None), (TargetLang.FR, None)])
def test_spacing_retains_fonts_and_respects_word_collapsed_paragraph_gap(tmp_path, lang, multiplier):
    folder = tmp_path / "pages"
    _save(folder, [_page(1), _page(2)], arabic=lang == TargetLang.AR)
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=lang, page_breaks=False))
    body = [p for p in doc.paragraphs if p.text]
    assert body[0].paragraph_format.line_spacing == multiplier
    assert body[0].style.paragraph_format.space_after.pt == 3
    assert body[1].paragraph_format.space_before.pt == 18  # max(before18, after3), not15+3.
    size = 11 if lang == TargetLang.AR else 10.5
    assert all(run.font.size.pt == size for p in body for run in p.runs)
    assert doc.sections[0].top_margin.pt >= 18 + doc.sections[0].header_distance.pt + size
    assert doc.sections[0].bottom_margin.pt > 42
    assert mapping["pages"][0]["source_spacing"]["line_height_evidence"] == "unavailable"
    reserve = mapping["section_furniture"]["sections"][0]["render_reserve"]
    assert reserve["header_body_gap_pt"] == 18 and reserve["footer_body_gap_pt"] == 6
    assert reserve["basis"] == "conservative_text_wrap_estimate"
    assert reserve["gap_policy"] == "readable_section_reserve_v2"
    assert reserve["word_render_verified"] is False


@pytest.mark.parametrize("lang,font_size", [(TargetLang.AR, 11), (TargetLang.EN, 10.5), (TargetLang.FR, 10.5)])
def test_footer_reserve_keeps_conservative_wrap_page_and_fonts_with_six_point_extra(lang, font_size):
    from docx.shared import Cm

    size = (595.3, 841.7)
    plan = {"parts": {"header": {"canonical_blocks": [{"text": "Synthetic court"}]},
                      "footer": {"canonical_blocks": [{"text": "x" * 100}, {"text": "Short contact"}]}}}
    width = size[0] - 2 * Cm(1.7).pt
    assert writer._furniture_line_estimate(plan["parts"]["footer"]["canonical_blocks"], width, font_size) == 3
    top, bottom = writer._furniture_margins(plan, size, lang)
    assert top == pytest.approx(Cm(0.7).pt + font_size * 1.2 + 18)
    assert bottom == pytest.approx(Cm(0.7).pt + 4 * font_size * 1.2 + 6)  # Three estimated lines plus PAGE.
    assert writer._font_profile(lang)[1] == font_size
    assert writer._furniture_margins(None, size, lang) == (Cm(1.5).pt, Cm(1.5).pt)


def test_explicit_page_matching_keeps_single_spacing_and_old_breaks(tmp_path):
    folder = tmp_path / "pages"
    _save(folder, [_page(1), _page(2)], arabic=True)
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=True))
    assert "section_furniture" not in mapping
    assert all("source_spacing" not in p for p in mapping["pages"])
    assert all(p.paragraph_format.line_spacing is None for p in doc.paragraphs)
    assert len(doc._element.xpath(".//w:br[@w:type='page']")) == 1


def test_confirmed_sentence_join_remains_valid_when_furniture_is_in_parts(tmp_path):
    from legalpdf_translate.formatting_support import link_source_continuations
    sources = [_page(1, body=[dict(text="A parte deve comparecer na audiência de julgamento marcada para", role="paragraph",
                                  bbox=(40, 740, 550, 780), alignment="justify")]),
               _page(2, body=[dict(text="o dia indicado na notificação.", role="paragraph", bbox=(40, 145, 550, 170), alignment="justify")])]
    linked = link_source_continuations({s["page_number"]: PageStructure.from_dict(s) for s in sources})
    folder = tmp_path / "pages"
    _save(folder, [s.to_dict() for s in linked.values()])
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False))
    body = [p.text for p in doc.paragraphs if p.text]
    assert body == ["A parte deve comparecer na audiência de julgamento marcada para o dia indicado na notificação."]
    first, second = [p["blocks"][1] for p in mapping["pages"]]
    assert first["location"] == second["location"]
    assert second["joined_to_block_id"] == first["block_id"]
    assert all(not p.get("list_furniture_reflows") for p in mapping["pages"])


def test_insufficient_body_reserve_falls_back_without_shrinking_or_losing_text(tmp_path):
    sources = [_page(1), _page(2)]
    for source in sources:
        source["height_pt"] = 260
        for block, box in zip(source["blocks"], [(100, 20, 500, 40), (40, 80, 550, 100),
                                               (40, 120, 550, 140), (160, 224, 410, 234), (100, 238, 500, 248)]):
            block["bbox"] = list(box)
    folder = tmp_path / "pages"
    targets = _save(folder, sources)
    for target in targets:
        target["blocks"][0]["text"] += " Traduction" * 25  # Bounded but cannot fit this short page.
        text = "\n".join(b["text"] for b in target["blocks"])
        target["translation_sha256"] = text_sha256(text)
        p = folder / f"page_{target['page_number']:04d}.txt"
        p.write_text(text, encoding="utf-8")
        p.with_suffix(".structure.json").write_text(json.dumps(target), encoding="utf-8")
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False))
    assert mapping["layout_review_required"]
    assert all("section_furniture_reserve_exceeds_page" in p["layout_warnings"] for p in mapping["pages"])
    assert not mapping["section_furniture"]["sections"][0]["consolidated"]
    assert all(b["location"]["kind"] == "body_paragraph" for p in mapping["pages"] for b in p["blocks"])
    assert sum("Traduction" in p.text for p in doc.paragraphs) == 2
    assert all(r.font.size.pt == 11 for p in doc.paragraphs for r in p.runs)
    assert doc.sections[0].top_margin.pt < 43 and doc.sections[0].bottom_margin.pt < 43


def test_missing_middle_source_ends_previous_parts_and_keeps_unproved_occurrences(tmp_path):
    folder = tmp_path / "pages"
    _save(folder, [_page(n) for n in (1, 2, 3, 4)])
    (folder / "page_0003.source_structure.json").unlink()
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False))
    assert len(doc.sections) == 2
    assert doc.sections[0].header.paragraphs[0].text
    assert not doc.sections[1].header.paragraphs[0].text
    assert "E-mail" not in doc.sections[1].footer._element.xml
    assert sum("Tribunal" in p.text for p in doc.paragraphs) == 2
    assert [p["section_furniture_adopted"] for p in mapping["pages"]] == [True, True, False, False]
    assert not doc._element.xpath(".//w:br[@w:type='page']")


def test_equal_furniture_across_independent_documents_has_separate_canonical_parts(tmp_path):
    folder = tmp_path / "pages"
    _save(folder, [_page(n, start=n in (1, 3)) for n in (1, 2, 3, 4)], variants=True)
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False))
    assert len(doc.sections) == 2
    assert "variant 1" in doc.sections[0].header.paragraphs[0].text
    assert "variant 3" in doc.sections[1].header.paragraphs[0].text
    assert [p["section_index"] for p in mapping["pages"]] == [0, 0, 1, 1]
    assert len({str(s.header.part.partname) for s in doc.sections}) == 2


def test_table_before_section_boundary_retains_cell_locator_without_extra_page_break(tmp_path):
    sources = [_page(1), _page(2, body=[dict(text="Cell meaning", role="table_cell", table_id="data",
                                          row=0, col=0, bbox=(40, 150, 550, 180))]),
               _page(3, header=None, contact=None)]
    folder = tmp_path / "pages"
    _save(folder, sources)
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False))
    assert len(doc.sections) == 2 and len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "Cell meaning"
    cell = next(b for b in mapping["pages"][1]["blocks"] if b["role"] == "table_cell")
    assert cell["location"]["table_index"] == 0
    assert not doc._element.xpath(".//w:br[@w:type='page']")


def test_legacy_txt_does_not_gain_sections_or_source_spacing(tmp_path):
    folder = tmp_path / "pages"; folder.mkdir()
    for n in (1, 2):
        (folder / f"page_{n:04d}.txt").write_text("النص القديم", encoding="utf-8")
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False))
    assert len(doc.sections) == 1 and "section_furniture" not in mapping
    assert all("source_spacing" not in p for p in mapping["pages"])
    assert all(p.paragraph_format.line_spacing is None for p in doc.paragraphs)


def test_all_oversized_furniture_rejections_keep_review_reasons_and_full_body(tmp_path):
    folder = tmp_path / "pages"
    targets = _save(folder, [_page(1), _page(2)])
    for target in targets:
        target["blocks"][0]["text"] = "Synthetic translated heading " * 60
        text = "\n".join(b["text"] for b in target["blocks"])
        target["translation_sha256"] = text_sha256(text)
        p = folder / f"page_{target['page_number']:04d}.txt"
        p.write_text(text, encoding="utf-8")
        p.with_suffix(".structure.json").write_text(json.dumps(target), encoding="utf-8")
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False))
    assert mapping["layout_review_required"]
    assert all("section_furniture_source_evidence_unavailable" in p["layout_warnings"] for p in mapping["pages"])
    assert "section_furniture" in mapping
    assert not any(s["consolidated"] for s in mapping["section_furniture"]["sections"])
    assert sum("Synthetic translated heading" in p.text for p in doc.paragraphs) == 2
    assert all(b["location"]["kind"] == "body_paragraph" for p in mapping["pages"] for b in p["blocks"])


def _cohesion_page(kind):
    if kind == "signature":
        rows = [dict(text="Cidade, na data indicada", role="signature", alignment="center", bbox=(225, 220, 375, 232)),
                dict(text="A pessoa responsável,", role="signature", alignment="center", bbox=(225, 290, 375, 302)),
                dict(text="João Exemplo", role="signature", alignment="center", bbox=(260, 318, 340, 330))]
    else:
        rows = [dict(text="Observe-se:", role="paragraph", alignment="left", bbox=(40, 150, 120, 162)),
                dict(text="• Primeira condição completa.", role="list_item", alignment="justify", bbox=(40, 171, 520, 196)),
                dict(text="• Segunda condição completa.", role="list_item", alignment="justify", bbox=(40, 205, 520, 230))]
    return _page(1, header=None, contact=None, body=rows)


def _save_cohesion(folder, source, kind):
    source["source_sha256"] = source["source_text_sha256"] = text_sha256("\n".join(b["text"] for b in source["blocks"]))
    target = _save(folder, [source])[0]
    translations = (["[[Cidade]], data", "الشخص المسؤول،", "[[João Exemplo]]"] if kind == "signature"
                    else ["لننظر:", "• الشرط الأول كامل.", "• الشرط الثاني كامل."])
    for block, text in zip(target["blocks"], translations):
        block["text"] = text
    text = "\n".join(b["text"] for b in target["blocks"])
    target["translation_sha256"] = text_sha256(text)
    path = folder / "page_0001.txt"
    path.write_text(text, encoding="utf-8")
    path.with_suffix(".structure.json").write_text(json.dumps(target), encoding="utf-8")
    return target


def test_centered_signature_role_and_name_stay_together_without_city_drag_or_realignment(tmp_path):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    folder = tmp_path / "pages"
    _save_cohesion(folder, _cohesion_page("signature"), "signature")
    inputs = {p.name: p.read_bytes() for p in folder.iterdir()}
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False))
    city, role, name = doc.paragraphs
    assert not city.paragraph_format.keep_with_next
    assert role.paragraph_format.keep_with_next is True
    assert name.paragraph_format.keep_with_next is False
    assert all(p.alignment == WD_ALIGN_PARAGRAPH.CENTER for p in (city, role, name))
    assert [p.text for p in doc.paragraphs] == ["Cidade, data", "الشخص المسؤول،", "João Exemplo"]
    assert len(name.runs) == 1 and name.runs[0]._r.rPr.find(qn("w:rtl")).get(qn("w:val")) == "0"
    assert all(r.font.size.pt == 11 for p in doc.paragraphs for r in p.runs)
    assert mapping["pages"][0]["source_cohesion"][0]["kind"] == "centered_signature"
    assert inputs == {p.name: p.read_bytes() for p in folder.iterdir()}


@pytest.mark.parametrize("guard", ["explicit", "missing_source", "source_hash", "target_geometry", "uncertain_page",
                                  "uncertain_name", "document_start", "distant_name", "different_center", "wrong_role"])
def test_centered_signature_cohesion_requires_complete_bound_nearby_same_group(tmp_path, guard):
    source = _cohesion_page("signature")
    name = source["blocks"][2]
    if guard == "uncertain_page":
        source["uncertain"] = True
    if guard == "uncertain_name":
        name["uncertain"] = True
    if guard == "document_start":
        name["document_start"] = True
    if guard == "distant_name":
        name["bbox"] = [260, 350, 340, 362]
    if guard == "different_center":
        name["bbox"] = [280, 318, 360, 330]
    if guard == "wrong_role":
        name["role"] = "reference"
    folder = tmp_path / "pages"
    target = _save_cohesion(folder, source, "signature")
    source_path = folder / "page_0001.source_structure.json"
    if guard == "missing_source":
        source_path.unlink()
    if guard == "source_hash":
        source["blocks"][0]["text"] += " stale"
        source_path.write_text(json.dumps(source), encoding="utf-8")
    if guard == "target_geometry":
        target["width_pt"] += 1
        (folder / "page_0001.structure.json").write_text(json.dumps(target), encoding="utf-8")
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.AR,
                                            page_breaks=guard == "explicit"))
    role = next(p for p in doc.paragraphs if p.text == "الشخص المسؤول،")
    assert not role.paragraph_format.keep_with_next
    assert "source_cohesion" not in mapping["pages"][0]
    assert "João Exemplo" in [p.text for p in doc.paragraphs]


def test_source_confirmed_short_list_introduction_keeps_only_first_item(tmp_path):
    folder = tmp_path / "pages"
    _save_cohesion(folder, _cohesion_page("list"), "list")
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False))
    intro, first, second = doc.paragraphs
    assert intro.paragraph_format.keep_with_next is True
    assert not first.paragraph_format.keep_with_next and not second.paragraph_format.keep_with_next
    assert [p.text for p in doc.paragraphs] == ["لننظر:", "• الشرط الأول كامل.", "• الشرط الثاني كامل."]
    assert len(mapping["pages"][0]["blocks"]) == 3
    assert mapping["pages"][0]["source_cohesion"][0]["kind"] == "short_list_introduction"
    assert all(r.font.size.pt == 11 for p in doc.paragraphs for r in p.runs)


@pytest.mark.parametrize("guard", ["explicit", "missing_source", "no_source_colon", "long_intro", "multiline_intro",
                                  "missing_marker", "not_list", "uncertain_item", "document_start", "large_gap", "other_column"])
def test_list_intro_cohesion_never_infers_from_target_colon_alone_or_chains_other_groups(tmp_path, guard):
    source = _cohesion_page("list")
    intro, item = source["blocks"][:2]
    if guard == "no_source_colon":
        intro["text"] = "Observe-se."
    if guard == "long_intro":
        intro["text"] = "Longa explicação completa " * 5 + ":"
    if guard == "multiline_intro":
        intro["text"] = "Outra indicação\nObserve-se:"
    if guard == "missing_marker":
        item["text"] = "Primeira condição completa."
    if guard == "not_list":
        item["role"] = "paragraph"
    if guard == "uncertain_item":
        item["uncertain"] = True
    if guard == "document_start":
        item["document_start"] = True
    if guard == "large_gap":
        item["bbox"] = [40, 200, 520, 220]
    if guard == "other_column":
        item["bbox"] = [300, 171, 520, 196]
    folder = tmp_path / "pages"
    _save_cohesion(folder, source, "list")
    if guard == "missing_source":
        (folder / "page_0001.source_structure.json").unlink()
    doc, mapping = _read(writer.assemble_docx(folder, tmp_path / "out.docx", lang=TargetLang.AR,
                                            page_breaks=guard == "explicit"))
    assert not doc.paragraphs[0].paragraph_format.keep_with_next
    assert "source_cohesion" not in mapping["pages"][0]
    assert len([p for p in doc.paragraphs if p.text]) == 3
