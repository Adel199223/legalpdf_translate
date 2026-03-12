from __future__ import annotations

import json
import os
import sys
from datetime import date, datetime
from pathlib import Path
from types import SimpleNamespace

if os.name != "nt" and "DISPLAY" not in os.environ:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PySide6.QtCore import QDate, QRect, Qt, QUrl
from PySide6.QtWidgets import QApplication, QMessageBox

import legalpdf_translate.user_settings as user_settings
from legalpdf_translate.honorarios_docx import (
    HonorariosDraft,
    HonorariosKind,
    build_honorarios_draft,
    build_interpretation_honorarios_draft,
    build_honorarios_paragraph_texts,
    default_interpretation_recipient_block,
    default_honorarios_filename,
    format_portuguese_date,
    generate_honorarios_docx,
)
from legalpdf_translate.user_profile import default_primary_profile
from legalpdf_translate.joblog_db import insert_job_run, open_job_log
from legalpdf_translate.qt_gui.dialogs import (
    JobLogSeed,
    QtHonorariosExportDialog,
    QtJobLogWindow,
    QtProfileManagerDialog,
    QtSaveToJobLogDialog,
    _default_documents_dir,
)
from legalpdf_translate.qt_gui.guarded_inputs import CALENDAR_WEEKEND_COLOR, GuardedDateEdit


def _write_docx_with_paragraphs(path: Path, *paragraphs: str) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def _profile(**overrides: str):
    profile = default_primary_profile(email="adel@example.com")
    for key, value in overrides.items():
        setattr(profile, key, value)
    return profile


def _patch_settings_file(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> Path:
    settings_file = tmp_path / "settings.json"
    monkeypatch.setattr(user_settings, "settings_path", lambda: settings_file)
    return settings_file


def _write_legacy_blank_primary_settings(settings_file: Path) -> None:
    profile = default_primary_profile()
    settings_file.parent.mkdir(parents=True, exist_ok=True)
    settings_file.write_text(
        json.dumps(
            {
                "settings_schema_version": 8,
                "profiles": [
                    {
                        "id": profile.id,
                        "first_name": profile.first_name,
                        "last_name": profile.last_name,
                        "document_name_override": profile.document_name_override,
                        "email": profile.email,
                        "phone_number": profile.phone_number,
                        "postal_address": profile.postal_address,
                        "iban": profile.iban,
                        "iva_text": profile.iva_text,
                        "irs_text": profile.irs_text,
                        "travel_origin_label": "",
                        "travel_distances_by_city": {},
                    }
                ],
                "primary_profile_id": profile.id,
            }
        ),
        encoding="utf-8",
    )


def test_format_portuguese_date_uses_full_month_name() -> None:
    assert format_portuguese_date(date(2026, 3, 6)) == "06 de março de 2026"


def test_build_honorarios_paragraph_texts_uses_locked_template() -> None:
    profile = _profile()
    draft = build_honorarios_draft(
        case_number="109/26.0PBBJA",
        word_count=1666,
        case_entity="Juízo Local Criminal de Beja",
        case_city="Beja",
        profile=profile,
        today=date(2026, 3, 6),
    )

    paragraphs = build_honorarios_paragraph_texts(draft)
    assert paragraphs == [
        ("Número de processo: 109/26.0PBBJA", "left"),
        ("", "left"),
        ("Exmo. Sr(a). Procurador(a) da república do Juízo Local Criminal de Beja", "address"),
        ("", "left"),
        (f"Nome: {profile.document_name}", "left"),
        (f"Morada: {profile.postal_address}", "left"),
        ("", "left"),
        (
            "Venho por este meio requerer o pagamento dos honorários devidos, em virtude de ter sido nomeado "
            "tradutor no âmbito do processo acima identificado.",
            "left",
        ),
        ("O documento traduzido contém 1666 palavras.", "left"),
        ("Este serviço inclui a taxa IVA de 23% e não tem retenção de IRS.", "left"),
        (f"O Pagamento deverá ser efetuado para o seguinte IBAN: {profile.iban}", "left"),
        ("", "left"),
        ("Melhores cumprimentos,", "left"),
        ("", "left"),
        ("Espera deferimento,", "center"),
        ("", "left"),
        ("Beja, 06 de março de 2026", "center"),
        ("", "left"),
        (profile.document_name, "center"),
    ]


def test_generate_honorarios_docx_has_expected_text_and_alignment(tmp_path: Path) -> None:
    profile = _profile()
    draft = HonorariosDraft(
        case_number="109/26.0PBBJA",
        word_count=1666,
        case_entity="Juízo Local Criminal de Beja",
        case_city="Beja",
        date_pt="06 de março de 2026",
        profile=profile,
    )
    output = tmp_path / "honorarios.docx"

    generate_honorarios_docx(draft, output)

    doc = Document(output)
    paragraphs = doc.paragraphs
    assert paragraphs[0].text == "Número de processo: 109/26.0PBBJA"
    assert paragraphs[2].text == "Exmo. Sr(a). Procurador(a) da república do Juízo Local Criminal de Beja"
    assert "\n" not in paragraphs[2].text
    assert paragraphs[12].text == "Melhores cumprimentos,"
    assert paragraphs[14].text == "Espera deferimento,"
    assert paragraphs[16].text == "Beja, 06 de março de 2026"
    assert paragraphs[18].text == profile.document_name
    assert paragraphs[14].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraphs[16].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraphs[18].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraphs[0].runs[0].font.name == "Arial"


def test_generate_honorarios_docx_avoids_overwriting_existing_translation(tmp_path: Path) -> None:
    profile = _profile()
    draft = HonorariosDraft(
        case_number="21/25.0FBPTM",
        word_count=264,
        case_entity="Tribunal Judicial da Comarca de Beja",
        case_city="Beja",
        date_pt="08 de março de 2026",
        profile=profile,
    )
    translated = tmp_path / "21-25_AR_20260308_064339.docx"
    translated.write_bytes(b"translated-bytes")
    expected_name = Path(default_honorarios_filename(draft.case_number)).stem

    saved = generate_honorarios_docx(draft, translated)

    assert translated.read_bytes() == b"translated-bytes"
    assert saved != translated.resolve()
    assert saved.stem.startswith(expected_name)
    assert saved.exists()
    doc = Document(saved)
    assert any("264 palavras" in paragraph.text for paragraph in doc.paragraphs)


def test_build_honorarios_paragraph_texts_keeps_case_city_only_in_closing_date() -> None:
    draft = build_honorarios_draft(
        case_number="21/25.0FBPTM",
        word_count=306,
        case_entity="Tribunal Judicial da Comarca de Beja",
        case_city="Moura",
        profile=_profile(),
        today=date(2026, 3, 7),
    )

    paragraphs = build_honorarios_paragraph_texts(draft)

    assert paragraphs[2] == (
        "Exmo. Sr(a). Procurador(a) da república do Tribunal Judicial da Comarca de Beja",
        "address",
    )
    assert "\n" not in paragraphs[2][0]
    assert "de Moura" not in paragraphs[2][0]
    assert paragraphs[16] == ("Moura, 07 de março de 2026", "center")


def test_build_honorarios_paragraph_texts_uses_case_city_for_plain_ministerio_publico() -> None:
    draft = build_honorarios_draft(
        case_number="1117/25.4T8BJA",
        word_count=120,
        case_entity="Ministério Público",
        case_city="Beja",
        profile=_profile(),
        today=date(2026, 3, 10),
    )

    paragraphs = build_honorarios_paragraph_texts(draft)

    assert paragraphs[2] == (
        "Exmo. Sr(a). Procurador(a) da república do Ministério Público de Beja",
        "address",
    )
    assert "\n" not in paragraphs[2][0]


def test_default_honorarios_filename_sanitizes_case_number() -> None:
    assert (
        default_honorarios_filename("109/26.0PBBJA", today=date(2026, 3, 6))
        == "Requerimento_Honorarios_109_26.0PBBJA_20260306.docx"
    )


def test_default_honorarios_filename_adds_interpretation_suffix() -> None:
    assert (
        default_honorarios_filename(
            "109/26.0PBBJA",
            today=date(2026, 3, 6),
            kind=HonorariosKind.INTERPRETATION,
        )
        == "Requerimento_Honorarios_Interpretacao_109_26.0PBBJA_20260306.docx"
    )


def test_build_interpretation_honorarios_paragraph_texts_uses_interpretation_template() -> None:
    profile = _profile(travel_origin_label="Marmelar")
    draft = build_interpretation_honorarios_draft(
        case_number="000055/25.5GAFAL",
        case_entity="Ministério Público de Beja",
        case_city="Beja",
        service_date="2025-04-09",
        service_entity="GNR",
        service_city="Vidigueira",
        use_service_location_in_honorarios=True,
        travel_km_outbound=50,
        travel_km_return=50,
        recipient_block=default_interpretation_recipient_block("Ministério Público de Beja"),
        profile=profile,
        today=date(2025, 4, 28),
    )

    paragraphs = build_honorarios_paragraph_texts(draft)

    assert paragraphs[2] == ("Exmo. Senhor Procurador do Ministério Público de Beja", "address")
    assert "intérprete" in paragraphs[7][0]
    assert "no dia 09/04/2025, na GNR de Vidigueira" in paragraphs[7][0]
    assert "entre Marmelar e Vidigueira" in paragraphs[7][0]
    assert "50 km em cada sentido" in paragraphs[7][0]
    assert "não está sujeito a retenção de IRS" in paragraphs[8][0]
    assert paragraphs[13] == ("Melhores cumprimentos,", "left")
    assert paragraphs[14] == ("Pede deferimento.", "left")
    assert paragraphs[16] == ("Beja, 09 de abril de 2025", "center")
    assert paragraphs[18] == ("O Requerente,", "center")
    assert paragraphs[20] == (profile.document_name, "center")


def test_default_interpretation_recipient_block_uses_case_city_for_plain_ministerio_publico() -> None:
    recipient = default_interpretation_recipient_block("Ministério Público", "Beja")

    assert recipient == "Exmo. Senhor Procurador do Ministério Público de Beja"
    assert "\n" not in recipient


def test_build_interpretation_honorarios_draft_falls_back_to_service_city_for_case_city() -> None:
    draft = build_interpretation_honorarios_draft(
        case_number="1117/25.4TBJA",
        case_entity="Ministério Público",
        case_city="",
        service_date="2026-03-26",
        service_entity="Ministério Público",
        service_city="Beja",
        profile=_profile(travel_origin_label="Marmelar"),
    )

    paragraphs = build_honorarios_paragraph_texts(draft)

    assert draft.case_city == "Beja"
    assert paragraphs[2] == ("Exmo. Senhor Procurador do Ministério Público de Beja", "address")
    assert paragraphs[16] == ("Beja, 26 de março de 2026", "center")


def test_build_interpretation_honorarios_draft_uses_service_date_for_footer_when_valid() -> None:
    draft = build_interpretation_honorarios_draft(
        case_number="109/26.0PBBJA",
        case_entity="Juízo Local Criminal de Beja",
        case_city="Beja",
        service_date="2026-03-09",
        service_entity="Juízo Local Criminal de Beja",
        service_city="Beja",
        travel_km_outbound=39,
        travel_km_return=39,
        profile=_profile(),
        today=date(2026, 3, 10),
    )

    assert draft.date_pt == "09 de março de 2026"


def test_build_interpretation_honorarios_short_form_omits_service_location_phrase() -> None:
    profile = _profile(travel_origin_label="Marmelar")
    draft = build_interpretation_honorarios_draft(
        case_number="000055/25.5GAFAL",
        case_entity="Juízo Local Criminal de Beja",
        case_city="Beja",
        service_date="2025-04-09",
        service_entity="Juízo Local Criminal",
        service_city="Beja",
        use_service_location_in_honorarios=False,
        travel_km_outbound=39,
        travel_km_return=39,
        profile=profile,
        today=date(2025, 4, 28),
    )

    body_text = next(
        text
        for text, _kind in build_honorarios_paragraph_texts(draft)
        if "Venho, por este meio, requerer o pagamento dos honorários devidos" in text
    )
    assert "na GNR de" not in body_text
    assert "na cidade de" not in body_text
    assert "entre Marmelar e Beja" in body_text


def test_honorarios_dialog_validates_required_fields(tmp_path: Path, monkeypatch) -> None:
    _patch_settings_file(monkeypatch, tmp_path)
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    dialog = QtHonorariosExportDialog(
        parent=None,
        draft=HonorariosDraft("", 0, "", "", "06 de março de 2026", _profile()),
        default_directory=tmp_path,
    )
    try:
        with pytest.raises(ValueError, match="Número de processo"):
            dialog._build_draft()
        dialog.case_number_edit.setText("109/26.0PBBJA")
        dialog.case_entity_edit.setText("Juízo Local Criminal de Beja")
        dialog.case_city_edit.setText("Beja")
        with pytest.raises(ValueError, match="greater than zero"):
            dialog._build_draft()
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()


def test_honorarios_dialog_reports_auto_renamed_save_path(tmp_path: Path, monkeypatch) -> None:
    _patch_settings_file(monkeypatch, tmp_path)
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    translated = tmp_path / "21-25_AR_20260308_064339.docx"
    translated.write_bytes(b"translated-bytes")
    requested = str(translated)
    infos: list[str] = []
    expected_name = Path(default_honorarios_filename("21/25.0FBPTM")).stem

    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QFileDialog.getSaveFileName",
        lambda *args, **kwargs: (requested, "Word Document (*.docx)"),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QMessageBox.information",
        lambda *args, **kwargs: infos.append(args[2] if len(args) > 2 else kwargs.get("text", "")),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QMessageBox.question",
        lambda *args, **kwargs: 65536,
    )

    dialog = QtHonorariosExportDialog(
        parent=None,
        draft=HonorariosDraft(
            case_number="21/25.0FBPTM",
            word_count=264,
            case_entity="Tribunal Judicial da Comarca de Beja",
            case_city="Beja",
            date_pt="08 de março de 2026",
            profile=_profile(),
        ),
        default_directory=tmp_path,
    )
    try:
        dialog._generate()
        assert dialog.saved_path is not None
        assert dialog.saved_path != translated.resolve()
        assert dialog.saved_path.stem.startswith(expected_name)
        assert infos
        assert str(dialog.saved_path) in infos[0]
        assert translated.read_bytes() == b"translated-bytes"
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()


def test_honorarios_dialog_defaults_to_primary_profile_and_secondary_override_is_one_off(
    tmp_path: Path,
    monkeypatch,
) -> None:
    _patch_settings_file(monkeypatch, tmp_path)
    primary = _profile(id="primary", email="primary@example.com")
    secondary = _profile(
        id="secondary",
        first_name="Jane",
        last_name="Doe",
        document_name_override="Jane Doe",
        email="jane@example.com",
        postal_address="Rua B",
        iban="PT50003506490000832760030",
        iva_text="6%",
        irs_text="Isento",
    )
    user_settings.save_profile_settings(profiles=[primary, secondary], primary_profile_id=primary.id)

    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    draft = HonorariosDraft(
        case_number="109/26.0PBBJA",
        word_count=1666,
        case_entity="Juízo Local Criminal de Beja",
        case_city="Beja",
        date_pt="06 de março de 2026",
        profile=primary,
    )

    first_dialog = QtHonorariosExportDialog(parent=None, draft=draft, default_directory=tmp_path)
    second_dialog = None
    try:
        assert first_dialog.profile_combo.currentData() == "primary"
        first_dialog.profile_combo.setCurrentIndex(first_dialog.profile_combo.findData("secondary"))
        rebuilt = first_dialog._build_draft()
        assert rebuilt.profile.id == "secondary"
        assert rebuilt.profile.document_name == "Jane Doe"

        second_dialog = QtHonorariosExportDialog(parent=None, draft=draft, default_directory=tmp_path)
        assert second_dialog.profile_combo.currentData() == "primary"
    finally:
        first_dialog.close()
        first_dialog.deleteLater()
        if second_dialog is not None:
            second_dialog.close()
            second_dialog.deleteLater()
        if owns_app:
            app.quit()


def test_honorarios_dialog_small_screen_uses_scrollable_body_and_fixed_action_bar(
    tmp_path: Path, monkeypatch
) -> None:
    _patch_settings_file(monkeypatch, tmp_path)
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.window_adaptive.available_screen_geometry",
        lambda _widget: QRect(0, 0, 820, 620),
    )

    draft = build_interpretation_honorarios_draft(
        case_number="109/26.0PBBJA",
        case_entity="Juízo Local Criminal de Beja",
        case_city="Beja",
        service_date="2026-03-09",
        service_entity="Juízo Local Criminal de Beja",
        service_city="Beja",
        use_service_location_in_honorarios=False,
        travel_km_outbound=0.0,
        travel_km_return=0.0,
        recipient_block=default_interpretation_recipient_block("Juízo Local Criminal de Beja"),
        profile=_profile(),
    )

    dialog = QtHonorariosExportDialog(parent=None, draft=draft, default_directory=tmp_path)
    try:
        dialog.show()
        app.processEvents()
        assert dialog.profile_combo.isEditable() is False
        assert dialog.generate_btn.objectName() == "PrimaryButton"
        assert dialog.width() <= 672
        assert dialog.height() <= 545
        assert dialog.form_scroll_area.widgetResizable() is True
        assert dialog.cancel_btn.parentWidget() is dialog.action_bar
        assert dialog.generate_btn.parentWidget() is dialog.action_bar
        assert dialog.recipient_block_edit.minimumHeight() == 72
        assert dialog.service_same_check.isChecked() is True
        assert dialog.travel_km_return_edit is dialog.travel_km_outbound_edit
        assert dialog.travel_km_outbound_edit.text() == "39"
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()


def test_honorarios_dialog_interpretation_defaults_service_same_and_one_way_distance(
    tmp_path: Path, monkeypatch
) -> None:
    _patch_settings_file(monkeypatch, tmp_path)
    user_settings.save_profile_settings(profiles=[_profile()], primary_profile_id="primary")
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    draft = build_interpretation_honorarios_draft(
        case_number="109/26.0PBBJA",
        case_entity="Juízo Local Criminal de Beja",
        case_city="Beja",
        service_date="2026-03-09",
        service_entity="",
        service_city="",
        use_service_location_in_honorarios=False,
        travel_km_outbound=0.0,
        travel_km_return=0.0,
        recipient_block=default_interpretation_recipient_block("Juízo Local Criminal de Beja"),
        profile=_profile(),
    )

    dialog = QtHonorariosExportDialog(parent=None, draft=draft, default_directory=tmp_path)
    try:
        dialog.show()
        app.processEvents()
        assert isinstance(dialog.service_date_edit, GuardedDateEdit)
        dialog.service_date_edit.setCalendarDate(QDate(2026, 3, 11))
        assert dialog.service_date_edit.text() == "2026-03-11"
        calendar = dialog.service_date_edit.calendarWidget()
        assert calendar.firstDayOfWeek() == Qt.DayOfWeek.Monday
        assert calendar.minimumWidth() >= 336
        calendar.setCurrentPage(2026, 3)
        app.processEvents()
        assert (
            calendar.dateTextFormat(QDate(2026, 2, 28)).foreground().color().name().lower()
            == CALENDAR_WEEKEND_COLOR.lower()
        )
        assert dialog.service_same_check.isChecked() is True
        assert dialog.service_entity_edit.text() == "Juízo Local Criminal de Beja"
        assert dialog.service_city_edit.text() == "Beja"
        assert dialog.service_entity_edit.isEnabled() is False
        assert dialog.service_city_edit.isEnabled() is False
        assert dialog.travel_km_outbound_edit.text() == "39"
        assert dialog.travel_km_return_edit is dialog.travel_km_outbound_edit
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()


def test_honorarios_dialog_interpretation_recipient_default_tracks_case_city(
    tmp_path: Path,
    monkeypatch,
) -> None:
    _patch_settings_file(monkeypatch, tmp_path)
    user_settings.save_profile_settings(profiles=[_profile()], primary_profile_id="primary")
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    draft = build_interpretation_honorarios_draft(
        case_number="1117/25.4TBJA",
        case_entity="Ministério Público",
        case_city="Beja",
        service_date="2026-03-26",
        service_entity="Ministério Público",
        service_city="Beja",
        use_service_location_in_honorarios=False,
        travel_km_outbound=39.0,
        travel_km_return=39.0,
        recipient_block="",
        profile=_profile(),
    )

    dialog = QtHonorariosExportDialog(parent=None, draft=draft, default_directory=tmp_path)
    try:
        dialog.show()
        app.processEvents()
        assert dialog.recipient_block_edit.toPlainText() == "Exmo. Senhor Procurador do Ministério Público de Beja"
        dialog.case_city_edit.setText("Serpa")
        app.processEvents()
        assert dialog.recipient_block_edit.toPlainText() == "Exmo. Senhor Procurador do Ministério Público de Serpa"
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()


def test_honorarios_dialog_interpretation_uses_repaired_legacy_primary_profile_distance(
    tmp_path: Path, monkeypatch
) -> None:
    settings_file = _patch_settings_file(monkeypatch, tmp_path)
    _write_legacy_blank_primary_settings(settings_file)
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    draft = build_interpretation_honorarios_draft(
        case_number="109/26.0PBBJA",
        case_entity="Juízo Local Criminal de Beja",
        case_city="Beja",
        service_date="2026-03-09",
        service_entity="",
        service_city="",
        use_service_location_in_honorarios=False,
        travel_km_outbound=0.0,
        travel_km_return=0.0,
        recipient_block=default_interpretation_recipient_block("Juízo Local Criminal de Beja"),
        profile=_profile(),
    )

    dialog = QtHonorariosExportDialog(parent=None, draft=draft, default_directory=tmp_path)
    try:
        dialog.show()
        app.processEvents()
        assert dialog.service_city_edit.text() == "Beja"
        assert dialog.travel_km_outbound_edit.text() == "39"
        profiles, primary_profile_id = user_settings.load_profile_settings()
        stored_profile = next(profile for profile in profiles if profile.id == primary_profile_id)
        assert stored_profile.travel_origin_label == "Marmelar"
        assert stored_profile.travel_distances_by_city["Beja"] == 39.0
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()


def test_honorarios_dialog_interpretation_profile_change_refreshes_known_distance(
    tmp_path: Path, monkeypatch
) -> None:
    _patch_settings_file(monkeypatch, tmp_path)
    primary = _profile(id="primary")
    primary.travel_distances_by_city = {"Beja": 39.0}
    secondary = _profile(
        id="secondary",
        first_name="Jane",
        last_name="Doe",
        document_name_override="Jane Doe",
        postal_address="Rua B",
        iban="PT50003506490000832760030",
        travel_origin_label="Marmelar",
        travel_distances_by_city={"Beja": 44.0},
    )
    user_settings.save_profile_settings(profiles=[primary, secondary], primary_profile_id=primary.id)
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    draft = build_interpretation_honorarios_draft(
        case_number="109/26.0PBBJA",
        case_entity="Juízo Local Criminal de Beja",
        case_city="Beja",
        service_date="2026-03-09",
        service_entity="",
        service_city="",
        use_service_location_in_honorarios=False,
        travel_km_outbound=0.0,
        travel_km_return=0.0,
        recipient_block=default_interpretation_recipient_block("Juízo Local Criminal de Beja"),
        profile=primary,
    )

    dialog = QtHonorariosExportDialog(parent=None, draft=draft, default_directory=tmp_path)
    try:
        dialog.show()
        app.processEvents()
        assert dialog.travel_km_outbound_edit.text() == "39"
        dialog.profile_combo.setCurrentIndex(dialog.profile_combo.findData("secondary"))
        app.processEvents()
        assert dialog.travel_km_outbound_edit.text() == "44"
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()


def test_honorarios_dialog_interpretation_service_city_change_updates_distance_and_persists_one_way(
    tmp_path: Path, monkeypatch
) -> None:
    _patch_settings_file(monkeypatch, tmp_path)
    profile = _profile()
    profile.travel_distances_by_city = {"Beja": 39.0, "Cuba": 26.0}
    user_settings.save_profile_settings(profiles=[profile], primary_profile_id=profile.id)
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    draft = build_interpretation_honorarios_draft(
        case_number="109/26.0PBBJA",
        case_entity="Juízo Local Criminal de Beja",
        case_city="Beja",
        service_date="2026-03-09",
        service_entity="Juízo Local Criminal de Beja",
        service_city="Beja",
        use_service_location_in_honorarios=False,
        travel_km_outbound=39.0,
        travel_km_return=39.0,
        recipient_block=default_interpretation_recipient_block("Juízo Local Criminal de Beja"),
        profile=profile,
    )

    dialog = QtHonorariosExportDialog(parent=None, draft=draft, default_directory=tmp_path)
    try:
        dialog.show()
        app.processEvents()
        dialog.service_same_check.setChecked(False)
        dialog.service_date_edit.setCalendarDate(QDate(2026, 3, 10))
        dialog.service_city_edit.setText("Cuba")
        app.processEvents()
        assert dialog.travel_km_outbound_edit.text() == "26"
        dialog.travel_km_outbound_edit.setText("27")
        rebuilt = dialog._build_draft()
        assert rebuilt.service_date == "2026-03-10"
        assert rebuilt.travel_km_outbound == 27.0
        assert rebuilt.travel_km_return == 27.0
        profiles, primary_profile_id = user_settings.load_profile_settings()
        stored_profile = next(profile for profile in profiles if profile.id == primary_profile_id)
        assert stored_profile.travel_distances_by_city["Cuba"] == 27.0
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()


def test_profile_manager_save_rejects_missing_required_fields(tmp_path: Path, monkeypatch) -> None:
    _patch_settings_file(monkeypatch, tmp_path)
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    saved: dict[str, object] = {}
    criticals: list[str] = []
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QMessageBox.critical",
        lambda *args, **kwargs: criticals.append(args[2] if len(args) > 2 else kwargs.get("text", "")),
    )

    dialog = QtProfileManagerDialog(
        parent=None,
        settings=user_settings.load_gui_settings(),
        save_callback=lambda profiles, primary_id: saved.update(
            profiles=profiles,
            primary_profile_id=primary_id,
        ),
    )
    try:
        dialog._new_profile()
        dialog.first_name_edit.setText("Jane")
        dialog.last_name_edit.setText("Doe")
        dialog.phone_edit.setText("+351912345678")
        dialog._save()
        assert "profiles" not in saved
        assert criticals
        assert "Postal address" in criticals[0]
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()


def test_profile_manager_can_set_primary_and_prevents_last_profile_delete(tmp_path: Path, monkeypatch) -> None:
    _patch_settings_file(monkeypatch, tmp_path)
    primary = _profile(id="primary")
    secondary = _profile(
        id="secondary",
        first_name="Jane",
        last_name="Doe",
        document_name_override="Jane Doe",
        email="jane@example.com",
        phone_number="+351911111111",
        postal_address="Rua B",
        iban="PT50003506490000832760030",
        iva_text="6%",
        irs_text="Isento",
    )
    user_settings.save_profile_settings(profiles=[primary, secondary], primary_profile_id=primary.id)

    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    saved: dict[str, object] = {}
    infos: list[str] = []
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QMessageBox.question",
        lambda *args, **kwargs: QMessageBox.StandardButton.Yes,
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QMessageBox.information",
        lambda *args, **kwargs: infos.append(args[2] if len(args) > 2 else kwargs.get("text", "")),
    )

    dialog = QtProfileManagerDialog(
        parent=None,
        settings=user_settings.load_gui_settings(),
        save_callback=lambda profiles, primary_id: saved.update(
            profiles=profiles,
            primary_profile_id=primary_id,
        ),
    )
    try:
        assert dialog.save_btn.objectName() == "PrimaryButton"
        assert dialog.delete_profile_btn.objectName() == "DangerButton"
        assert dialog.delete_distance_btn.objectName() == "DangerButton"
        dialog.profile_list.setCurrentRow(1)
        dialog._set_primary()
        dialog._save()
        assert saved["primary_profile_id"] == "secondary"
        assert saved["profiles"][1].phone_number == "+351911111111"

        last_dialog = QtProfileManagerDialog(
            parent=None,
            settings={"profiles": [user_settings.load_gui_settings()["profiles"][0]], "primary_profile_id": "primary"},
            save_callback=lambda profiles, primary_id: None,
        )
        try:
            last_dialog._delete_profile()
            assert infos
            assert "At least one profile must remain." in infos[-1]
        finally:
            last_dialog.close()
            last_dialog.deleteLater()
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()


def test_save_to_joblog_dialog_opens_honorarios_dialog_with_current_values(tmp_path: Path, monkeypatch) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    seed = JobLogSeed(
        completed_at=datetime.now().isoformat(timespec="seconds"),
        translation_date="2026-03-06",
        job_type="Translation",
        case_number="seed-case",
        court_email="",
        case_entity="Seed Entity",
        case_city="Seed City",
        service_entity="Seed Entity",
        service_city="Seed City",
        service_date="2026-03-06",
        lang="AR",
        pages=7,
        word_count=1666,
        rate_per_word=0.09,
        expected_total=149.94,
        amount_paid=0.0,
        api_cost=0.56,
        run_id="20260306_165834",
        target_lang="AR",
        total_tokens=57126,
        estimated_api_cost=0.56,
        quality_risk_score=0.1754,
        profit=149.38,
        pdf_path=tmp_path / "sample.pdf",
        output_docx=tmp_path / "out" / "translated.docx",
    )
    seed.output_docx.parent.mkdir(parents=True)
    seed.output_docx.write_bytes(b"docx")

    captured: dict[str, object] = {}

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            captured["parent"] = parent
            captured["draft"] = draft
            captured["default_directory"] = default_directory

        def exec(self) -> int:
            captured["exec"] = True
            return 0

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_downloads_dir", lambda: tmp_path / "downloads")
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_documents_dir", lambda: tmp_path / "documents")

    dialog = QtSaveToJobLogDialog(parent=None, db_path=tmp_path / "joblog.sqlite3", seed=seed)
    try:
        dialog.case_number_edit.setText("109/26.0PBBJA")
        dialog.case_entity_combo.addItem("Juízo Local Criminal de Beja")
        dialog.case_entity_combo.setCurrentText("Juízo Local Criminal de Beja")
        dialog.case_city_combo.addItem("Beja")
        dialog.case_city_combo.setCurrentText("Beja")
        dialog._open_honorarios_dialog()
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()

    draft = captured["draft"]
    assert isinstance(draft, HonorariosDraft)
    assert draft.case_number == "109/26.0PBBJA"
    assert draft.word_count == 1666
    assert draft.case_entity == "Juízo Local Criminal de Beja"
    assert draft.case_city == "Beja"
    assert captured["default_directory"] == seed.output_docx.parent.resolve()
    assert captured["exec"] is True


def test_joblog_window_generates_honorarios_from_selected_row(tmp_path: Path, monkeypatch) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    db_path = tmp_path / "joblog.sqlite3"
    with open_job_log(db_path) as conn:
        insert_job_run(
            conn,
            {
                "completed_at": "2026-03-06T16:58:34",
                "translation_date": "2026-03-06",
                "job_type": "Translation",
                "case_number": "109/26.0PBBJA",
                "court_email": "beja.judicial@tribunais.org.pt",
                "case_entity": "Juízo Local Criminal de Beja",
                "case_city": "Beja",
                "service_entity": "Juízo Local Criminal de Beja",
                "service_city": "Beja",
                "service_date": "2026-03-06",
                "lang": "AR",
                "target_lang": "AR",
                "run_id": "20260306_165834",
                "pages": 7,
                "word_count": 1666,
                "total_tokens": 57126,
                "rate_per_word": 0.09,
                "expected_total": 149.94,
                "amount_paid": 0.0,
                "api_cost": 0.56,
                "estimated_api_cost": 0.56,
                "quality_risk_score": 0.1754,
                "profit": 149.38,
            },
        )

    captured: dict[str, object] = {}

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            captured["draft"] = draft
            captured["default_directory"] = default_directory

        def exec(self) -> int:
            captured["exec"] = True
            return 0

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_downloads_dir", lambda: tmp_path / "downloads")
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_documents_dir", lambda: tmp_path / "documents")

    window = QtJobLogWindow(parent=None, db_path=db_path)
    try:
        assert window.honorarios_btn.isEnabled() is False
        window.table.selectRow(0)
        QApplication.processEvents()
        assert window.honorarios_btn.isEnabled() is True
        window._open_honorarios_dialog()
    finally:
        window.close()
        window.deleteLater()
        if owns_app:
            app.quit()


def test_save_to_joblog_dialog_offers_gmail_draft_for_current_run_export(
    tmp_path: Path, monkeypatch
) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    translated = tmp_path / "out" / "translated.docx"
    translated.parent.mkdir(parents=True)
    translated.write_bytes(b"docx")
    honorarios = tmp_path / "honorarios.docx"
    honorarios.write_bytes(b"docx")
    seed = JobLogSeed(
        completed_at=datetime.now().isoformat(timespec="seconds"),
        translation_date="2026-03-06",
        job_type="Translation",
        case_number="seed-case",
        court_email="beja.judicial@tribunais.org.pt",
        case_entity="Seed Entity",
        case_city="Seed City",
        service_entity="Seed Entity",
        service_city="Seed City",
        service_date="2026-03-06",
        lang="AR",
        pages=7,
        word_count=1666,
        rate_per_word=0.09,
        expected_total=149.94,
        amount_paid=0.0,
        api_cost=0.56,
        run_id="20260306_165834",
        target_lang="AR",
        total_tokens=57126,
        estimated_api_cost=0.56,
        quality_risk_score=0.1754,
        profit=149.38,
        pdf_path=tmp_path / "sample.pdf",
        output_docx=translated,
    )

    captured: dict[str, object] = {}

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            self.saved_path = honorarios

        def exec(self) -> int:
            return 1

    class _FakeRequest:
        def __init__(self) -> None:
            self.account_email = "adel.belghali@gmail.com"
            self.to_email = "beja.judicial@tribunais.org.pt"
            self.subject = "subject"
            self.body = "body"
            self.attachments = (translated, honorarios)
            self.gog_path = Path(r"C:\gog.exe")

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_downloads_dir", lambda: tmp_path / "downloads")
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_documents_dir", lambda: tmp_path / "documents")
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: type(
            "ReadyStatus",
            (),
            {
                "ready": True,
                "message": "ready",
                "gog_path": Path(r"C:\gog.exe"),
                "account_email": "adel.belghali@gmail.com",
                "accounts": ("adel.belghali@gmail.com",),
            },
        )(),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.validate_translated_docx_artifacts_for_gmail_draft",
        lambda **kwargs: tuple(kwargs["translated_docxs"]),
    )

    def _fake_build_request(**kwargs):
        captured["request_kwargs"] = kwargs
        return _FakeRequest()

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.build_honorarios_gmail_request", _fake_build_request)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.create_gmail_draft_via_gog",
        lambda request: type(
            "DraftResult",
            (),
            {"ok": True, "message": "ok", "stdout": "", "stderr": "", "payload": {"id": "1"}},
        )(),
    )
    answers = iter([16384, 16384])
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.question", lambda *args, **kwargs: next(answers))
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.critical", lambda *args, **kwargs: None)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QDesktopServices.openUrl", lambda url: captured.setdefault("url", url))

    dialog = QtSaveToJobLogDialog(parent=None, db_path=tmp_path / "joblog.sqlite3", seed=seed)
    try:
        dialog.case_number_edit.setText("109/26.0PBBJA")
        dialog.case_entity_combo.setCurrentText("Juízo Local Criminal de Beja")
        dialog.case_city_combo.setCurrentText("Beja")
        dialog._open_honorarios_dialog()
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()

    assert captured["request_kwargs"]["to_email"] == "beja.judicial@tribunais.org.pt"
    assert captured["request_kwargs"]["case_number"] == "109/26.0PBBJA"
    assert captured["request_kwargs"]["translation_docx"] == translated
    assert captured["request_kwargs"]["honorarios_docx"] == honorarios
    assert captured["request_kwargs"]["profile"].document_name == "Adel Belghali"
    opened = captured["url"]
    assert isinstance(opened, QUrl)
    assert opened.toString() == "https://mail.google.com/mail/u/0/#drafts"


def test_save_to_joblog_dialog_uses_partial_docx_for_gmail_when_final_output_missing(
    tmp_path: Path, monkeypatch
) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    partial = tmp_path / "out" / "partial.docx"
    partial.parent.mkdir(parents=True)
    partial.write_bytes(b"docx")
    honorarios = tmp_path / "honorarios.docx"
    honorarios.write_bytes(b"docx")
    seed = JobLogSeed(
        completed_at=datetime.now().isoformat(timespec="seconds"),
        translation_date="2026-03-06",
        job_type="Translation",
        case_number="seed-case",
        court_email="beja.judicial@tribunais.org.pt",
        case_entity="Seed Entity",
        case_city="Seed City",
        service_entity="Seed Entity",
        service_city="Seed City",
        service_date="2026-03-06",
        lang="AR",
        pages=7,
        word_count=1666,
        rate_per_word=0.09,
        expected_total=149.94,
        amount_paid=0.0,
        api_cost=0.56,
        run_id="20260306_165834",
        target_lang="AR",
        total_tokens=57126,
        estimated_api_cost=0.56,
        quality_risk_score=0.1754,
        profit=149.38,
        pdf_path=tmp_path / "sample.pdf",
        output_docx=tmp_path / "out" / "missing_final.docx",
        partial_docx=partial,
    )

    captured: dict[str, object] = {}

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            self.saved_path = honorarios

        def exec(self) -> int:
            return 1

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_downloads_dir", lambda: tmp_path / "downloads")
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_documents_dir", lambda: tmp_path / "documents")
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: type(
            "ReadyStatus",
            (),
            {
                "ready": True,
                "message": "ready",
                "gog_path": Path(r"C:\gog.exe"),
                "account_email": "adel.belghali@gmail.com",
                "accounts": ("adel.belghali@gmail.com",),
            },
        )(),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.validate_translated_docx_artifacts_for_gmail_draft",
        lambda **kwargs: tuple(kwargs["translated_docxs"]),
    )
    def _fake_build_request(**kwargs):
        captured["request_kwargs"] = kwargs
        return type(
            "Request",
            (),
            {
                "account_email": "adel.belghali@gmail.com",
                "to_email": kwargs["to_email"],
                "subject": "subject",
                "body": "body",
                "attachments": (kwargs["translation_docx"], kwargs["honorarios_docx"]),
                "gog_path": kwargs["gog_path"],
            },
        )()

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.build_honorarios_gmail_request", _fake_build_request)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.create_gmail_draft_via_gog",
        lambda request: type(
            "DraftResult",
            (),
            {"ok": True, "message": "ok", "stdout": "", "stderr": "", "payload": {"id": "1"}},
        )(),
    )
    answers = iter([16384, 65536])
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.question", lambda *args, **kwargs: next(answers))
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.critical", lambda *args, **kwargs: None)

    dialog = QtSaveToJobLogDialog(parent=None, db_path=tmp_path / "joblog.sqlite3", seed=seed)
    try:
        dialog.case_number_edit.setText("109/26.0PBBJA")
        dialog.case_entity_combo.setCurrentText("Juízo Local Criminal de Beja")
        dialog.case_city_combo.setCurrentText("Beja")
        dialog._open_honorarios_dialog()
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()

    assert captured["request_kwargs"]["translation_docx"] == partial
    assert captured["request_kwargs"]["honorarios_docx"] == honorarios


def test_save_to_joblog_dialog_skips_gmail_when_court_email_missing(
    tmp_path: Path, monkeypatch
) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    translated = tmp_path / "out" / "translated.docx"
    translated.parent.mkdir(parents=True)
    translated.write_bytes(b"docx")
    honorarios = tmp_path / "honorarios.docx"
    honorarios.write_bytes(b"docx")
    seed = JobLogSeed(
        completed_at=datetime.now().isoformat(timespec="seconds"),
        translation_date="2026-03-06",
        job_type="Translation",
        case_number="seed-case",
        court_email="",
        case_entity="Seed Entity",
        case_city="Seed City",
        service_entity="Seed Entity",
        service_city="Seed City",
        service_date="2026-03-06",
        lang="AR",
        pages=7,
        word_count=1666,
        rate_per_word=0.09,
        expected_total=149.94,
        amount_paid=0.0,
        api_cost=0.56,
        run_id="20260306_165834",
        target_lang="AR",
        total_tokens=57126,
        estimated_api_cost=0.56,
        quality_risk_score=0.1754,
        profit=149.38,
        pdf_path=tmp_path / "sample.pdf",
        output_docx=translated,
    )

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            self.saved_path = honorarios

        def exec(self) -> int:
            return 1

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: (_ for _ in ()).throw(AssertionError("should not assess Gmail prereqs")),
    )

    dialog = QtSaveToJobLogDialog(parent=None, db_path=tmp_path / "joblog.sqlite3", seed=seed)
    try:
        dialog._open_honorarios_dialog()
    finally:
        dialog.close()
        dialog.deleteLater()
        if owns_app:
            app.quit()


def test_joblog_window_honorarios_export_offers_gmail_draft_for_selected_row(
    tmp_path: Path, monkeypatch
) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    db_path = tmp_path / "joblog.sqlite3"
    translated = tmp_path / "translated.docx"
    translated.write_bytes(b"docx")
    with open_job_log(db_path) as conn:
        insert_job_run(
            conn,
            {
                "completed_at": "2026-03-06T16:58:34",
                "translation_date": "2026-03-06",
                "job_type": "Translation",
                "case_number": "109/26.0PBBJA",
                "court_email": "beja.judicial@tribunais.org.pt",
                "case_entity": "Juízo Local Criminal de Beja",
                "case_city": "Beja",
                "service_entity": "Juízo Local Criminal de Beja",
                "service_city": "Beja",
                "service_date": "2026-03-06",
                "lang": "AR",
                "target_lang": "AR",
                "run_id": "20260306_165834",
                "pages": 7,
                "word_count": 1666,
                "total_tokens": 57126,
                "rate_per_word": 0.09,
                "expected_total": 149.94,
                "amount_paid": 0.0,
                "api_cost": 0.56,
                "estimated_api_cost": 0.56,
                "quality_risk_score": 0.1754,
                "profit": 149.38,
                "output_docx_path": str(translated),
            },
            )

    captured: dict[str, object] = {}

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            captured["draft"] = draft
            captured["default_directory"] = default_directory
            self.saved_path = tmp_path / "historical.docx"

        def exec(self) -> int:
            captured["exec"] = True
            return 1

    class _FakeRequest:
        def __init__(self) -> None:
            self.account_email = "adel.belghali@gmail.com"
            self.to_email = "beja.judicial@tribunais.org.pt"
            self.subject = "subject"
            self.body = "body"
            self.attachments = (translated, tmp_path / "historical.docx")
            self.gog_path = Path(r"C:\gog.exe")

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: type(
            "ReadyStatus",
            (),
            {
                "ready": True,
                "message": "ready",
                "gog_path": Path(r"C:\gog.exe"),
                "account_email": "adel.belghali@gmail.com",
                "accounts": ("adel.belghali@gmail.com",),
            },
        )(),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.validate_translated_docx_artifacts_for_gmail_draft",
        lambda **kwargs: tuple(kwargs["translated_docxs"]),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QFileDialog.getOpenFileName",
        lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("picker should not open when output_docx_path is stored")),
    )

    def _fake_build_request(**kwargs):
        captured["request_kwargs"] = kwargs
        return _FakeRequest()

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.build_honorarios_gmail_request", _fake_build_request)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.create_gmail_draft_via_gog",
        lambda request: type(
            "DraftResult",
            (),
            {"ok": True, "message": "ok", "stdout": "", "stderr": "", "payload": {"id": "1"}},
        )(),
    )
    answers = iter([16384, 16384])
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.question", lambda *args, **kwargs: next(answers))
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.critical", lambda *args, **kwargs: None)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.warning", lambda *args, **kwargs: None)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QDesktopServices.openUrl",
        lambda url: captured.setdefault("url", url),
    )

    window = QtJobLogWindow(parent=None, db_path=db_path)
    try:
        window.table.selectRow(0)
        QApplication.processEvents()
        window._open_honorarios_dialog()
    finally:
        window.close()
        window.deleteLater()
        if owns_app:
            app.quit()

    draft = captured["draft"]
    assert isinstance(draft, HonorariosDraft)
    assert draft.case_number == "109/26.0PBBJA"
    assert captured["request_kwargs"]["to_email"] == "beja.judicial@tribunais.org.pt"
    assert captured["request_kwargs"]["case_number"] == "109/26.0PBBJA"
    assert captured["request_kwargs"]["translation_docx"] == translated
    assert captured["request_kwargs"]["honorarios_docx"] == tmp_path / "historical.docx"
    assert captured["request_kwargs"]["profile"].document_name == "Adel Belghali"
    opened = captured["url"]
    assert isinstance(opened, QUrl)
    assert opened.toString() == "https://mail.google.com/mail/u/0/#drafts"


def test_joblog_window_honorarios_export_falls_back_to_stored_partial_docx(
    tmp_path: Path, monkeypatch
) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    db_path = tmp_path / "joblog.sqlite3"
    partial = tmp_path / "partial.docx"
    partial.write_bytes(b"docx")
    with open_job_log(db_path) as conn:
        insert_job_run(
            conn,
            {
                "completed_at": "2026-03-06T16:58:34",
                "translation_date": "2026-03-06",
                "job_type": "Translation",
                "case_number": "109/26.0PBBJA",
                "court_email": "beja.judicial@tribunais.org.pt",
                "case_entity": "Juízo Local Criminal de Beja",
                "case_city": "Beja",
                "service_entity": "Juízo Local Criminal de Beja",
                "service_city": "Beja",
                "service_date": "2026-03-06",
                "lang": "AR",
                "target_lang": "AR",
                "run_id": "20260306_165834",
                "pages": 7,
                "word_count": 1666,
                "total_tokens": 57126,
                "rate_per_word": 0.09,
                "expected_total": 149.94,
                "amount_paid": 0.0,
                "api_cost": 0.56,
                "estimated_api_cost": 0.56,
                "quality_risk_score": 0.1754,
                "profit": 149.38,
                "partial_docx_path": str(partial),
            },
        )

    captured: dict[str, object] = {}

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            self.saved_path = tmp_path / "historical.docx"

        def exec(self) -> int:
            return 1

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: type(
            "ReadyStatus",
            (),
            {
                "ready": True,
                "message": "ready",
                "gog_path": Path(r"C:\gog.exe"),
                "account_email": "adel.belghali@gmail.com",
                "accounts": ("adel.belghali@gmail.com",),
            },
        )(),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.validate_translated_docx_artifacts_for_gmail_draft",
        lambda **kwargs: tuple(kwargs["translated_docxs"]),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QFileDialog.getOpenFileName",
        lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("picker should not open when partial_docx_path is stored")),
    )
    def _fake_build_request(**kwargs):
        captured["request_kwargs"] = kwargs
        return type(
            "Request",
            (),
            {
                "account_email": "adel.belghali@gmail.com",
                "to_email": kwargs["to_email"],
                "subject": "subject",
                "body": "body",
                "attachments": (kwargs["translation_docx"], kwargs["honorarios_docx"]),
                "gog_path": kwargs["gog_path"],
            },
        )()

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.build_honorarios_gmail_request", _fake_build_request)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.create_gmail_draft_via_gog",
        lambda request: type(
            "DraftResult",
            (),
            {"ok": True, "message": "ok", "stdout": "", "stderr": "", "payload": {"id": "1"}},
        )(),
    )
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.question", lambda *args, **kwargs: 16384)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.critical", lambda *args, **kwargs: None)

    window = QtJobLogWindow(parent=None, db_path=db_path)
    try:
        window.table.selectRow(0)
        QApplication.processEvents()
        window._open_honorarios_dialog()
    finally:
        window.close()
        window.deleteLater()
        if owns_app:
            app.quit()

    assert captured["request_kwargs"]["translation_docx"] == partial


def test_joblog_window_honorarios_export_informs_when_court_email_missing(tmp_path: Path, monkeypatch) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    db_path = tmp_path / "joblog.sqlite3"
    with open_job_log(db_path) as conn:
        insert_job_run(
            conn,
            {
                "completed_at": "2026-03-06T16:58:34",
                "translation_date": "2026-03-06",
                "job_type": "Translation",
                "case_number": "109/26.0PBBJA",
                "court_email": "",
                "case_entity": "Juízo Local Criminal de Beja",
                "case_city": "Beja",
                "service_entity": "Juízo Local Criminal de Beja",
                "service_city": "Beja",
                "service_date": "2026-03-06",
                "lang": "AR",
                "target_lang": "AR",
                "run_id": "20260306_165834",
                "pages": 7,
                "word_count": 1666,
                "total_tokens": 57126,
                "rate_per_word": 0.09,
                "expected_total": 149.94,
                "amount_paid": 0.0,
                "api_cost": 0.56,
                "estimated_api_cost": 0.56,
                "quality_risk_score": 0.1754,
                "profit": 149.38,
            },
        )

    captured: dict[str, object] = {}

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            self.saved_path = tmp_path / "historical.docx"

        def exec(self) -> int:
            return 1

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: (_ for _ in ()).throw(AssertionError("court-email missing should stop before Gmail prereqs")),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QMessageBox.information",
        lambda *args, **kwargs: captured.setdefault("info", args[2] if len(args) > 2 else kwargs.get("text")),
    )

    window = QtJobLogWindow(parent=None, db_path=db_path)
    try:
        window.table.selectRow(0)
        QApplication.processEvents()
        window._open_honorarios_dialog()
    finally:
        window.close()
        window.deleteLater()
        if owns_app:
            app.quit()

    assert captured["info"] == "Court Email is missing for this Job Log entry. The Gmail draft was not created."


def test_joblog_window_honorarios_export_stops_when_translation_docx_picker_cancelled(
    tmp_path: Path, monkeypatch
) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    db_path = tmp_path / "joblog.sqlite3"
    with open_job_log(db_path) as conn:
        insert_job_run(
            conn,
            {
                "completed_at": "2026-03-06T16:58:34",
                "translation_date": "2026-03-06",
                "job_type": "Translation",
                "case_number": "109/26.0PBBJA",
                "court_email": "beja.judicial@tribunais.org.pt",
                "case_entity": "Juízo Local Criminal de Beja",
                "case_city": "Beja",
                "service_entity": "Juízo Local Criminal de Beja",
                "service_city": "Beja",
                "service_date": "2026-03-06",
                "lang": "AR",
                "target_lang": "AR",
                "run_id": "20260306_165834",
                "pages": 7,
                "word_count": 1666,
                "total_tokens": 57126,
                "rate_per_word": 0.09,
                "expected_total": 149.94,
                "amount_paid": 0.0,
                "api_cost": 0.56,
                "estimated_api_cost": 0.56,
                "quality_risk_score": 0.1754,
                "profit": 149.38,
            },
        )

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            self.saved_path = tmp_path / "historical.docx"

        def exec(self) -> int:
            return 1

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_downloads_dir", lambda: tmp_path / "downloads")
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_documents_dir", lambda: tmp_path / "documents")
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: type(
            "ReadyStatus",
            (),
            {
                "ready": True,
                "message": "ready",
                "gog_path": Path(r"C:\gog.exe"),
                "account_email": "adel.belghali@gmail.com",
                "accounts": ("adel.belghali@gmail.com",),
            },
        )(),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.validate_translated_docx_artifacts_for_gmail_draft",
        lambda **kwargs: tuple(kwargs["translated_docxs"]),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QFileDialog.getOpenFileName",
        lambda *args, **kwargs: ("", ""),
    )
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.question", lambda *args, **kwargs: 16384)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.build_honorarios_gmail_request",
        lambda **kwargs: (_ for _ in ()).throw(AssertionError("request builder should not run when picker is cancelled")),
    )

    window = QtJobLogWindow(parent=None, db_path=db_path)
    try:
        window.table.selectRow(0)
        QApplication.processEvents()
        window._open_honorarios_dialog()
    finally:
        window.close()
        window.deleteLater()
        if owns_app:
            app.quit()


def test_joblog_window_honorarios_export_persists_selected_translation_docx_for_legacy_row(
    tmp_path: Path, monkeypatch
) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    db_path = tmp_path / "joblog.sqlite3"
    with open_job_log(db_path) as conn:
        insert_job_run(
            conn,
            {
                "completed_at": "2026-03-06T16:58:34",
                "translation_date": "2026-03-06",
                "job_type": "Translation",
                "case_number": "109/26.0PBBJA",
                "court_email": "beja.judicial@tribunais.org.pt",
                "case_entity": "Juízo Local Criminal de Beja",
                "case_city": "Beja",
                "service_entity": "Juízo Local Criminal de Beja",
                "service_city": "Beja",
                "service_date": "2026-03-06",
                "lang": "AR",
                "target_lang": "AR",
                "run_id": "20260306_165834",
                "pages": 7,
                "word_count": 1666,
                "total_tokens": 57126,
                "rate_per_word": 0.09,
                "expected_total": 149.94,
                "amount_paid": 0.0,
                "api_cost": 0.56,
                "estimated_api_cost": 0.56,
                "quality_risk_score": 0.1754,
                "profit": 149.38,
            },
        )

    translated = tmp_path / "legacy-selected.docx"
    translated.write_bytes(b"docx")

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            self.saved_path = tmp_path / "historical.docx"

        def exec(self) -> int:
            return 1

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_downloads_dir", lambda: tmp_path / "downloads")
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_documents_dir", lambda: tmp_path / "documents")
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: type(
            "ReadyStatus",
            (),
            {
                "ready": True,
                "message": "ready",
                "gog_path": Path(r"C:\gog.exe"),
                "account_email": "adel.belghali@gmail.com",
                "accounts": ("adel.belghali@gmail.com",),
            },
        )(),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.validate_translated_docx_artifacts_for_gmail_draft",
        lambda **kwargs: tuple(kwargs["translated_docxs"]),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QFileDialog.getOpenFileName",
        lambda *args, **kwargs: (str(translated), "Word Document (*.docx)"),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.build_honorarios_gmail_request",
        lambda **kwargs: type(
            "Request",
            (),
            {
                "account_email": "adel.belghali@gmail.com",
                "to_email": kwargs["to_email"],
                "subject": "subject",
                "body": "body",
                "attachments": (kwargs["translation_docx"], kwargs["honorarios_docx"]),
                "gog_path": kwargs["gog_path"],
            },
        )(),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.create_gmail_draft_via_gog",
        lambda request: type(
            "DraftResult",
            (),
            {"ok": True, "message": "ok", "stdout": "", "stderr": "", "payload": {"id": "1"}},
        )(),
    )
    answers = iter([16384, 65536])
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.question", lambda *args, **kwargs: next(answers))
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.critical", lambda *args, **kwargs: None)

    window = QtJobLogWindow(parent=None, db_path=db_path)
    try:
        window.table.selectRow(0)
        QApplication.processEvents()
        window._open_honorarios_dialog()
    finally:
        window.close()
        window.deleteLater()
        if owns_app:
            app.quit()

    with open_job_log(db_path) as conn:
        row = conn.execute(
            "SELECT output_docx_path, partial_docx_path FROM job_runs WHERE case_number = ?",
            ("109/26.0PBBJA",),
        ).fetchone()

    assert row is not None
    assert row["output_docx_path"] == str(translated.resolve())
    assert row["partial_docx_path"] is None


def test_joblog_window_honorarios_export_recovers_exact_run_id_translation_docx(
    tmp_path: Path, monkeypatch
) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    db_path = tmp_path / "joblog.sqlite3"
    run_id = "20260307_005451"
    with open_job_log(db_path) as conn:
        insert_job_run(
            conn,
            {
                "completed_at": "2026-03-07T00:54:51",
                "translation_date": "2026-03-07",
                "job_type": "Translation",
                "case_number": "109/26.0PBBJA",
                "court_email": "beja.judicial@tribunais.org.pt",
                "case_entity": "Tribunal do Trabalho",
                "case_city": "Beja",
                "service_entity": "Tribunal do Trabalho",
                "service_city": "Beja",
                "service_date": "2026-03-07",
                "lang": "EN",
                "target_lang": "EN",
                "run_id": run_id,
                "pages": 1,
                "word_count": 230,
                "total_tokens": 1000,
                "rate_per_word": 0.08,
                "expected_total": 18.4,
                "amount_paid": 0.0,
                "api_cost": 0.5,
                "estimated_api_cost": 0.5,
                "quality_risk_score": 0.1,
                "profit": 17.9,
            },
        )

    downloads_dir = tmp_path / "downloads"
    downloads_dir.mkdir()
    translated = downloads_dir / f"1_PDFsam_auto (1)_EN_{run_id}.docx"
    translated.write_bytes(b"docx")

    captured: dict[str, object] = {}

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            self.saved_path = downloads_dir / "Requerimento_Honorarios_109_26.0PBBJA_20260307.docx"
            self.saved_path.write_bytes(b"docx")

        def exec(self) -> int:
            return 1

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_downloads_dir", lambda: downloads_dir)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_documents_dir", lambda: tmp_path / "documents")
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: type(
            "ReadyStatus",
            (),
            {
                "ready": True,
                "message": "ready",
                "gog_path": Path(r"C:\gog.exe"),
                "account_email": "adel.belghali@gmail.com",
                "accounts": ("adel.belghali@gmail.com",),
            },
        )(),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.validate_translated_docx_artifacts_for_gmail_draft",
        lambda **kwargs: tuple(kwargs["translated_docxs"]),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QFileDialog.getOpenFileName",
        lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("picker should not open when exact run_id recovery succeeds")),
    )
    def _build_request(**kwargs):
        captured["request_kwargs"] = kwargs
        return type(
            "Request",
            (),
            {
                "account_email": "adel.belghali@gmail.com",
                "to_email": kwargs["to_email"],
                "subject": "subject",
                "body": "body",
                "attachments": (kwargs["translation_docx"], kwargs["honorarios_docx"]),
                "gog_path": kwargs["gog_path"],
            },
        )()

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.build_honorarios_gmail_request", _build_request)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.create_gmail_draft_via_gog",
        lambda request: type(
            "DraftResult",
            (),
            {"ok": True, "message": "ok", "stdout": "", "stderr": "", "payload": {"id": "1"}},
        )(),
    )
    answers = iter([16384, 65536])
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.question", lambda *args, **kwargs: next(answers))
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.critical", lambda *args, **kwargs: None)

    window = QtJobLogWindow(parent=None, db_path=db_path)
    try:
        window.table.selectRow(0)
        QApplication.processEvents()
        window._open_honorarios_dialog()
    finally:
        window.close()
        window.deleteLater()
        if owns_app:
            app.quit()

    assert captured["request_kwargs"]["translation_docx"] == translated.resolve()
    with open_job_log(db_path) as conn:
        row = conn.execute(
            "SELECT output_docx_path, partial_docx_path FROM job_runs WHERE case_number = ?",
            ("109/26.0PBBJA",),
        ).fetchone()
    assert row is not None
    assert row["output_docx_path"] == str(translated.resolve())
    assert row["partial_docx_path"] is None


def test_joblog_window_honorarios_export_recovers_exact_run_id_partial_docx(
    tmp_path: Path, monkeypatch
) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    db_path = tmp_path / "joblog.sqlite3"
    run_id = "20260307_005451"
    with open_job_log(db_path) as conn:
        insert_job_run(
            conn,
            {
                "completed_at": "2026-03-07T00:54:51",
                "translation_date": "2026-03-07",
                "job_type": "Translation",
                "case_number": "109/26.0PBBJA",
                "court_email": "beja.judicial@tribunais.org.pt",
                "case_entity": "Tribunal do Trabalho",
                "case_city": "Beja",
                "service_entity": "Tribunal do Trabalho",
                "service_city": "Beja",
                "service_date": "2026-03-07",
                "lang": "EN",
                "target_lang": "EN",
                "run_id": run_id,
                "pages": 1,
                "word_count": 230,
                "total_tokens": 1000,
                "rate_per_word": 0.08,
                "expected_total": 18.4,
                "amount_paid": 0.0,
                "api_cost": 0.5,
                "estimated_api_cost": 0.5,
                "quality_risk_score": 0.1,
                "profit": 17.9,
            },
        )

    downloads_dir = tmp_path / "downloads"
    downloads_dir.mkdir()
    partial = downloads_dir / f"1_PDFsam_auto (1)_EN_{run_id}_PARTIAL.docx"
    partial.write_bytes(b"docx")

    captured: dict[str, object] = {}

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            self.saved_path = downloads_dir / "Requerimento_Honorarios_109_26.0PBBJA_20260307.docx"
            self.saved_path.write_bytes(b"docx")

        def exec(self) -> int:
            return 1

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_downloads_dir", lambda: downloads_dir)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_documents_dir", lambda: tmp_path / "documents")
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: type(
            "ReadyStatus",
            (),
            {
                "ready": True,
                "message": "ready",
                "gog_path": Path(r"C:\gog.exe"),
                "account_email": "adel.belghali@gmail.com",
                "accounts": ("adel.belghali@gmail.com",),
            },
        )(),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.validate_translated_docx_artifacts_for_gmail_draft",
        lambda **kwargs: tuple(kwargs["translated_docxs"]),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QFileDialog.getOpenFileName",
        lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("picker should not open when exact partial run_id recovery succeeds")),
    )
    def _build_request(**kwargs):
        captured["request_kwargs"] = kwargs
        return type(
            "Request",
            (),
            {
                "account_email": "adel.belghali@gmail.com",
                "to_email": kwargs["to_email"],
                "subject": "subject",
                "body": "body",
                "attachments": (kwargs["translation_docx"], kwargs["honorarios_docx"]),
                "gog_path": kwargs["gog_path"],
            },
        )()

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.build_honorarios_gmail_request", _build_request)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.create_gmail_draft_via_gog",
        lambda request: type(
            "DraftResult",
            (),
            {"ok": True, "message": "ok", "stdout": "", "stderr": "", "payload": {"id": "1"}},
        )(),
    )
    answers = iter([16384, 65536])
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.question", lambda *args, **kwargs: next(answers))
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.critical", lambda *args, **kwargs: None)

    window = QtJobLogWindow(parent=None, db_path=db_path)
    try:
        window.table.selectRow(0)
        QApplication.processEvents()
        window._open_honorarios_dialog()
    finally:
        window.close()
        window.deleteLater()
        if owns_app:
            app.quit()

    assert captured["request_kwargs"]["translation_docx"] == partial.resolve()
    with open_job_log(db_path) as conn:
        row = conn.execute(
            "SELECT output_docx_path, partial_docx_path FROM job_runs WHERE case_number = ?",
            ("109/26.0PBBJA",),
        ).fetchone()
    assert row is not None
    assert row["output_docx_path"] is None
    assert row["partial_docx_path"] == str(partial.resolve())


def test_joblog_window_honorarios_export_falls_back_to_picker_when_run_id_matches_are_ambiguous(
    tmp_path: Path, monkeypatch
) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    db_path = tmp_path / "joblog.sqlite3"
    run_id = "20260307_005451"
    with open_job_log(db_path) as conn:
        insert_job_run(
            conn,
            {
                "completed_at": "2026-03-07T00:54:51",
                "translation_date": "2026-03-07",
                "job_type": "Translation",
                "case_number": "109/26.0PBBJA",
                "court_email": "beja.judicial@tribunais.org.pt",
                "case_entity": "Tribunal do Trabalho",
                "case_city": "Beja",
                "service_entity": "Tribunal do Trabalho",
                "service_city": "Beja",
                "service_date": "2026-03-07",
                "lang": "EN",
                "target_lang": "EN",
                "run_id": run_id,
                "pages": 1,
                "word_count": 230,
                "total_tokens": 1000,
                "rate_per_word": 0.08,
                "expected_total": 18.4,
                "amount_paid": 0.0,
                "api_cost": 0.5,
                "estimated_api_cost": 0.5,
                "quality_risk_score": 0.1,
                "profit": 17.9,
            },
        )

    downloads_dir = tmp_path / "downloads"
    downloads_dir.mkdir()
    first = downloads_dir / f"1_PDFsam_auto (1)_EN_{run_id}.docx"
    second = downloads_dir / f"another_EN_{run_id}.docx"
    first.write_bytes(b"docx")
    second.write_bytes(b"docx")
    picked = downloads_dir / "manual-selection.docx"
    picked.write_bytes(b"docx")

    captured: dict[str, object] = {}

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            self.saved_path = downloads_dir / "Requerimento_Honorarios_109_26.0PBBJA_20260307.docx"
            self.saved_path.write_bytes(b"docx")

        def exec(self) -> int:
            return 1

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_downloads_dir", lambda: downloads_dir)
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs._default_documents_dir", lambda: tmp_path / "documents")
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: type(
            "ReadyStatus",
            (),
            {
                "ready": True,
                "message": "ready",
                "gog_path": Path(r"C:\gog.exe"),
                "account_email": "adel.belghali@gmail.com",
                "accounts": ("adel.belghali@gmail.com",),
            },
        )(),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.validate_translated_docx_artifacts_for_gmail_draft",
        lambda **kwargs: tuple(kwargs["translated_docxs"]),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QFileDialog.getOpenFileName",
        lambda *args, **kwargs: (str(picked), "Word Document (*.docx)"),
    )
    def _build_request(**kwargs):
        captured["request_kwargs"] = kwargs
        return type(
            "Request",
            (),
            {
                "account_email": "adel.belghali@gmail.com",
                "to_email": kwargs["to_email"],
                "subject": "subject",
                "body": "body",
                "attachments": (kwargs["translation_docx"], kwargs["honorarios_docx"]),
                "gog_path": kwargs["gog_path"],
            },
        )()

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.build_honorarios_gmail_request", _build_request)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.create_gmail_draft_via_gog",
        lambda request: type(
            "DraftResult",
            (),
            {"ok": True, "message": "ok", "stdout": "", "stderr": "", "payload": {"id": "1"}},
        )(),
    )
    answers = iter([16384, 65536])
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.question", lambda *args, **kwargs: next(answers))
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.critical", lambda *args, **kwargs: None)

    window = QtJobLogWindow(parent=None, db_path=db_path)
    try:
        window.table.selectRow(0)
        QApplication.processEvents()
        window._open_honorarios_dialog()
    finally:
        window.close()
        window.deleteLater()
        if owns_app:
            app.quit()

    assert captured["request_kwargs"]["translation_docx"] == picked.resolve()


def test_joblog_window_honorarios_export_warns_when_gmail_not_ready(tmp_path: Path, monkeypatch) -> None:
    app = QApplication.instance()
    owns_app = app is None
    if app is None:
        app = QApplication(sys.argv[:1])

    db_path = tmp_path / "joblog.sqlite3"
    with open_job_log(db_path) as conn:
        insert_job_run(
            conn,
            {
                "completed_at": "2026-03-06T16:58:34",
                "translation_date": "2026-03-06",
                "job_type": "Translation",
                "case_number": "109/26.0PBBJA",
                "court_email": "beja.judicial@tribunais.org.pt",
                "case_entity": "Juízo Local Criminal de Beja",
                "case_city": "Beja",
                "service_entity": "Juízo Local Criminal de Beja",
                "service_city": "Beja",
                "service_date": "2026-03-06",
                "lang": "AR",
                "target_lang": "AR",
                "run_id": "20260306_165834",
                "pages": 7,
                "word_count": 1666,
                "total_tokens": 57126,
                "rate_per_word": 0.09,
                "expected_total": 149.94,
                "amount_paid": 0.0,
                "api_cost": 0.56,
                "estimated_api_cost": 0.56,
                "quality_risk_score": 0.1754,
                "profit": 149.38,
            },
        )

    captured: dict[str, object] = {}

    class _FakeHonorariosDialog:
        def __init__(self, *, parent, draft, default_directory) -> None:
            self.saved_path = tmp_path / "historical.docx"

        def exec(self) -> int:
            return 1

    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QtHonorariosExportDialog", _FakeHonorariosDialog)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: type(
            "NotReadyStatus",
            (),
            {
                "ready": False,
                "message": "No Gmail account is authenticated in gog.",
                "gog_path": Path(r"C:\gog.exe"),
                "account_email": None,
                "accounts": (),
            },
        )(),
    )
    monkeypatch.setattr("legalpdf_translate.qt_gui.dialogs.QMessageBox.question", lambda *args, **kwargs: 16384)
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QMessageBox.warning",
        lambda *args, **kwargs: captured.setdefault("warning", args[2] if len(args) > 2 else kwargs.get("text")),
    )

    window = QtJobLogWindow(parent=None, db_path=db_path)
    try:
        window.table.selectRow(0)
        QApplication.processEvents()
        window._open_honorarios_dialog()
    finally:
        window.close()
        window.deleteLater()
        if owns_app:
            app.quit()

    assert "No Gmail account is authenticated in gog." in captured["warning"]


def test_save_to_joblog_gmail_draft_blocks_contaminated_translation_docx(tmp_path: Path, monkeypatch) -> None:
    translated = tmp_path / "translated.docx"
    honorarios = tmp_path / "honorarios.docx"
    _write_docx_with_paragraphs(
        translated,
        "Venho por este meio requerer o pagamento dos honorários devidos.",
        "O documento traduzido contém 264 palavras.",
        "O Pagamento deverá ser efetuado para o seguinte IBAN: PT50003506490000832760029",
    )
    _write_docx_with_paragraphs(honorarios, "Requerimento de honorários distinto.")
    captured: dict[str, object] = {}

    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: SimpleNamespace(
            ready=True,
            message="ready",
            gog_path=Path(r"C:\gog.exe"),
            account_email="adel.belghali@gmail.com",
            accounts=("adel.belghali@gmail.com",),
        ),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QMessageBox.question",
        lambda *args, **kwargs: 16384,
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QMessageBox.critical",
        lambda *args, **kwargs: captured.setdefault(
            "critical",
            args[2] if len(args) > 2 else kwargs.get("text"),
        ),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.build_honorarios_gmail_request",
        lambda **kwargs: (_ for _ in ()).throw(AssertionError("draft request should not be built for contaminated translation artifacts")),
    )

    fake = SimpleNamespace(
        _gui_settings={"gmail_gog_path": "", "gmail_account_email": ""},
        court_email_combo=SimpleNamespace(currentText=lambda: "beja.judicial@tribunais.org.pt"),
        case_number_edit=SimpleNamespace(text=lambda: "21/25.0FBPTM"),
        _current_translation_docx_path=lambda: translated.resolve(),
    )

    QtSaveToJobLogDialog._offer_gmail_draft_for_honorarios(fake, honorarios.resolve(), default_primary_profile())

    assert "contaminated with honorários content" in captured["critical"]
    assert str(translated.resolve()) in captured["critical"]


def test_joblog_gmail_draft_blocks_contaminated_historical_translation_docx(
    tmp_path: Path,
    monkeypatch,
) -> None:
    translated = tmp_path / "translated.docx"
    honorarios = tmp_path / "honorarios.docx"
    _write_docx_with_paragraphs(
        translated,
        "Venho por este meio requerer o pagamento dos honorários devidos.",
        "O documento traduzido contém 264 palavras.",
    )
    _write_docx_with_paragraphs(honorarios, "Requerimento de honorários distinto.")
    captured: dict[str, object] = {}

    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.assess_gmail_draft_prereqs",
        lambda **kwargs: SimpleNamespace(
            ready=True,
            message="ready",
            gog_path=Path(r"C:\gog.exe"),
            account_email="adel.belghali@gmail.com",
            accounts=("adel.belghali@gmail.com",),
        ),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QMessageBox.question",
        lambda *args, **kwargs: 16384,
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.QMessageBox.critical",
        lambda *args, **kwargs: captured.setdefault(
            "critical",
            args[2] if len(args) > 2 else kwargs.get("text"),
        ),
    )
    monkeypatch.setattr(
        "legalpdf_translate.qt_gui.dialogs.build_honorarios_gmail_request",
        lambda **kwargs: (_ for _ in ()).throw(AssertionError("draft request should not be built for contaminated historical artifacts")),
    )

    fake = SimpleNamespace(
        _gui_settings={"gmail_gog_path": "", "gmail_account_email": ""},
        _historical_translation_docx_path=lambda row, honorarios_docx=None: translated.resolve(),
        _persist_historical_translation_docx=lambda row, path: None,
    )
    row = {
        "court_email": "beja.judicial@tribunais.org.pt",
        "case_number": "21/25.0FBPTM",
    }

    QtJobLogWindow._offer_gmail_draft_for_honorarios(fake, row, honorarios.resolve(), default_primary_profile())

    assert "contaminated with honorários content" in captured["critical"]
    assert str(translated.resolve()) in captured["critical"]
