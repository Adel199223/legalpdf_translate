from __future__ import annotations

import os
import zipfile
from pathlib import Path
from xml.etree import ElementTree

import fitz
import pytest
from docx import Document

from legalpdf_translate import word_pdf_artifacts as artifacts


def _docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("Requerimento de honorários — João Guerreiro")
    document.save(path)
    return path


def _pdf(path: Path, text: str = "Honorarios canary", *, encrypted: bool = False) -> Path:
    with fitz.open() as document:
        document.new_page().insert_text((72, 72), text)
        options = (
            {"encryption": fitz.PDF_ENCRYPT_AES_256, "owner_pw": "owner", "user_pw": "reader"}
            if encrypted
            else {}
        )
        document.save(path, **options)
    return path


def _zero_page_pdf(path: Path) -> Path:
    data = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, content in enumerate(
        (b"<< /Type /Catalog /Pages 2 0 R >>", b"<< /Type /Pages /Count 0 /Kids [] >>"), 1
    ):
        offsets.append(len(data))
        data.extend(f"{index} 0 obj\n".encode() + content + b"\nendobj\n")
    xref_offset = len(data)
    data.extend(b"xref\n0 3\n0000000000 65535 f \n")
    for offset in offsets[1:]:
        data.extend(f"{offset:010d} 00000 n \n".encode())
    data.extend(
        f"trailer\n<< /Size 3 /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode()
    )
    path.write_bytes(data)
    return path


@pytest.mark.parametrize("name", ["request.docx", "João d'Ávila requerimento.docx"])
def test_staging_promotes_fresh_verified_pdf_and_keeps_source(tmp_path: Path, name: str) -> None:
    source = _docx(tmp_path / name)
    original = source.read_bytes()
    destination = _pdf(tmp_path / "honorários.pdf", "Previous PDF")
    old_pdf = destination.read_bytes()

    with artifacts.staged_pdf_export(source, destination, expected_text="New canary") as (copy, output):
        stage_directory = copy.parent
        assert stage_directory.parent == destination.parent
        assert copy != source and copy.read_bytes() == original
        assert output.parent == stage_directory and not output.exists()
        assert destination.read_bytes() == old_pdf
        _pdf(output, "New canary")
        fresh_pdf = output.read_bytes()

    assert source.read_bytes() == original
    assert destination.read_bytes() == fresh_pdf
    assert not stage_directory.exists()


def test_staging_can_create_new_destination(tmp_path: Path) -> None:
    source = _docx(tmp_path / "request.docx")
    destination = tmp_path / "request.pdf"
    with artifacts.staged_pdf_export(source, destination) as (_, output):
        _pdf(output)
    artifacts.validate_pdf(destination)


@pytest.mark.parametrize("failure", ["worker", "missing", "empty", "corrupt", "wrong_text"])
def test_failure_preserves_existing_pdf_and_docx(tmp_path: Path, failure: str) -> None:
    source = _docx(tmp_path / "request.docx")
    destination = _pdf(tmp_path / "request.pdf", "Previous PDF")
    source_before, pdf_before = source.read_bytes(), destination.read_bytes()
    with pytest.raises((RuntimeError, ValueError)):
        with artifacts.staged_pdf_export(source, destination, expected_text="New output") as (_, output):
            stage_directory = output.parent
            if failure == "worker":
                raise RuntimeError("Simulated worker failure")
            if failure == "empty":
                output.touch()
            elif failure == "corrupt":
                output.write_bytes(b"%PDF-1.7\nnot a valid PDF\n%%EOF")
            elif failure == "wrong_text":
                _pdf(output, "Other output")
    assert source.read_bytes() == source_before
    assert destination.read_bytes() == pdf_before
    assert not stage_directory.exists()


@pytest.mark.parametrize("target", ["source", "destination", "new_destination"])
def test_concurrent_document_change_prevents_publication(tmp_path: Path, target: str) -> None:
    source = _docx(tmp_path / "request.docx")
    destination = tmp_path / "request.pdf"
    if target != "new_destination":
        _pdf(destination, "Old PDF")
    with pytest.raises(ValueError, match="changed"):
        with artifacts.staged_pdf_export(source, destination) as (_, output):
            _pdf(output)
            if target == "source":
                source.write_bytes(source.read_bytes() + b"changed")
            else:
                destination.write_bytes(b"User changed the destination")
            expected_source = source.read_bytes()
            expected_output = destination.read_bytes()
    assert source.read_bytes() == expected_source
    assert destination.read_bytes() == expected_output


def test_atomic_publication_failure_preserves_old_pdf(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    source = _docx(tmp_path / "request.docx")
    destination = _pdf(tmp_path / "request.pdf", "Old PDF")
    before = destination.read_bytes()

    def locked(*args: object, **kwargs: object) -> None:
        raise PermissionError("Private user path must not appear in the message")

    monkeypatch.setattr(artifacts.os, "replace", locked)
    with pytest.raises(ValueError, match="publish") as failure:
        with artifacts.staged_pdf_export(source, destination) as (_, output):
            _pdf(output)
    assert "Private user path" not in str(failure.value)
    assert destination.read_bytes() == before


@pytest.mark.parametrize("kind", ["missing", "empty", "header_only", "corrupt", "truncated", "encrypted", "zero_pages", "repair"])
def test_invalid_pdfs_are_rejected(tmp_path: Path, kind: str) -> None:
    path = tmp_path / "sensitive-person.pdf"
    if kind == "empty":
        path.touch()
    elif kind == "header_only":
        path.write_bytes(b"%PDF-1.7")
    elif kind == "corrupt":
        path.write_bytes(b"%PDF-1.7\nsecret private document\n%%EOF")
    elif kind == "truncated":
        _pdf(path)
        path.write_bytes(path.read_bytes()[:-30])
    elif kind == "encrypted":
        _pdf(path, encrypted=True)
    elif kind == "zero_pages":
        _zero_page_pdf(path)
    elif kind == "repair":
        _pdf(path)
        data = path.read_bytes()
        start, _ = data.rsplit(b"startxref\n", 1)
        path.write_bytes(start + b"startxref\n0\n%%EOF\n")
    with pytest.raises(ValueError) as failure:
        artifacts.validate_pdf(path)
    assert path.name not in str(failure.value)
    assert "private document" not in str(failure.value)


def test_expected_pdf_text_can_span_line_breaks(tmp_path: Path) -> None:
    path = _pdf(tmp_path / "valid.pdf", "Honorarios\ncanary")
    artifacts.validate_pdf(path, expected_text="Honorarios canary")
    with pytest.raises(ValueError, match="expected text"):
        artifacts.validate_pdf(path, expected_text="Missing synthetic phrase")


def test_pdf_with_owner_only_encryption_is_rejected(tmp_path: Path) -> None:
    path = tmp_path / "encrypted.pdf"
    with fitz.open() as document:
        document.new_page().insert_text((72, 72), "Canary")
        document.save(path, encryption=fitz.PDF_ENCRYPT_AES_256, owner_pw="owner", user_pw="")
    with pytest.raises(ValueError, match="encrypted"):
        artifacts.validate_pdf(path)


def test_pdf_validates_every_page_and_sanitizes_page_reader_errors(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    path = _pdf(tmp_path / "request.pdf")
    loaded: list[int] = []
    closed: list[bool] = []

    class BrokenDocument:
        is_pdf = True
        is_repaired = False
        is_encrypted = False
        needs_pass = False
        metadata = {}
        page_count = 2

        def load_page(self, number: int) -> object:
            loaded.append(number)
            if number == 1:
                raise ValueError("Secret document text and private file path")
            return type("Page", (), {"rect": fitz.Rect(0, 0, 600, 800), "get_text": lambda self: "Canary"})()

        def close(self) -> None:
            closed.append(True)

    monkeypatch.setattr(fitz, "open", lambda **kwargs: BrokenDocument())
    with pytest.raises(ValueError, match="unreadable pages") as failure:
        artifacts.validate_pdf(path)
    assert loaded == [0, 1] and closed == [True]
    assert "Secret" not in str(failure.value)


@pytest.mark.parametrize("kind", ["invalid_suffix", "not_zip", "empty_zip", "macro", "encrypted", "zip_corruption"])
def test_only_valid_non_macro_docx_is_staged(tmp_path: Path, kind: str) -> None:
    source = tmp_path / ("request.docm" if kind == "invalid_suffix" else "request.docx")
    _docx(source)
    if kind == "not_zip":
        source.write_bytes(b"encrypted compound or plain text")
    elif kind == "empty_zip":
        with zipfile.ZipFile(source, "w"):
            pass
    elif kind == "macro":
        with zipfile.ZipFile(source, "a") as archive:
            archive.writestr("word/vbaProject.bin", b"not executed")
    elif kind == "encrypted":
        # An encrypted member flag must be rejected before copying to Word.
        content = bytearray(source.read_bytes())
        local = content.index(b"PK\x03\x04")
        central = content.index(b"PK\x01\x02")
        content[local + 6] |= 1
        content[central + 8] |= 1
        source.write_bytes(content)
    elif kind == "zip_corruption":
        content = bytearray(source.read_bytes())
        content[100] ^= 255
        source.write_bytes(content)
    with pytest.raises(ValueError, match="DOCX"):
        with artifacts.staged_pdf_export(source, tmp_path / "request.pdf"):
            pytest.fail("Invalid input reached worker")


@pytest.mark.parametrize("invalid", ["same_path", "wrong_output_suffix", "missing_directory", "destination_directory"])
def test_bad_output_paths_are_rejected(tmp_path: Path, invalid: str) -> None:
    source = _docx(tmp_path / "request.docx")
    destination = {
        "same_path": source,
        "wrong_output_suffix": tmp_path / "request.txt",
        "missing_directory": tmp_path / "missing" / "request.pdf",
        "destination_directory": tmp_path / "folder.pdf",
    }[invalid]
    if invalid == "destination_directory":
        destination.mkdir()
    with pytest.raises(ValueError):
        with artifacts.staged_pdf_export(source, destination):
            pytest.fail("Invalid output reached worker")


def test_input_output_hardlink_alias_is_rejected(tmp_path: Path) -> None:
    source = _docx(tmp_path / "request.docx")
    destination = tmp_path / "request.pdf"
    os.link(source, destination)
    with pytest.raises(ValueError, match="same file"):
        with artifacts.staged_pdf_export(source, destination):
            pytest.fail("Aliased input reached worker")


def test_renderer_cannot_rewrite_staged_docx_and_publish(tmp_path: Path) -> None:
    source = _docx(tmp_path / "request.docx")
    original = source.read_bytes()
    destination = _pdf(tmp_path / "request.pdf", "Old PDF")
    old_pdf = destination.read_bytes()
    with pytest.raises(ValueError, match="staged DOCX changed"):
        with artifacts.staged_pdf_export(source, destination) as (copy, output):
            _pdf(output)
            copy.write_bytes(b"Renderer must not rewrite this source")
    assert source.read_bytes() == original
    assert destination.read_bytes() == old_pdf


def test_candidate_change_after_validation_prevents_publication(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    source = _docx(tmp_path / "request.docx")
    destination = _pdf(tmp_path / "request.pdf", "Old PDF")
    old_pdf = destination.read_bytes()
    original_validator = artifacts.validate_pdf

    def tampered(path: Path, *, expected_text: str | None = None) -> None:
        original_validator(path, expected_text=expected_text)
        path.write_bytes(b"Changed after valid output")

    monkeypatch.setattr(artifacts, "validate_pdf", tampered)
    with pytest.raises(ValueError, match="staged PDF changed"):
        with artifacts.staged_pdf_export(source, destination) as (_, output):
            _pdf(output)
    assert destination.read_bytes() == old_pdf


@pytest.mark.parametrize("target", ["source", "destination", "candidate"])
def test_symbolic_link_files_are_rejected(tmp_path: Path, target: str) -> None:
    source = _docx(tmp_path / "request.docx")
    destination = tmp_path / "request.pdf"
    real_file = _pdf(tmp_path / "other.pdf")
    link = tmp_path / "link.docx" if target == "source" else destination
    if target != "candidate":
        try:
            link.symlink_to(source if target == "source" else real_file)
        except OSError:
            pytest.skip("Host does not permit synthetic symlinks")
        if target == "source":
            source = link
        with pytest.raises(ValueError, match="link"):
            with artifacts.staged_pdf_export(source, destination):
                pytest.fail("Linked target reached worker")
    else:
        with pytest.raises(ValueError, match="link"):
            with artifacts.staged_pdf_export(source, destination) as (_, candidate):
                try:
                    candidate.symlink_to(real_file)
                except OSError:
                    pytest.skip("Host does not permit synthetic symlinks")


@pytest.mark.skipif(os.name != "nt", reason="Windows alternate data stream regression")
@pytest.mark.parametrize("zone", ["", "[ZoneTransfer]\nZoneId=0\n", "[ZoneTransfer]\nZoneId=3\n", "[ZoneTransfer]\nZoneId=4\n"])
def test_marked_source_is_not_copied_without_origin_security(tmp_path: Path, zone: str) -> None:
    source = _docx(tmp_path / "request.docx")
    origin_stream = Path(str(source) + ":Zone.Identifier")
    origin_stream.write_text(zone, encoding="utf-8")
    original = source.read_bytes()
    with pytest.raises(ValueError, match="origin security"):
        with artifacts.staged_pdf_export(source, tmp_path / "request.pdf"):
            pytest.fail("Origin-marked DOCX reached a mark-free staging copy")
    assert source.read_bytes() == original
    assert origin_stream.read_text(encoding="utf-8") == zone
    assert not list(tmp_path.glob(".honorarios-pdf-*"))


@pytest.mark.skipif(os.name != "nt", reason="Windows alternate data stream regression")
def test_origin_mark_added_after_staging_prevents_publication(tmp_path: Path) -> None:
    source = _docx(tmp_path / "request.docx")
    destination = _pdf(tmp_path / "request.pdf", "Old PDF")
    old_pdf = destination.read_bytes()
    with pytest.raises(ValueError, match="origin security"):
        with artifacts.staged_pdf_export(source, destination) as (_, output):
            _pdf(output)
            Path(str(source) + ":Zone.Identifier").write_text("[ZoneTransfer]\nZoneId=3\n", encoding="utf-8")
    assert destination.read_bytes() == old_pdf


@pytest.mark.skipif(os.name != "nt", reason="Windows origin inspection regression")
def test_inaccessible_origin_stream_is_fail_closed(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    source = _docx(tmp_path / "request.docx")
    original_open = Path.open

    def inaccessible(path: Path, *args: object, **kwargs: object):
        if str(path).endswith(":Zone.Identifier"):
            raise PermissionError("Private stream metadata must not leak")
        return original_open(path, *args, **kwargs)

    monkeypatch.setattr(Path, "open", inaccessible)
    with pytest.raises(ValueError, match="origin security") as failure:
        with artifacts.staged_pdf_export(source, tmp_path / "request.pdf"):
            pytest.fail("Unchecked origin reached Word")
    assert "Private stream" not in str(failure.value)


def test_alternate_data_stream_paths_are_not_export_targets(tmp_path: Path) -> None:
    source = _docx(tmp_path / "request.docx")
    with pytest.raises(ValueError, match="alternate data streams"):
        with artifacts.staged_pdf_export(source, tmp_path / "host-file:request.pdf"):
            pytest.fail("An alternate data stream reached the exporter")


def _add_external_relationship(path: Path, kind: str, *, mode: str = "External") -> None:
    relationships = ElementTree.Element("Relationships", xmlns="http://schemas.openxmlformats.org/package/2006/relationships")
    ElementTree.SubElement(
        relationships, "Relationship", Id="syntheticExternal",
        Type=f"http://schemas.openxmlformats.org/officeDocument/2006/relationships/{kind}",
        Target="https://example.invalid/sensitive-target", TargetMode=mode,
    )
    with zipfile.ZipFile(path, "a") as package:
        package.writestr("word/_rels/synthetic.xml.rels", ElementTree.tostring(relationships))


@pytest.mark.parametrize("kind", ["attachedTemplate", "oleObject", "image", "aFChunk", "package", "unknown"])
def test_externally_loaded_relationships_are_rejected_without_rewriting(tmp_path: Path, kind: str) -> None:
    source = _docx(tmp_path / "request.docx")
    _add_external_relationship(source, kind)
    original = source.read_bytes()
    with pytest.raises(ValueError, match="external content") as failure:
        with artifacts.staged_pdf_export(source, tmp_path / "request.pdf"):
            pytest.fail("External content reached Word")
    assert source.read_bytes() == original
    assert "sensitive-target" not in str(failure.value)


def test_external_target_cannot_omit_external_mode(tmp_path: Path) -> None:
    source = _docx(tmp_path / "request.docx")
    _add_external_relationship(source, "image", mode="Internal")
    with pytest.raises(ValueError, match="external content"):
        with artifacts.staged_pdf_export(source, tmp_path / "request.pdf"):
            pytest.fail("External URL with incorrect mode reached Word")


def test_hyperlink_relationship_is_preserved_verbatim(tmp_path: Path) -> None:
    source = _docx(tmp_path / "request.docx")
    _add_external_relationship(source, "hyperlink")
    original = source.read_bytes()
    with artifacts.staged_pdf_export(source, tmp_path / "request.pdf") as (copy, output):
        assert copy.read_bytes() == original
        _pdf(output)
    assert source.read_bytes() == original


@pytest.mark.parametrize("field_code", ["INCLUDETEXT", "INCLUDEPICTURE", "DDE", "DDEAUTO", "LINK", "DATABASE", "RD"])
@pytest.mark.parametrize("simple", [True, False])
def test_external_field_instructions_are_rejected(tmp_path: Path, field_code: str, simple: bool) -> None:
    source = _docx(tmp_path / "request.docx")
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = ElementTree.Element(f"{{{namespace}}}hdr")
    paragraph = ElementTree.SubElement(root, f"{{{namespace}}}p")
    instruction = f'{field_code} "https://example.invalid/sensitive-target"'
    if simple:
        ElementTree.SubElement(paragraph, f"{{{namespace}}}fldSimple", {f"{{{namespace}}}instr": instruction})
    else:
        ElementTree.SubElement(paragraph, f"{{{namespace}}}fldChar", {f"{{{namespace}}}fldCharType": "begin"})
        # Word can split a field keyword across multiple runs.
        for part in (instruction[:3], instruction[3:]):
            ElementTree.SubElement(paragraph, f"{{{namespace}}}instrText").text = part
        ElementTree.SubElement(paragraph, f"{{{namespace}}}fldChar", {f"{{{namespace}}}fldCharType": "end"})
    with zipfile.ZipFile(source, "a") as package:
        package.writestr("word/synthetic-header.xml", ElementTree.tostring(root))
    original = source.read_bytes()
    with pytest.raises(ValueError, match="external content") as failure:
        with artifacts.staged_pdf_export(source, tmp_path / "request.pdf"):
            pytest.fail("External field reached Word")
    assert source.read_bytes() == original
    assert "sensitive-target" not in str(failure.value)
