"""Synthetic writer tests; never open Word or copy private reference content."""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
import pytest

from legalpdf_translate.document_structure import PageStructure, StructureBlock, text_sha256
from legalpdf_translate.document_layout import derive_page_layout, validate_page_layout
from legalpdf_translate.docx_writer import (
    _segment_rtl_placeholder_aware_runs,
    assemble_docx,
    sanitize_bidi_controls,
    unwrap_internal_placeholders,
)
from legalpdf_translate.types import TargetLang


def _save_page(folder: Path, blocks: list[dict], *, metadata=None, number=1):
    folder.mkdir(exist_ok=True)
    rows = [StructureBlock(id=f"p{number:04d}_b{index:04d}", **block)
            for index, block in enumerate(blocks, 1)]
    page = PageStructure(page_number=number, source_sha256=text_sha256("synthetic source"),
                         source_file_sha256="a" * 64, source_text_sha256=text_sha256("synthetic source"),
                         blocks=rows, metadata=metadata or {})
    page.translation_sha256 = text_sha256(page.text)
    path = folder / f"page_{number:04d}.txt"
    path.write_text(page.text, encoding="utf-8")
    path.with_suffix(".structure.json").write_text(json.dumps(page.to_dict(), ensure_ascii=False), encoding="utf-8")
    return path, page.to_dict()


def _visible(value):
    return sanitize_bidi_controls(unwrap_internal_placeholders(value))


@pytest.mark.parametrize("line", [
    "[[Rua Luís de Camões]] [[Nr.]]: [[17]]",
    "[[7000-123]] [[ÉVORA]]",
    "[[AB]] [[123456789]] [[PT]]",
    "[[123/26.4ABC]] [ [[9876543]] ]",
    " [[João Exemplo]] ",
    "Rua Luís de Camões Nr.: 17",
])
def test_latin_only_line_is_one_ltr_sequence_including_neutral_separators(line):
    runs, mixed = _segment_rtl_placeholder_aware_runs(line, strip_bidi_controls=True)
    assert runs == [("ltr", _visible(line))]
    assert mixed is False


@pytest.mark.parametrize("strip_controls", [False, True])
def test_legacy_multiline_placeholder_keeps_visible_text_and_no_internal_delimiters(strip_controls):
    text = "السيد \u2066[[João\nExemplo]]\u2069 حاضر."
    runs, mixed = _segment_rtl_placeholder_aware_runs(text, strip_bidi_controls=strip_controls)
    joined = "".join(chunk for _, chunk in runs)
    assert "[[" not in joined and "]]" not in joined
    assert sanitize_bidi_controls(joined) == _visible(text)
    assert mixed
    assert any(kind == "ltr" and "João\nExemplo" in chunk for kind, chunk in runs)


def test_multiline_arabic_address_keeps_each_whole_latin_line_and_exact_spacing(tmp_path):
    text = "السيد المحترم\n[[João Exemplo]]\n[[Rua Luís de Camões]] [[Nr.]]: [[17]]\n[[7000-123]] [[ÉVORA]]"
    pages = tmp_path / "pages"
    _save_page(pages, [{"text": text, "role": "address", "alignment": "left"}])
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    paragraph = Document(output).paragraphs[0]
    assert _visible(paragraph.text) == _visible(text)
    for expected in ("João Exemplo", "Rua Luís de Camões Nr.: 17", "7000-123 ÉVORA"):
        runs = [run for run in paragraph.runs if expected in _visible(run.text)]
        assert len(runs) == 1
        assert runs[0]._r.rPr.find(qn("w:rtl")).get(qn("w:val")) == "0"
    assert paragraph._p.pPr.jc.get(qn("w:val")) == "end"


def test_mixed_arabic_prose_keeps_boundary_spaces_outside_name_run(tmp_path):
    text = "حضر [[João Exemplo]] إلى المحكمة في [[10]] أيام."
    pages = tmp_path / "pages"
    _save_page(pages, [{"text": text}])
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    paragraph = Document(output).paragraphs[0]
    assert _visible(paragraph.text) == _visible(text)
    name = next(run for run in paragraph.runs if "João" in run.text)
    assert _visible(name.text) == "João Exemplo"
    assert name._r.rPr.find(qn("w:rtl")).get(qn("w:val")) == "0"
    assert len([run for run in paragraph.runs if "João" in run.text]) == 1


@pytest.mark.parametrize("role", ["heading", "header", "signature"])
@pytest.mark.parametrize("alignment,expected", [("left", "start"), ("right", "start"),
                                               ("center", "center"), ("justify", "both")])
def test_arabic_labels_follow_reading_direction_but_preserve_centers(tmp_path, role, alignment, expected):
    pages = tmp_path / "pages"
    _save_page(pages, [{"text": "نص عربي", "role": role, "alignment": alignment}])
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    assert Document(output).paragraphs[0]._p.pPr.jc.get(qn("w:val")) == expected


@pytest.mark.parametrize("role", ["heading", "header", "signature"])
def test_latin_only_labels_retain_explicit_physical_left(tmp_path, role):
    pages = tmp_path / "pages"
    _save_page(pages, [{"text": "[[João Exemplo]]", "role": role, "alignment": "left"}])
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    assert Document(output).paragraphs[0]._p.pPr.jc.get(qn("w:val")) == "end"


def test_invalid_optional_layout_retains_complete_structure_ids_and_text(tmp_path):
    pages = tmp_path / "pages"
    _, page = _save_page(pages, [{"text": "First retained obligation."}, {"text": "Second retained obligation."}],
                         metadata={"layout": {"version": 999}})
    stats = {}
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.EN, page_breaks=False, stats=stats)
    assert [paragraph.text for paragraph in Document(output).paragraphs] == [block["text"] for block in page["blocks"]]
    source_map = json.loads(output.with_suffix(".source_map.json").read_text())
    mapped = source_map["pages"][0]
    assert mapped["structure_status"] == "validated"
    assert mapped["source_block_coverage_status"] == "complete"
    assert mapped["layout_status"] == "invalid_layout_flow_fallback"
    assert mapped["layout_review_required"] is source_map["layout_review_required"] is True
    assert [block["block_id"] for block in mapped["blocks"]] == [block["id"] for block in page["blocks"]]
    assert stats["layout_review_page_count"] == 1


def _region_page(folder, *, shade=True, region_alignment=None, data_table=False):
    blocks = [
        {"text": "العنوان الأيسر", "role": "heading", "alignment": "left", "bbox": [25, 45, 225, 65]},
        {"text": "الالتزام في اللوحة [[10]] أيام", "bbox": [25, 75, 225, 100]},
        {"text": "العنوان الأيمن", "role": "heading", "alignment": "center", "bbox": [285, 45, 565, 65]},
        {"text": "نص مستقل في العمود الأيمن", "alignment": "left", "bbox": [285, 75, 565, 110]},
    ]
    if data_table:
        blocks[1:2] = [
            {"text": "CODE-A", "role": "table_cell", "table_id": "data", "row": 0, "col": 0,
             "alignment": "left", "bbox": [25, 75, 105, 90]},
            {"text": "DATE-B", "role": "table_cell", "table_id": "data", "row": 0, "col": 1,
             "alignment": "right", "bbox": [125, 75, 225, 90]},
            {"text": "VALUE-C", "role": "table_cell", "table_id": "data", "row": 1, "col": 0,
             "alignment": "left", "bbox": [25, 90, 105, 110]},
            {"text": "VALUE-D", "role": "table_cell", "table_id": "data", "row": 1, "col": 1,
             "alignment": "right", "bbox": [125, 90, 225, 110]},
        ]
    path, structure = _save_page(folder, blocks)
    image_bytes = None
    if shade:
        from PIL import Image, ImageDraw

        image = Image.new("RGB", (600, 842), "white")
        ImageDraw.Draw(image).rectangle((20, 40, 230, 115), fill="#E6E6E6")
        buffer = BytesIO()
        image.save(buffer, format="PNG")
        image_bytes = buffer.getvalue()
    layout = derive_page_layout(structure, image_bytes=image_bytes)
    assert layout["status"] == "regions"
    if region_alignment:
        for region in layout["bands"][0]["regions"]:
            region["alignment"] = region_alignment
    structure["metadata"]["layout"] = validate_page_layout(layout, structure)
    path.with_suffix(".structure.json").write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
    return path, structure


def _body_text(document):
    return [_visible("".join(node.text or "" for node in p.iter(qn("w:t"))))
            for p in document._element.body.iter(qn("w:p"))
            if "".join(node.text or "" for node in p.iter(qn("w:t"))).strip()]


def test_region_columns_and_shaded_panel_are_editable_physical_not_mirrored(tmp_path):
    pages = tmp_path / "pages"
    _, structure = _region_page(pages)
    original = (pages / "page_0001.structure.json").read_bytes()
    stats = {}
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False, stats=stats)
    document = Document(output)
    assert len(document.tables) == 1
    outer = document.tables[0]
    assert outer._tbl.tblPr.find(qn("w:bidiVisual")).get(qn("w:val")) == "0"
    assert outer.autofit is False
    assert len(outer.columns) == 2
    assert outer.columns[1].width > outer.columns[0].width
    assert len(outer.cell(0, 0).tables) == 1
    panel = outer.cell(0, 0).tables[0]
    assert panel.cell(0, 0)._tc.tcPr.find(qn("w:shd")).get(qn("w:fill")) == "E6E6E6"
    assert all(child.get(qn("w:w")) == "100" for child in panel.cell(0, 0)._tc.tcPr.find(qn("w:tcMar")))
    assert _body_text(document) == [_visible(block["text"]) for block in structure["blocks"]]
    assert not document._element.body.xpath(".//w:cantSplit|.//w:trHeight")
    for table in (outer, panel):
        assert all(edge.get(qn("w:val")) == "nil" for edge in table._tbl.tblPr.find(qn("w:tblBorders")))
    for run in document._element.body.iter(qn("w:r")):
        assert run.find("./" + qn("w:rPr") + "/" + qn("w:szCs")).get(qn("w:val")) == "22"
    mapped = json.loads(output.with_suffix(".source_map.json").read_text())["pages"][0]
    assert mapped["layout_status"] == "regions"
    assert mapped["source_block_coverage_status"] == "complete"
    assert [row["block_id"] for row in mapped["blocks"]] == [row["id"] for row in structure["blocks"]]
    assert len({row["block_id"] for row in mapped["blocks"]}) == 4
    assert all(row["location"]["layout_region_id"] for row in mapped["blocks"])
    assert stats["layout_region_page_count"] == 1
    assert (pages / "page_0001.structure.json").read_bytes() == original


def test_regions_without_pixel_style_evidence_render_columns_but_require_review(tmp_path):
    pages = tmp_path / "pages"
    _region_page(pages, shade=False)
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    document = Document(output)
    assert len(document.tables) == 1
    assert not list(document._element.body.iter(qn("w:shd")))
    mapped = json.loads(output.with_suffix(".source_map.json").read_text())["pages"][0]
    assert mapped["layout_status"] == "regions" and mapped["layout_review_required"]
    assert mapped["layout_warnings"] == ["panel_style_not_evaluated"]


def test_explicit_physical_region_alignment_and_block_center_are_preserved(tmp_path):
    pages = tmp_path / "pages"
    _region_page(pages, shade=False, region_alignment="left")
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    outer = Document(output).tables[0]
    assert [p._p.pPr.jc.get(qn("w:val")) for p in outer.cell(0, 0).paragraphs] == ["end", "end"]
    assert [p._p.pPr.jc.get(qn("w:val")) for p in outer.cell(0, 1).paragraphs] == ["center", "end"]


def test_semantic_table_remains_independent_inside_physical_region(tmp_path):
    pages = tmp_path / "pages"
    _, structure = _region_page(pages, shade=False, data_table=True)
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    document = Document(output)
    outer = document.tables[0]
    data = outer.cell(0, 0).tables[0]
    assert outer._tbl.tblPr.find(qn("w:bidiVisual")).get(qn("w:val")) == "0"
    assert data._tbl.tblPr.find(qn("w:bidiVisual")).get(qn("w:val")) == "1"
    assert [[_visible(c.text) for c in row.cells] for row in data.rows] == [["CODE-A", "DATE-B"], ["VALUE-C", "VALUE-D"]]
    assert data.cell(0, 0).paragraphs[0]._p.pPr.jc.get(qn("w:val")) == "end"
    assert data.cell(0, 1).paragraphs[0]._p.pPr.jc.get(qn("w:val")) == "start"
    assert _body_text(document) == [_visible(block["text"]) for block in structure["blocks"]]
    mapped = json.loads(output.with_suffix(".source_map.json").read_text())["pages"][0]["blocks"]
    assert len(mapped) == len(structure["blocks"])
    assert all(row["location"]["kind"] == "table_cell" for row in mapped[1:5])


@pytest.mark.parametrize("invalid", [False, True])
def test_format_only_rebuild_derivative_does_not_rewrite_translation_evidence(tmp_path, invalid):
    from legalpdf_translate.layout_integration import LAYOUT_DERIVATION_VERSION

    pages = tmp_path / "pages"
    path, structure = _region_page(pages, shade=False)
    layout = structure["metadata"].pop("layout")
    path.with_suffix(".structure.json").write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
    before = (path.read_bytes(), path.with_suffix(".structure.json").read_bytes())
    derivative = {"version": 1, "derivation_version": LAYOUT_DERIVATION_VERSION,
                  "translation_sha256": "0" * 64 if invalid else structure["translation_sha256"],
                  "source_file_sha256": structure["source_file_sha256"], "layout": layout}
    path.with_suffix(".layout.json").write_text(json.dumps(derivative), encoding="utf-8")
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    document = Document(output)
    assert _body_text(document) == [_visible(block["text"]) for block in structure["blocks"]]
    assert len(document.tables) == (0 if invalid else 1)
    mapped = json.loads(output.with_suffix(".source_map.json").read_text())["pages"][0]
    assert mapped["structure_status"] == "validated" and mapped["source_block_coverage_status"] == "complete"
    assert mapped["layout_status"] == ("needs_review_flow_fallback" if invalid else "regions")
    if invalid:
        assert "invalid_layout_derivative" in mapped["layout_warnings"]
    assert before == (path.read_bytes(), path.with_suffix(".structure.json").read_bytes())


def test_full_width_bands_and_physical_columns_reflow_without_fixed_heights(tmp_path):
    pages = tmp_path / "pages"
    path, structure = _save_page(pages, [
        {"text": "عنوان كامل العرض", "role": "heading", "alignment": "center", "bbox": [25, 20, 565, 40]},
        {"text": "رأس العمود الأيسر", "role": "header", "bbox": [25, 65, 225, 80]},
        {"text": "تفاصيل العمود الأيسر " * 60, "bbox": [25, 90, 225, 120]},
        {"text": "رأس العمود الأيمن", "role": "header", "bbox": [285, 65, 565, 80]},
        {"text": "تفاصيل العمود الأيمن", "bbox": [285, 90, 565, 120]},
        {"text": "التزام ختامي كامل العرض", "bbox": [25, 145, 565, 170]},
    ])
    layout = derive_page_layout(structure)
    assert layout["status"] == "regions" and len(layout["bands"]) == 3
    structure["metadata"]["layout"] = layout
    path.with_suffix(".structure.json").write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    document = Document(output)
    assert len(document.tables) == 3
    assert [len(table.rows[0]._tr.tc_lst) for table in document.tables] == [1, 2, 1]
    assert _body_text(document) == [_visible(block["text"]) for block in structure["blocks"]]
    assert not document._element.body.xpath(".//w:trHeight|.//w:cantSplit")
    mapped = json.loads(output.with_suffix(".source_map.json").read_text())["pages"][0]["blocks"]
    assert len(mapped) == len({row["block_id"] for row in mapped}) == 6


@pytest.mark.parametrize("page_breaks,document_start,expected_sections", [(False, False, 1), (True, False, 2),
                                                                      (False, True, 2), (True, True, 2)])
def test_layout_page_preserves_explicit_and_independent_document_boundaries(
        tmp_path, page_breaks, document_start, expected_sections):
    pages = tmp_path / "pages"
    _, first = _region_page(pages, shade=False)
    path, second = _save_page(pages, [{"text": "نص الوثيقة التالية", "document_start": document_start}], number=2)
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=page_breaks)
    document = Document(output)
    assert len(document.sections) == expected_sections
    assert _body_text(document) == [_visible(block["text"]) for block in first["blocks"] + second["blocks"]]
    source_map = json.loads(output.with_suffix(".source_map.json").read_text())
    assert source_map["source_page_count"] == 2 and source_map["rendered_page_count"] is None
    assert all(page["source_block_coverage_status"] == "complete" for page in source_map["pages"])
    assert all(section.left_margin.twips == 964 and section.right_margin.twips == 964
               and section.top_margin.twips == 850 and section.bottom_margin.twips == 850
               for section in document.sections)
    with ZipFile(output) as archive:
        assert " PAGE " in archive.read("word/footer1.xml").decode("utf-8")


def _signature_blocks():
    return [
        {"text": "موظفة المحكمة،", "role": "signature", "alignment": "left", "bbox": [28, 150, 120, 158]},
        {"text": "[[Maria Silva]]", "role": "signature", "alignment": "left", "bbox": [29, 170, 150, 179]},
        {"text": "تقنية العدل", "role": "signature", "alignment": "left", "bbox": [29, 184, 100, 193]},
    ]


def test_confirmed_mixed_script_signature_stays_one_right_aligned_group(tmp_path):
    pages = tmp_path / "pages"
    path, structure = _save_page(pages, _signature_blocks())
    before = (path.read_bytes(), path.with_suffix(".structure.json").read_bytes())
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    paragraphs = Document(output).paragraphs
    assert [p._p.pPr.jc.get(qn("w:val")) for p in paragraphs] == ["start"] * 3
    assert [p.paragraph_format.keep_with_next for p in paragraphs] == [True, True, False]
    assert [_visible(p.text) for p in paragraphs] == [_visible(block["text"]) for block in structure["blocks"]]
    assert len(paragraphs[1].runs) == 1
    assert paragraphs[1].runs[0]._r.rPr.find(qn("w:rtl")).get(qn("w:val")) == "0"
    assert all(run.font.size.pt == 11 for p in paragraphs for run in p.runs)
    assert before == (path.read_bytes(), path.with_suffix(".structure.json").read_bytes())
    mapped = json.loads(output.with_suffix(".source_map.json").read_text())["pages"][0]["blocks"]
    assert [block["block_id"] for block in mapped] == [block["id"] for block in structure["blocks"]]


@pytest.mark.parametrize("case", ["missing_geometry", "far_apart", "different_anchor", "uncertain_block",
                                  "uncertain_page", "document_boundary", "different_role", "single_latin"])
def test_signature_grouping_requires_positive_role_and_geometry_evidence(tmp_path, case):
    pages = tmp_path / "pages"
    blocks = _signature_blocks()[:2]
    if case == "missing_geometry":
        blocks[1].pop("bbox")
    elif case == "far_apart":
        blocks[1]["bbox"] = [29, 300, 150, 309]
    elif case == "different_anchor":
        blocks[1]["bbox"] = [250, 170, 370, 179]
    elif case == "uncertain_block":
        blocks[1]["uncertain"] = True
    elif case == "document_boundary":
        blocks[1]["document_start"] = True
    elif case == "different_role":
        blocks[0]["role"] = "paragraph"
    elif case == "single_latin":
        blocks = blocks[1:]
    path, structure = _save_page(pages, blocks)
    if case == "uncertain_page":
        structure["uncertain"] = True
        path.with_suffix(".structure.json").write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    paragraph = next(p for p in Document(output).paragraphs if "Maria Silva" in p.text)
    assert paragraph._p.pPr.jc.get(qn("w:val")) == "end"
    assert paragraph.paragraph_format.keep_with_next is None


@pytest.mark.parametrize("lang,alignment,expected", [(TargetLang.AR, "center", "center"),
                                                    (TargetLang.EN, "left", "left"),
                                                    (TargetLang.FR, "left", "left")])
def test_signature_grouping_preserves_centered_and_non_arabic_layout(tmp_path, lang, alignment, expected):
    pages = tmp_path / "pages"
    blocks = _signature_blocks()
    for block in blocks:
        block["alignment"] = alignment
    _save_page(pages, blocks)
    output = assemble_docx(pages, tmp_path / "out.docx", lang=lang, page_breaks=False)
    paragraphs = Document(output).paragraphs
    assert [p._p.pPr.jc.get(qn("w:val")) for p in paragraphs] == [expected] * 3
    assert all(p.paragraph_format.keep_with_next is None for p in paragraphs)


@pytest.mark.parametrize("physical_alignment,expected,keep", [(None, "start", True), ("left", "end", None)])
def test_signature_grouping_is_region_local_and_respects_explicit_physical_alignment(
        tmp_path, physical_alignment, expected, keep):
    pages = tmp_path / "pages"
    blocks = _signature_blocks() + [
        {"text": "رأس العمود المقابل", "role": "header", "bbox": [285, 150, 565, 164]},
        {"text": "نص مستقل في العمود المقابل", "bbox": [285, 174, 565, 194]},
    ]
    path, structure = _save_page(pages, blocks)
    layout = derive_page_layout(structure)
    assert layout["status"] == "regions"
    for region in layout["bands"][0]["regions"]:
        region["alignment"] = physical_alignment
    structure["metadata"]["layout"] = validate_page_layout(layout, structure)
    path.with_suffix(".structure.json").write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    paragraphs = Document(output).tables[0].cell(0, 0).paragraphs
    assert [p._p.pPr.jc.get(qn("w:val")) for p in paragraphs] == [expected] * 3
    assert paragraphs[1].paragraph_format.keep_with_next is keep
    assert [_visible(p.text) for p in paragraphs] == [_visible(block["text"]) for block in blocks[:3]]


@pytest.mark.parametrize("panel", [False, True])
@pytest.mark.parametrize("physical_alignment", [None, "left", "center", "right"])
def test_arabic_literal_list_markers_cannot_hang_outside_layout_cell(tmp_path, panel, physical_alignment):
    pages = tmp_path / "pages"
    path, structure = _region_page(pages, shade=panel, region_alignment=physical_alignment)
    for block, marker in zip(structure["blocks"], ["•", "1.", "•", "a)"]):
        block["role"] = "list_item"
        block["text"] = f"{marker} " + block["text"]
    # Regenerate geometry binding after role changes; retain pixel-proven panel
    # annotations solely for this synthetic image-backed fixture.
    old_layout = structure["metadata"].pop("layout")
    layout = derive_page_layout(structure)
    for region, old_region in zip(layout["bands"][0]["regions"], old_layout["bands"][0]["regions"]):
        region["panels"] = old_region["panels"]
        region["alignment"] = physical_alignment
    structure["metadata"]["layout"] = validate_page_layout(layout, structure)
    text = "\n".join(block["text"] for block in structure["blocks"])
    structure["translation_sha256"] = text_sha256(text)
    path.write_text(text, encoding="utf-8")
    path.with_suffix(".structure.json").write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    document = Document(output)
    paragraphs = [p for p in document._element.body.iter(qn("w:p")) if list(p.iter(qn("w:t")))]
    assert len(paragraphs) == 4
    assert _body_text(document) == [_visible(block["text"]) for block in structure["blocks"]]
    for paragraph in paragraphs:
        ind = paragraph.find("./" + qn("w:pPr") + "/" + qn("w:ind"))
        assert {key: ind.get(qn(f"w:{key}")) for key in ("left", "right", "hanging")} == {
            "left": "0", "right": "0", "hanging": "0"}
        assert ind.get(qn("w:firstLine")) is None
        assert paragraph.find("./" + qn("w:pPr") + "/" + qn("w:numPr")) is None
        assert all(run.find("./" + qn("w:rPr") + "/" + qn("w:szCs")).get(qn("w:val")) == "22"
                   for run in paragraph.iter(qn("w:r")))
    # The global style is deliberately unchanged; it is used by full-width lists.
    assert document.styles["LegalPDF List"].element.pPr.ind.get(qn("w:hanging")) == "198"


def test_layout_list_indent_override_does_not_change_full_width_arabic_lists(tmp_path):
    pages = tmp_path / "pages"
    text = "• قائمة عربية خارج خلايا التخطيط"
    _save_page(pages, [{"role": "list_item", "text": text, "alignment": "left"}])
    output = assemble_docx(pages, tmp_path / "out.docx", lang=TargetLang.AR, page_breaks=False)
    paragraph = Document(output).paragraphs[0]
    assert paragraph._p.pPr.find(qn("w:ind")) is None
    assert paragraph.style.element.pPr.ind.get(qn("w:right")) == "198"
    assert paragraph.style.element.pPr.ind.get(qn("w:hanging")) == "198"
    assert _visible(paragraph.text) == text


@pytest.mark.parametrize("lang", [TargetLang.EN, TargetLang.FR])
def test_layout_list_indent_override_does_not_change_non_arabic_lists(tmp_path, lang):
    pages = tmp_path / "pages"
    path, structure = _region_page(pages, shade=False)
    for block in structure["blocks"]:
        block["role"] = "list_item"
        block["text"] = "• Literal item preserved."
    structure["metadata"]["layout"] = derive_page_layout(structure)
    text = "\n".join(block["text"] for block in structure["blocks"])
    structure["translation_sha256"] = text_sha256(text)
    path.write_text(text, encoding="utf-8")
    path.with_suffix(".structure.json").write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
    output = assemble_docx(pages, tmp_path / "out.docx", lang=lang, page_breaks=False)
    document = Document(output)
    paragraphs = [p for p in document._element.body.iter(qn("w:p")) if list(p.iter(qn("w:t")))]
    assert len(paragraphs) == 4
    assert all(p.find("./" + qn("w:pPr") + "/" + qn("w:ind")) is None for p in paragraphs)
    assert document.styles["LegalPDF List"].element.pPr.ind.get(qn("w:left")) == "198"
