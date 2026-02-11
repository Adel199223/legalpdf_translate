from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

import legalpdf_translate.output_paths as output_paths_module
import legalpdf_translate.workflow as workflow_module
from legalpdf_translate.checkpoint import (
    build_run_paths,
    ensure_run_dirs,
    load_run_state,
    mark_page_done,
    new_run_state,
    save_run_state_atomic,
    sha256_of_text,
)
from legalpdf_translate.docx_writer import assemble_docx
from legalpdf_translate.output_paths import (
    build_output_paths,
    require_writable_output_dir,
    require_writable_output_dir_text,
)
from legalpdf_translate.types import ImageMode, PageStatus, ReasoningEffort, RunConfig, TargetLang
from legalpdf_translate.workflow import TranslationWorkflow, _PageOutcome


def _make_config(pdf: Path, outdir: Path) -> RunConfig:
    return RunConfig(
        pdf_path=pdf,
        output_dir=outdir,
        target_lang=TargetLang.EN,
        effort=ReasoningEffort.HIGH,
        image_mode=ImageMode.AUTO,
        max_pages=1,
        resume=False,
        page_breaks=True,
        keep_intermediates=True,
    )


def _write_page(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def test_outdir_required_and_writable(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    with pytest.raises(ValueError):
        require_writable_output_dir_text("")

    with pytest.raises(ValueError):
        require_writable_output_dir(tmp_path / "missing")

    unwritable = tmp_path / "unwritable"
    unwritable.mkdir()

    def _raise_probe(_path: Path) -> None:
        raise OSError("no write")

    monkeypatch.setattr(output_paths_module, "_write_probe_file", _raise_probe)
    with pytest.raises(ValueError):
        require_writable_output_dir(unwritable)


def test_paths_frozen(tmp_path: Path) -> None:
    outdir_initial = tmp_path / "out_a"
    outdir_changed = tmp_path / "out_b"
    outdir_initial.mkdir()
    outdir_changed.mkdir()
    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(b"%PDF-1.4")

    frozen_paths = build_output_paths(
        outdir_initial,
        pdf,
        TargetLang.EN,
        run_started_at="20260211_030303",
    )

    assert frozen_paths.frozen_outdir == outdir_initial.resolve()
    assert frozen_paths.final_docx_path.parent == outdir_initial.resolve()
    assert frozen_paths.final_docx_path.parent != outdir_changed.resolve()
    assert frozen_paths.final_docx_path.name == "sample_EN_20260211_030303.docx"


def test_atomic_docx_write_success(tmp_path: Path) -> None:
    pages_dir = tmp_path / "pages"
    _write_page(pages_dir / "page_0001.txt", "Line one\nLine two")
    final_docx = tmp_path / "final.docx"

    assemble_docx(pages_dir, final_docx, lang=TargetLang.EN, page_breaks=True)

    assert final_docx.exists()
    assert final_docx.stat().st_size > 0
    assert not (tmp_path / "final.docx.tmp").exists()
    Document(final_docx)


def test_no_completed_without_file(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(b"not a real pdf")
    outdir = tmp_path / "out"
    outdir.mkdir()
    config = _make_config(pdf, outdir)

    def _fake_process_page(self, *, client, config, paths, instructions, context_text, page_number, total_pages):  # type: ignore[no-untyped-def]
        _write_page(paths.pages_dir / "page_0001.txt", "ready")
        return _PageOutcome(
            status=PageStatus.DONE,
            image_used=False,
            retry_used=False,
            usage={},
            error=None,
        )

    def _fail_docx(*args, **kwargs):  # type: ignore[no-untyped-def]
        raise RuntimeError("save failed")

    monkeypatch.setattr(workflow_module, "load_environment", lambda: None)
    monkeypatch.setattr(workflow_module, "get_page_count", lambda _pdf: 1)
    monkeypatch.setattr(TranslationWorkflow, "_process_page", _fake_process_page)
    monkeypatch.setattr(workflow_module, "assemble_docx", _fail_docx)

    workflow = TranslationWorkflow(client=object())
    summary = workflow.run(config)

    assert summary.success is False
    assert summary.error == "docx_write_failed"
    assert summary.output_docx is None
    assert summary.attempted_output_docx is not None
    assert summary.attempted_output_docx.parent == outdir.resolve()
    assert summary.attempted_output_docx.exists() is False


def test_rebuild_docx_from_pages(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    outdir = tmp_path / "out"
    outdir.mkdir()
    config = _make_config(pdf, outdir)

    paths = build_run_paths(
        config.output_dir,
        config.pdf_path,
        config.target_lang,
        run_started_at="20260211_040404",
    )
    ensure_run_dirs(paths)
    _write_page(paths.pages_dir / "page_0001.txt", "First page")
    _write_page(paths.pages_dir / "page_0002.txt", "Second page")

    state = new_run_state(
        config=config,
        paths=paths,
        pdf_fingerprint="pdfhash",
        context_hash=sha256_of_text(None),
        total_pages=2,
        selected_pages=[1, 2],
    )
    mark_page_done(state, 1, image_used=False, retry_used=False, usage={})
    mark_page_done(state, 2, image_used=False, retry_used=False, usage={})
    save_run_state_atomic(paths.run_state_path, state)

    def _forbid_api_client(*args, **kwargs):  # type: ignore[no-untyped-def]
        raise AssertionError("API client must not be used in rebuild mode")

    monkeypatch.setattr(workflow_module, "OpenAIResponsesClient", _forbid_api_client)

    workflow = TranslationWorkflow()
    rebuilt_path = workflow.rebuild_docx(config)

    assert rebuilt_path.exists()
    assert rebuilt_path.stat().st_size > 0
    assert rebuilt_path.name == "sample_EN_20260211_040404.docx"

    doc = Document(rebuilt_path)
    assert [p.text for p in doc.paragraphs] == ["First page", "Second page"]

    loaded_state = load_run_state(paths.run_state_path)
    assert loaded_state is not None
    assert loaded_state.final_docx_path_abs == str(rebuilt_path.resolve())


def test_workflow_uses_real_pdf_page_numbers_for_selection(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    outdir = tmp_path / "out"
    outdir.mkdir()
    config = RunConfig(
        pdf_path=pdf,
        output_dir=outdir,
        target_lang=TargetLang.EN,
        effort=ReasoningEffort.HIGH,
        image_mode=ImageMode.AUTO,
        start_page=3,
        end_page=5,
        max_pages=2,
        resume=False,
        page_breaks=True,
        keep_intermediates=True,
    )

    called_pages: list[int] = []

    def _fake_process_page(self, *, client, config, paths, instructions, context_text, page_number, total_pages):  # type: ignore[no-untyped-def]
        called_pages.append(page_number)
        _write_page(paths.pages_dir / f"page_{page_number:04d}.txt", f"Page {page_number}")
        return _PageOutcome(
            status=PageStatus.DONE,
            image_used=False,
            retry_used=False,
            usage={},
            error=None,
        )

    monkeypatch.setattr(workflow_module, "load_environment", lambda: None)
    monkeypatch.setattr(workflow_module, "get_page_count", lambda _pdf: 6)
    monkeypatch.setattr(TranslationWorkflow, "_process_page", _fake_process_page)

    workflow = TranslationWorkflow(client=object())
    summary = workflow.run(config)

    assert summary.success is True
    assert called_pages == [3, 4]
    paths = build_run_paths(config.output_dir, config.pdf_path, config.target_lang)
    page_files = sorted(path.name for path in paths.pages_dir.glob("page_*.txt"))
    assert page_files == ["page_0003.txt", "page_0004.txt"]


def test_resume_mismatch_raises_explicit_error(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    outdir = tmp_path / "out"
    outdir.mkdir()

    checkpoint_config = RunConfig(
        pdf_path=pdf,
        output_dir=outdir,
        target_lang=TargetLang.EN,
        effort=ReasoningEffort.HIGH,
        image_mode=ImageMode.AUTO,
        start_page=1,
        end_page=3,
        max_pages=3,
        resume=True,
        page_breaks=True,
        keep_intermediates=True,
    )
    paths = build_run_paths(checkpoint_config.output_dir, checkpoint_config.pdf_path, checkpoint_config.target_lang)
    ensure_run_dirs(paths)
    state = new_run_state(
        config=checkpoint_config,
        paths=paths,
        pdf_fingerprint=sha256_of_text("pdf"),
        context_hash=sha256_of_text(None),
        total_pages=6,
        selected_pages=[1, 2, 3],
    )
    save_run_state_atomic(paths.run_state_path, state)

    run_config = RunConfig(
        pdf_path=pdf,
        output_dir=outdir,
        target_lang=TargetLang.EN,
        effort=ReasoningEffort.HIGH,
        image_mode=ImageMode.AUTO,
        start_page=2,
        end_page=4,
        max_pages=3,
        resume=True,
        page_breaks=True,
        keep_intermediates=True,
    )

    monkeypatch.setattr(workflow_module, "load_environment", lambda: None)
    monkeypatch.setattr(workflow_module, "get_page_count", lambda _pdf: 6)
    monkeypatch.setattr(workflow_module, "sha256_of_file", lambda _pdf: sha256_of_text("pdf"))

    workflow = TranslationWorkflow(client=object())
    with pytest.raises(ValueError, match="Checkpoint is incompatible"):
        workflow.run(run_config)
